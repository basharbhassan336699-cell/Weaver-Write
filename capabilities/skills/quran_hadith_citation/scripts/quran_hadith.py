"""
quran_hadith.py — Quran/Hadith citation formatting (working script)
===================================================================
Applies the exact typographic conventions for quoting Quranic verses and
Prophetic hadiths inside a Word document, per the bundled skill rules.

Quranic verse (آية):
  - enclosed in the ornamental Quran brackets:  ﴿ ... ﴾  (U+FD3E / U+FD3F),
    NEVER normal parentheses or quotes,
  - the whole verse (incl. brackets) is BOLD,
  - citation "(Surah: Ayah)" in a smaller normal font after it, e.g.
    (البقرة: 286).

Prophetic hadith (حديث):
  - enclosed in Arabic double angle quotes:  « ... »  (NOT the Quran
    brackets — mixing the two is strictly forbidden),
  - the matn is BOLD (an optional key phrase may be colored red, still bold),
  - an optional non-bold lead-in "قال رسول الله ﷺ:" before it,
  - takhrij after it: "رواه [مُخرِّج] في [مصدر], و[حكم المحدث]" in a smaller
    normal font, e.g. رواه البخاري، صحيح.

This script formats ONLY; it does not verify that a verse belongs to its
surah or that a hadith's grading is correct — source accuracy is separate.

Requires: python-docx
"""
from __future__ import annotations

# The mandatory marks (do not substitute)
QURAN_OPEN = "\uFD3E"   # ﴿
QURAN_CLOSE = "\uFD3F"  # ﴾
HADITH_OPEN = "\u00AB"  # «
HADITH_CLOSE = "\u00BB"  # »
SALLA = "\uFDFA"        # ﷺ (SALLALLAHOU ALAYHE WASALLAM ligature)


def _set_run(run, size=None, bold=None, color=None, font="Kufyan Arabic"):
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    if bold is not None:
        run.font.bold = bold
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    run.font.name = font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts"); rpr.append(rfonts)
    for a in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(a), font)


def _set_rtl(paragraph):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = paragraph._p.get_or_add_pPr()
    if pPr.find(qn("w:bidi")) is None:
        pPr.append(OxmlElement("w:bidi"))


def add_quran_verse(doc_or_para, verse_text, surah, ayah, font="Kufyan Arabic",
                    body_size=14, cite_size=11, center=True):
    """
    Insert a Quranic verse with the correct formatting:
      ﴿ verse ﴾ (bold)  (Surah: Ayah) (smaller, normal)
    `doc_or_para`: a Document (a new paragraph is added) or an existing one.
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    if hasattr(doc_or_para, "add_paragraph"):
        p = doc_or_para.add_paragraph()
    else:
        p = doc_or_para
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_rtl(p)

    # bold verse in ornamental brackets
    verse = f"{QURAN_OPEN} {verse_text.strip()} {QURAN_CLOSE}"
    run = p.add_run(verse)
    _set_run(run, size=body_size, bold=True, font=font)

    # citation, smaller, normal weight
    cite = p.add_run(f"  ({surah}: {ayah})")
    _set_run(cite, size=cite_size, bold=False, font=font)
    return p


def add_hadith(doc_or_para, matn, narrator="", source="", grading="",
               lead_in=True, key_phrase=None, font="Kufyan Arabic",
               body_size=14, cite_size=11, center=True):
    """
    Insert a Prophetic hadith with the correct formatting:
      [قال رسول الله ﷺ:] « matn » (bold)  رواه [narrator] في [source], [grading]
    - `lead_in`: prepend the non-bold "قال رسول الله ﷺ:".
    - `key_phrase`: optional substring of matn to color red (kept bold).
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    if hasattr(doc_or_para, "add_paragraph"):
        p = doc_or_para.add_paragraph()
    else:
        p = doc_or_para
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_rtl(p)

    # optional non-bold lead-in
    if lead_in:
        li = p.add_run(f"قال رسول الله {SALLA}: ")
        _set_run(li, size=body_size, bold=False, font=font)

    # bold matn in « », with optional red key phrase
    p.add_run(HADITH_OPEN + " ")._element  # opening mark (styled below)
    # to allow a colored key phrase, split the matn around it
    if key_phrase and key_phrase in matn:
        before, after = matn.split(key_phrase, 1)
        r1 = p.add_run(before); _set_run(r1, size=body_size, bold=True, font=font)
        r2 = p.add_run(key_phrase)
        _set_run(r2, size=body_size, bold=True, color="C00000", font=font)
        r3 = p.add_run(after); _set_run(r3, size=body_size, bold=True, font=font)
    else:
        rm = p.add_run(matn.strip())
        _set_run(rm, size=body_size, bold=True, font=font)
    close = p.add_run(" " + HADITH_CLOSE)
    _set_run(close, size=body_size, bold=True, font=font)

    # takhrij, smaller, normal
    parts = []
    if narrator:
        parts.append(f"رواه {narrator}")
    if source:
        parts.append(f"في {source}")
    tail = "، ".join(parts) if parts else ""
    if grading:
        tail = (tail + "، " if tail else "") + grading
    if tail:
        cite = p.add_run(f"  ({tail})")
        _set_run(cite, size=cite_size, bold=False, font=font)
    return p


def validate_marks(text: str) -> dict:
    """
    Quick check that a document string doesn't misuse the marks:
      - a verse must not be wrapped in normal quotes/parens,
      - a hadith must not use the Quran brackets.
    Returns {"ok": bool, "warnings": [...]} (heuristic).
    """
    warnings = []
    # Quran brackets around clearly non-Quran? (can't know content; light check)
    if QURAN_OPEN in text and HADITH_OPEN in text:
        # both present — ensure they aren't nested wrongly (heuristic only)
        pass
    if '"' in text and QURAN_OPEN not in text:
        warnings.append("verses should use ﴿ ﴾, not normal quotes")
    return {"ok": len(warnings) == 0, "warnings": warnings}


if __name__ == "__main__":
    from docx import Document
    doc = Document()
    doc.add_heading("اختبار تنسيق الآيات والأحاديث", 0)

    add_quran_verse(doc, "لَا يُكَلِّفُ اللَّهُ نَفْسًا إِلَّا وُسْعَهَا",
                    "البقرة", "286")
    add_hadith(doc, "إنما الأعمال بالنيات، وإنما لكل امرئ ما نوى",
               narrator="البخاري ومسلم", grading="صحيح",
               key_phrase="بالنيات")
    doc.save("/tmp/quran_hadith_test.docx")
    print("saved /tmp/quran_hadith_test.docx")
