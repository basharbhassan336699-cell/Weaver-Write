"""
validate_docx.py — verify Word document integrity (working script)
==================================================================
Pure-logic validator. Opens a .docx and checks structural soundness
before delivery: it has content, headings exist, references section
present, and (for Arabic) RTL paragraphs are set.

Usage:
    python validate_docx.py --file research.docx --lang ar
"""
from __future__ import annotations
import argparse


def validate_docx(path: str, lang: str = "ar") -> dict:
    """Structural checks. Returns issues and an overall ok flag."""
    from docx import Document
    from docx.oxml.ns import qn

    issues, stats = [], {}
    doc = Document(path)

    paras = doc.paragraphs
    stats["paragraphs"] = len(paras)
    stats["non_empty"] = sum(1 for p in paras if p.text.strip())
    if stats["non_empty"] == 0:
        issues.append("document has no text content")

    # headings
    headings = [p for p in paras if p.style.name.startswith("Heading")]
    stats["headings"] = len(headings)
    if len(headings) == 0:
        issues.append("no headings found (document may be unstructured)")

    # references section
    ref_markers = ["المراجع", "references", "works cited", "bibliography"]
    has_refs = any(any(m in p.text.lower() for m in ref_markers) for p in paras)
    stats["has_references"] = has_refs
    if not has_refs:
        issues.append("no references section detected")

    # RTL check for Arabic
    if lang == "ar":
        rtl_count = 0
        for p in paras:
            pPr = p._p.find(qn("w:pPr"))
            if pPr is not None and pPr.find(qn("w:bidi")) is not None:
                rtl_count += 1
        stats["rtl_paragraphs"] = rtl_count
        if rtl_count == 0 and stats["non_empty"] > 0:
            issues.append("Arabic document but no RTL paragraphs set")

    # tables
    stats["tables"] = len(doc.tables)

    return {"ok": len(issues) == 0, "issues": issues, "stats": stats}


def _main():
    p = argparse.ArgumentParser(description="Validate a Word document")
    p.add_argument("--file", required=True)
    p.add_argument("--lang", default="ar", choices=["ar", "en"])
    args = p.parse_args()
    r = validate_docx(args.file, args.lang)
    print(f"Status: {'OK' if r['ok'] else 'issues found'}")
    for k, v in r["stats"].items():
        print(f"  {k}: {v}")
    for i in r["issues"]:
        print(f"  ! {i}")


if __name__ == "__main__":
    _main()
