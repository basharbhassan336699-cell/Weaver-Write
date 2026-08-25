"""
docx_frontmatter.py — cover page + table of contents (working script)
=====================================================================
Two front-matter features for academic Word documents, correct for both
Arabic (RTL) and English (LTR):

1) COVER PAGE (add_cover_page):
   A mandatory title page for any real task — university/institution line,
   the title (large), subtitle, author/student info, supervisor, course,
   and date — centered, themed, direction-correct. Skipped only when the
   caller says so (working inside the task's own file, writing a fragment of
   earlier work, or an explicit "no cover page" instruction).

2) TABLE OF CONTENTS (add_toc_page):
   A real Word TOC field (Word builds it on "update fields"). Placement:
   - "after_cover" (default when TOC requested but position unspecified):
     page 2, right after the cover.
   - "end": last page of the document.
   Fully bilingual heading and RTL/LTR direction.

Both reuse the shared theme palette so front matter matches the body and any
slides on the same topic.

Requires: python-docx
"""
from __future__ import annotations
import os
import json

from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

_THEMES = os.path.join(
    os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))),
    "pptx_builder", "themes", "themes.json")


def _palette(theme_id="academic_navy"):
    try:
        with open(_THEMES, encoding="utf-8") as f:
            t = json.load(f)["themes"].get(theme_id, {})
        return (t.get("primary", "1B2A4A"), t.get("accent", "C8A04A"),
                t.get("text", "222A38"))
    except Exception:
        return ("1B2A4A", "C8A04A", "222A38")


def _set_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    if pPr.find(qn("w:bidi")) is None:
        pPr.append(OxmlElement("w:bidi"))


def _line(doc, text, size, color, bold=False, rtl=False, font=None,
          space_after=12, italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    if font:
        run.font.name = font
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts"); rpr.append(rfonts)
        for a in ("w:ascii", "w:hAnsi", "w:cs"):
            rfonts.set(qn(a), font)
    if rtl:
        _set_rtl(p)
    return p


# ── cover page labels (bilingual) ────────────────────────────
_LABELS = {
    "ar": {"supervisor": "إشراف", "student": "إعداد", "course": "المقرر",
           "date": "التاريخ", "prepared_by": "إعداد الطالب"},
    "en": {"supervisor": "Supervised by", "student": "Prepared by",
           "course": "Course", "date": "Date", "prepared_by": "Prepared by"},
}


def add_cover_page(doc, title, lang="ar", theme_id="academic_navy", font=None,
                   institution="", subtitle="", author="", supervisor="",
                   course="", date="", logo_path=None, page_break_after=True):
    """
    Build a centered, themed cover page. Only `title` is required; other lines
    appear when provided. Direction follows `lang`.
    """
    primary, accent, text_col = _palette(theme_id)
    rtl = (lang == "ar")
    font = font or ("Kufyan Arabic" if rtl else "Times New Roman")
    L = _LABELS["ar" if rtl else "en"]

    # top spacing
    for _ in range(2):
        doc.add_paragraph()

    # institution
    if institution:
        _line(doc, institution, 18, primary, bold=True, rtl=rtl, font=font,
              space_after=6)
    # optional logo
    if logo_path and os.path.exists(logo_path):
        from docx.shared import Inches
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(logo_path, width=Inches(1.6))

    # spacer
    for _ in range(3):
        doc.add_paragraph()

    # title (large)
    _line(doc, title, 30, primary, bold=True, rtl=rtl, font=font, space_after=10)
    # accent rule under title
    rule = doc.add_paragraph(); rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = rule.add_run("────────"); r.font.color.rgb = RGBColor.from_string(accent)
    r.font.size = Pt(14)
    # subtitle
    if subtitle:
        _line(doc, subtitle, 18, text_col, italic=True, rtl=rtl, font=font,
              space_after=8)

    # spacer
    for _ in range(4):
        doc.add_paragraph()

    # author / supervisor / course / date block
    def _kv(label, value):
        if value:
            _line(doc, f"{label}: {value}", 15, text_col, bold=False, rtl=rtl,
                  font=font, space_after=6)

    _kv(L["student"], author)
    _kv(L["supervisor"], supervisor)
    _kv(L["course"], course)
    _kv(L["date"], date)

    if page_break_after:
        doc.add_page_break()
    return doc


# ── table of contents ────────────────────────────────────────
def add_toc_page(doc, lang="ar", theme_id="academic_navy", font=None,
                 page_break_after=True):
    """
    Insert a real Word TOC field with a themed heading. Word builds the actual
    entries on "update fields". Bilingual + direction-correct.
    """
    primary, _, _ = _palette(theme_id)
    rtl = (lang == "ar")
    font = font or ("Kufyan Arabic" if rtl else "Times New Roman")
    title = "المحتويات" if rtl else "Table of Contents"

    # heading
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT
    run = h.add_run(title)
    run.font.size = Pt(18); run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(primary)
    run.font.name = font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts"); rpr.append(rfonts)
    for a in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(a), font)
    if rtl:
        _set_rtl(h)

    # TOC field
    p = doc.add_paragraph()
    if rtl:
        _set_rtl(p)
    run = p.add_run()
    fldBegin = OxmlElement("w:fldChar"); fldBegin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldSep = OxmlElement("w:fldChar"); fldSep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = ("اضغط بزر الماوس الأيمن ثم (تحديث الحقل) لعرض المحتويات"
                        if rtl else
                        "Right-click and choose Update Field to build the TOC")
    fldEnd = OxmlElement("w:fldChar"); fldEnd.set(qn("w:fldCharType"), "end")
    for el in (fldBegin, instr, fldSep, placeholder, fldEnd):
        run._r.append(el)

    if page_break_after:
        doc.add_page_break()
    return doc


# ── policy helper: should a cover page be added? ─────────────
def should_add_cover(task_card: dict) -> bool:
    """
    Cover page is MANDATORY for any real task unless:
      - working inside the task's own uploaded file (in_place),
      - writing a fragment/continuation of earlier work (fragment),
      - an explicit 'no cover page' instruction (no_cover).
    If nothing is said, default to True (add it).
    """
    if task_card.get("no_cover") or task_card.get("in_place") \
            or task_card.get("fragment") or task_card.get("continuation"):
        return False
    return True


def resolve_toc_position(task_card: dict) -> str:
    """
    Where the TOC goes:
      - explicit 'end' -> last page,
      - explicit 'after_cover' -> page 2,
      - requested but unspecified -> 'after_cover' (page 2, the default).
    Returns 'after_cover' | 'end' | None (not requested).
    """
    if not task_card.get("toc"):
        return None
    pos = task_card.get("toc_position", "after_cover")
    return "end" if pos == "end" else "after_cover"


if __name__ == "__main__":
    from docx import Document
    # AR cover + TOC
    doc = Document()
    add_cover_page(doc, "معوقات التنمية في دول العالم الثالث", lang="ar",
                   institution="جامعة ليوا", subtitle="دراسة تحليلية",
                   author="الطالب", supervisor="أ. خالد", course="التنمية",
                   date="2026")
    add_toc_page(doc, lang="ar")
    doc.add_heading("المقدمة", 1)
    doc.add_paragraph("نص...")
    doc.save("/tmp/frontmatter_ar.docx")

    # EN cover + TOC
    doc2 = Document()
    add_cover_page(doc2, "Barriers to Development", lang="en",
                   institution="Liwa University", author="Student",
                   supervisor="Dr. Khaled", course="Development", date="2026")
    add_toc_page(doc2, lang="en")
    doc2.save("/tmp/frontmatter_en.docx")
    print("saved AR + EN front-matter samples")
