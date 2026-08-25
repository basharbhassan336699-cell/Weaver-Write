"""
docx_advanced.py — professional Word features (working script)
==============================================================
Rich, theme-aware Word building that matches the quality Claude produces:
formatted tables, inline images, two-column layouts, headers/footers with
page numbers, a table-of-contents field, and colored headings — all with
correct direction (Arabic RTL / English LTR) at the XML level.

These are composable helpers that operate on a python-docx Document, plus a
high-level build_rich_docx() that assembles a full themed document.

Direction:
  - Arabic (lang="ar") -> paragraphs get <w:bidi/>, right alignment, and
    tables get <w:bidiVisual/> so columns read right-to-left.
  - English (lang="en") -> normal LTR.

Themes: reuses the presentation palettes (themes.json) so a Word report and
a deck on the same topic share colors.

Requires: python-docx  (+ a bundled font via fonts-core)
"""
from __future__ import annotations
import os
import json

from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── theme palette (shared with slides) ───────────────────────
_THEMES = os.path.join(
    os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))),
    "pptx_builder", "themes", "themes.json")


def load_palette(theme_id="academic_navy"):
    try:
        with open(_THEMES, encoding="utf-8") as f:
            t = json.load(f)["themes"].get(theme_id, {})
        return {
            "primary": t.get("primary", "1B2A4A"),
            "accent": t.get("accent", "C8A04A"),
            "text": t.get("text", "222A38"),
        }
    except Exception:
        return {"primary": "1B2A4A", "accent": "C8A04A", "text": "222A38"}


# ── direction helpers ────────────────────────────────────────
def set_paragraph_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    if pPr.find(qn("w:bidi")) is None:
        pPr.append(OxmlElement("w:bidi"))


def set_table_rtl(table):
    """Make a table read right-to-left (columns flow RTL)."""
    tblPr = table._tbl.tblPr
    if tblPr.find(qn("w:bidiVisual")) is None:
        tblPr.append(OxmlElement("w:bidiVisual"))


def _shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_run_font(run, font_name, size=None, color=None, bold=None):
    from docx.shared import Pt, RGBColor
    if font_name:
        run.font.name = font_name
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts"); rpr.append(rfonts)
        for a in ("w:ascii", "w:hAnsi", "w:cs"):
            rfonts.set(qn(a), font_name)
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.font.bold = bold


# ── formatted table ──────────────────────────────────────────
def add_table(doc, headers, rows, lang="ar", theme_id="academic_navy",
              font=None, totals_row=None):
    """
    Add a styled table: colored header row, borders, RTL-aware.
    headers: list[str]; rows: list[list]; totals_row: optional list.
    """
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    pal = load_palette(theme_id)
    rtl = (lang == "ar")
    font = font or ("Kufyan Arabic" if rtl else "Times New Roman")

    n_cols = len(headers)
    table = doc.add_table(rows=1, cols=n_cols)
    table.style = "Table Grid"
    if rtl:
        set_table_rtl(table)

    # header
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        _shade_cell(hdr[i], pal["primary"])
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(h))
        _set_run_font(run, font, 12, "FFFFFF", bold=True)
        if rtl:
            set_paragraph_rtl(p)

    # body
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            _set_run_font(run, font, 11, pal["text"])
            if rtl:
                set_paragraph_rtl(p)

    # totals
    if totals_row:
        cells = table.add_row().cells
        for i, val in enumerate(totals_row):
            _shade_cell(cells[i], pal["accent"])
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            _set_run_font(run, font, 11, "1A1A1A", bold=True)
            if rtl:
                set_paragraph_rtl(p)
    return table


# ── inline image ─────────────────────────────────────────────
def add_image(doc, image_path, caption="", width_inches=5.5, lang="ar",
              font=None):
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    rtl = (lang == "ar")
    if not os.path.exists(image_path):
        return False
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(image_path, width=Inches(width_inches))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        _set_run_font(run, font or ("Kufyan Arabic" if rtl else "Times New Roman"),
                      10, "666666")
        run.font.italic = True
        if rtl:
            set_paragraph_rtl(cap)
    return True


# ── two-column section ───────────────────────────────────────
def set_columns(section, num=2, space_twips=425):
    """Set a section to N newspaper-style columns."""
    sectPr = section._sectPr
    cols = sectPr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols"); sectPr.append(cols)
    cols.set(qn("w:num"), str(num))
    cols.set(qn("w:space"), str(space_twips))


def add_column_break(doc):
    from docx.enum.text import WD_BREAK
    doc.add_paragraph().add_run().add_break(WD_BREAK.COLUMN)


# ── header / footer with page numbers ────────────────────────
def add_page_numbers(section, lang="ar", text=""):
    """Add a footer with a page-number field (centered)."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if text:
        p.add_run(text + "   ")
    # PAGE field
    run = p.add_run()
    fldBegin = OxmlElement("w:fldChar"); fldBegin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fldEnd = OxmlElement("w:fldChar"); fldEnd.set(qn("w:fldCharType"), "end")
    run._r.append(fldBegin); run._r.append(instr); run._r.append(fldEnd)
    if lang == "ar":
        set_paragraph_rtl(p)


def set_header(section, text, lang="ar", font=None, theme_id="academic_navy"):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    pal = load_palette(theme_id)
    header = section.header
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if lang == "ar" else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    _set_run_font(run, font or ("Kufyan Arabic" if lang == "ar" else "Times New Roman"),
                  10, pal["primary"], bold=True)
    if lang == "ar":
        set_paragraph_rtl(p)


# ── table of contents field ──────────────────────────────────
def add_toc(doc, lang="ar", font=None):
    """Insert a TOC field. Word populates it on open (update fields)."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    title = "المحتويات" if lang == "ar" else "Table of Contents"
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.RIGHT if lang == "ar" else WD_ALIGN_PARAGRAPH.LEFT
    run = h.add_run(title)
    _set_run_font(run, font or ("Kufyan Arabic" if lang == "ar" else "Times New Roman"),
                  16, None, bold=True)
    if lang == "ar":
        set_paragraph_rtl(h)

    p = doc.add_paragraph()
    run = p.add_run()
    fldBegin = OxmlElement("w:fldChar"); fldBegin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldSep = OxmlElement("w:fldChar"); fldSep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = ("اضغط تحديث الحقول لعرض المحتويات" if lang == "ar"
                        else "Right-click > Update Field to build the TOC")
    fldEnd = OxmlElement("w:fldChar"); fldEnd.set(qn("w:fldCharType"), "end")
    run._r.append(fldBegin); run._r.append(instr); run._r.append(fldSep)
    run._r.append(placeholder); run._r.append(fldEnd)
    if lang == "ar":
        set_paragraph_rtl(p)


# ── colored heading ──────────────────────────────────────────
def add_colored_heading(doc, text, level=1, lang="ar", theme_id="academic_navy",
                        font=None):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    pal = load_palette(theme_id)
    h = doc.add_heading("", level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.RIGHT if lang == "ar" else WD_ALIGN_PARAGRAPH.LEFT
    run = h.add_run(text)
    size = 18 if level == 1 else 15
    _set_run_font(run, font or ("Kufyan Arabic" if lang == "ar" else "Times New Roman"),
                  size, pal["primary"], bold=True)
    if lang == "ar":
        set_paragraph_rtl(h)
    return h


if __name__ == "__main__":
    # self-test: build a document exercising every feature
    from docx import Document
    doc = Document()
    sec = doc.sections[0]
    set_header(sec, "تقرير تجريبي", lang="ar")
    add_page_numbers(sec, lang="ar", text="صفحة")
    add_toc(doc, lang="ar")
    add_colored_heading(doc, "المقدمة", 1, lang="ar")
    doc.add_paragraph("نص تجريبي للفقرة.")
    add_table(doc, ["البند", "القيمة"], [["أ", 10], ["ب", 20]],
              lang="ar", totals_row=["الإجمالي", 30])
    doc.save("/tmp/advanced_test.docx")
    print("saved /tmp/advanced_test.docx")


def build_rich_docx(title, sections, output_path="research.docx", lang="ar",
                    theme_id="academic_navy", font=None, subtitle="",
                    references=None, header_text=None, page_numbers=True,
                    toc=False, two_columns=False, cover=None,
                    toc_position="after_cover"):
    """
    High-level rich document builder (Claude-quality).

    sections: list of dicts, each may contain:
        {"heading": str, "body": str,
         "table": {"headers":[...], "rows":[...], "totals":[...]}?,
         "image": {"path": str, "caption": str}?}
    Options: header_text, page_numbers, toc, two_columns.
    """
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    rtl = (lang == "ar")
    font = font or ("Kufyan Arabic" if rtl else "Times New Roman")
    doc = Document()
    sec = doc.sections[0]

    # ── cover page (mandatory unless suppressed) + TOC placement ──
    try:
        import sys, os
        here = os.path.dirname(os.path.abspath(__file__))
        if here not in sys.path:
            sys.path.insert(0, here)
        from docx_frontmatter import (add_cover_page, add_toc_page,
                                      should_add_cover, resolve_toc_position)
        _card = {"no_cover": (not cover) if cover is not None else False}
        if cover and should_add_cover(_card):
            cinfo = cover if isinstance(cover, dict) else {}
            add_cover_page(doc, title, lang=lang, theme_id=theme_id, font=font,
                           institution=cinfo.get("institution", ""),
                           subtitle=cinfo.get("subtitle", subtitle),
                           author=cinfo.get("author", ""),
                           supervisor=cinfo.get("supervisor", ""),
                           course=cinfo.get("course", ""),
                           date=cinfo.get("date", ""))
        _toc_pos = resolve_toc_position({"toc": toc,
                                         "toc_position": toc_position})
        if _toc_pos == "after_cover":
            add_toc_page(doc, lang=lang, theme_id=theme_id, font=font)
    except Exception:
        pass

    if header_text:
        set_header(sec, header_text, lang, font, theme_id)
    if page_numbers:
        add_page_numbers(sec, lang, "صفحة " if rtl else "Page ")

    # title
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    trun = tp.add_run(title)
    _set_run_font(trun, font, 24, load_palette(theme_id)["primary"], bold=True)
    if rtl:
        set_paragraph_rtl(tp)
    if subtitle:
        sp = doc.add_paragraph()
        sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        srun = sp.add_run(subtitle)
        _set_run_font(srun, font, 14, load_palette(theme_id)["accent"])
        if rtl:
            set_paragraph_rtl(sp)

    if toc:
        doc.add_page_break()
        add_toc(doc, lang, font)
        doc.add_page_break()

    # optionally switch body to two columns
    if two_columns:
        set_columns(sec, 2)

    # sections
    for s in sections:
        if s.get("heading"):
            add_colored_heading(doc, s["heading"], 1, lang, theme_id, font)
        if s.get("body"):
            bp = doc.add_paragraph()
            bp.alignment = WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT
            brun = bp.add_run(s["body"])
            _set_run_font(brun, font, 14, load_palette(theme_id)["text"])
            if rtl:
                set_paragraph_rtl(bp)
        if s.get("table"):
            t = s["table"]
            add_table(doc, t.get("headers", []), t.get("rows", []),
                      lang, theme_id, font, t.get("totals"))
        if s.get("image"):
            img = s["image"]
            add_image(doc, img.get("path", ""), img.get("caption", ""),
                      lang=lang, font=font)
        if s.get("equation"):
            # native Word equation (OMML)
            try:
                import sys, os
                here = os.path.dirname(os.path.abspath(__file__))
                if here not in sys.path:
                    sys.path.insert(0, here)
                from docx_math import add_equation
                eq = s["equation"]
                if isinstance(eq, str):
                    eq = [eq]
                for e in eq:
                    add_equation(doc, e, inline=False)
            except Exception:
                pass

    # references
    if references:
        add_colored_heading(doc, "المراجع" if rtl else "References", 1,
                            lang, theme_id, font)
        for ref in references:
            rp = doc.add_paragraph(style="List Number")
            rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT
            rrun = rp.add_run(ref)
            _set_run_font(rrun, font, 12, load_palette(theme_id)["text"])
            if rtl:
                set_paragraph_rtl(rp)

    # TOC at end, if requested
    try:
        if toc and toc_position == "end":
            from docx_frontmatter import add_toc_page
            doc.add_page_break()
            add_toc_page(doc, lang=lang, theme_id=theme_id, font=font,
                        page_break_after=False)
    except Exception:
        pass

    doc.save(output_path)
    return output_path
