"""
build_docx.py — build a professional Arabic RTL Word document (working script)
==============================================================================
Usage (as a module):
    from build_docx import build_academic_docx
    build_academic_docx(
        title="Research Title",
        sections=[{"heading": "Introduction", "body": "..."}],
        references=["Ref 1", "Ref 2"],
        output_path="research.docx",
        lang="ar",
    )

Requires: pip install python-docx

Note: output-facing strings (the "References" heading) switch between
Arabic and English based on `lang`, so the produced file reads naturally
in the task language.
"""
from __future__ import annotations
import argparse
import json


def _set_rtl(paragraph):
    """Set the paragraph direction to right-to-left."""
    from docx.oxml.ns import qn  # noqa
    from docx.oxml import OxmlElement
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    pPr.append(bidi)


def _set_font(paragraph, font_name, size=None):
    """Apply a font family (and optional size) to every run in a paragraph,
    setting the complex-script (cs) slot too so Arabic uses it."""
    from docx.shared import Pt
    from docx.oxml.ns import qn
    for run in paragraph.runs:
        run.font.name = font_name
        if size:
            run.font.size = Pt(size)
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = rpr.makeelement(qn("w:rFonts"), {})
            rpr.append(rfonts)
        for attr in ("w:ascii", "w:hAnsi", "w:cs"):
            rfonts.set(qn(attr), font_name)


def build_academic_docx(title, sections, references=None,
                        output_path="research.docx", lang="ar",
                        font=None, heading_size=16, body_size=14):
    """Build a complete academic research document.

    font: the font family to write into the document (e.g. "Kufyan Arabic
    Black", "Arial", "Simplified Arabic", "Times New Roman"). The NAME is
    preserved in the file so Word renders it with the real font on the
    device; if omitted, a sensible default per language is used.
    """
    from docx import Document
    from docx.shared import Pt  # noqa
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    is_rtl = (lang == "ar")
    # default font name written into the document
    if not font:
        font = "Kufyan Arabic Black" if is_rtl else "Times New Roman"

    # Title
    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(h, font, heading_size + 4)
    if is_rtl:
        _set_rtl(h)

    # Sections
    for sec in sections:
        heading = doc.add_heading(sec.get("heading", ""), level=1)
        _set_font(heading, font, heading_size)
        if is_rtl:
            heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _set_rtl(heading)
        body = doc.add_paragraph(sec.get("body", ""))
        _set_font(body, font, body_size)
        if is_rtl:
            body.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _set_rtl(body)

    # References — output-facing heading in the task language
    if references:
        ref_title = "المراجع" if is_rtl else "References"
        ref_heading = doc.add_heading(ref_title, level=1)
        _set_font(ref_heading, font, heading_size)
        if is_rtl:
            ref_heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _set_rtl(ref_heading)
        for ref in references:
            p = doc.add_paragraph(ref, style="List Number")
            _set_font(p, font, body_size)
            if is_rtl:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                _set_rtl(p)

    doc.save(output_path)
    return output_path


def _main():
    p = argparse.ArgumentParser(description="Build an academic Word document")
    p.add_argument("--json", required=True, help="JSON file with title/sections/references")
    p.add_argument("--output", default="research.docx")
    p.add_argument("--lang", default="ar")
    args = p.parse_args()

    with open(args.json, encoding="utf-8") as f:
        data = json.load(f)

    path = build_academic_docx(
        title=data["title"],
        sections=data.get("sections", []),
        references=data.get("references", []),
        output_path=args.output,
        lang=args.lang,
    )
    print(f"Created: {path}")


if __name__ == "__main__":
    _main()
