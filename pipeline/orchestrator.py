"""
pipeline/orchestrator.py
========================
المنسّق المركزي لـ Weaver Write.

يُدير ٥ مهام متوازية، كل مهمة في pipeline مستقل.
يُوزّع المهام على الطبقات ويُتابع حالة كل مهمة.
"""

from __future__ import annotations
import asyncio
import uuid
import time
from dataclasses import dataclass, field
from enum import Enum
from typing import Optional

try:
    from ..core.memory import MemoryManager, TaskMemory
    from ..core.sandbox import SandboxManager, TaskSandbox
except ImportError:
    import sys, os
    sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))
    from core.memory import MemoryManager, TaskMemory
    from core.sandbox import SandboxManager, TaskSandbox

MAX_TASKS = 5

# Built-in public SearXNG fallbacks, tried AUTOMATICALLY (no manual setup) only
# after the user's own instance/fallbacks AND DuckDuckGo have failed — so the
# reliable no-server path (DuckDuckGo) is never slowed by them. Public instances
# are flaky and many disable format=json, so each is attempted at most once per
# process (see _SEARX_DEAD) with a short timeout. Override/extend via the
# WEAVER_SEARXNG_FALLBACKS env var, which is tried earlier (before DuckDuckGo).
_DEFAULT_SEARXNG_FALLBACKS = [
    "https://searx.be",
    "https://search.inetol.net",
    "https://priv.au",
    "https://searx.tiekoetter.com",
]
_SEARX_DEAD = set()   # instances that failed this process — skipped next time

# Capability registry (Tools/Skills/Libraries) — Claude pattern
try:
    from capabilities import CapabilityRegistry
    _CAPABILITIES = CapabilityRegistry().load_all()
except Exception:
    _CAPABILITIES = None


class TaskStatus(str, Enum):
    QUEUED     = "في الطابور"
    LAYER_0    = "تنسيق"
    LAYER_1    = "بنية تحتية"
    LAYER_2    = "إدخال"
    LAYER_3    = "فهم"
    LAYER_4    = "بحث"
    LAYER_5    = "مصداقية"
    LAYER_6    = "صياغة"
    LAYER_6_5  = "إعادة صياغة"
    LAYER_7    = "تحقق"
    LAYER_8    = "إخراج"
    COMPLETED  = "مكتملة"
    FAILED     = "فشلت"


@dataclass
class Task:
    """مهمة بحثية واحدة."""
    task_id: str = field(default_factory=lambda: str(uuid.uuid4())[:8])
    description: str = ""
    input_files: list[str] = field(default_factory=list)
    status: TaskStatus = TaskStatus.QUEUED
    created_at: float = field(default_factory=time.time)
    started_at: Optional[float] = None
    completed_at: Optional[float] = None
    output_path: Optional[str] = None
    error: Optional[str] = None

    # أولوية الدور في الطابور — الأعلى يُنفَّذ أولاً (افتراضي 0)
    priority: int = 0

    # بطاقة المهمة (تُملأ في طبقة الفهم)
    task_card: dict = field(default_factory=dict)

    # ما يُوجَّه إليه في طبقة الفهم (Phase 3) ويُستهلك في الطبقات التالية
    tools: list = field(default_factory=list)     # أسماء الأدوات المطلوبة
    skills: list = field(default_factory=list)    # أسماء المهارات المطلوبة
    draft: str = ""                                # مسودة النص (طبقة ٦)
    sections: list = field(default_factory=list)   # أقسام الوثيقة النهائية

    def elapsed(self) -> float:
        if self.started_at:
            end = self.completed_at or time.time()
            return end - self.started_at
        return 0.0


class WeaverOrchestrator:
    """
    المنسّق المركزي — قلب Weaver Write.

    يُدير:
      - طابور المهام (Queue)
      - ٥ مهام نشطة بالتوازي
      - الذاكرة المعزولة لكل مهمة (UniMemory)
      - بيئة العزل لكل مهمة (OpenSandbox)
      - تتابع الطبقات من ٠ إلى ٨
    """

    def __init__(
        self,
        db_path: str = "./weaver_memory.db",
        sandbox_domain: str = "localhost:8080",
        sandbox_key: str = "",
        llm_fn=None,
        vision_fn=None,
    ):
        self.memory = MemoryManager(db_path=db_path)
        self.sandbox = SandboxManager(domain=sandbox_domain, api_key=sandbox_key)

        # The one LLM client, built from config/.env. May be None (no key) →
        # every layer then keeps its offline placeholder behaviour.
        try:
            from core.llm import get_llm_fn, get_vision_fn
            self.llm_fn = llm_fn or get_llm_fn()
            self.vision_fn = vision_fn or get_vision_fn()
        except Exception:
            self.llm_fn = llm_fn
            self.vision_fn = vision_fn
        self.caps = _CAPABILITIES

        # main system prompt + the professional-conduct rule (rule 10), so the
        # MODEL itself also stays calm under hostility
        try:
            from pipeline.prompts import SYSTEM_PROMPT_MAIN
            from capabilities.skills.conduct_guard.scripts.conduct_guard import (
                CONDUCT_SYSTEM_RULE)
            self.system_main = SYSTEM_PROMPT_MAIN + "\n\n" + CONDUCT_SYSTEM_RULE
        except Exception:
            try:
                from pipeline.prompts import SYSTEM_PROMPT_MAIN
                self.system_main = SYSTEM_PROMPT_MAIN
            except Exception:
                self.system_main = None

        self._queue: list[Task] = []
        self._active: dict[str, Task] = {}      # task_id → Task
        self._completed: list[Task] = []
        self._lock = asyncio.Lock()

    # ── إضافة مهمة ──

    async def submit(self, description: str, input_files: list[str] = None,
                     priority: int = 0) -> Task:
        """
        يُضيف مهمة جديدة. حتى ٥ مهام تعمل بالتوازي؛ الزائد يدخل طابور أولوية:
        الأعلى `priority` يُنفَّذ أولاً (وعند التساوي: الأقدم أولاً).
        """
        task = Task(
            description=description,
            input_files=input_files or [],
            priority=priority,
        )

        async with self._lock:
            if len(self._active) < MAX_TASKS:
                await self._start_task(task)
            else:
                # priority insert: place before the first lower-priority task
                idx = len(self._queue)
                for i, q in enumerate(self._queue):
                    if q.priority < task.priority:
                        idx = i
                        break
                self._queue.insert(idx, task)
                print(f"📋 مهمة [{task.task_id}] في الطابور "
                      f"(أولوية {task.priority}، {len(self._queue)} بالانتظار)")

        return task

    async def _start_task(self, task: Task):
        """يبدأ مهمة جديدة."""
        task.status = TaskStatus.LAYER_0
        task.started_at = time.time()
        self._active[task.task_id] = task

        # إنشاء ذاكرة معزولة
        self.memory.create_task(task.task_id)

        # إنشاء sandbox معزول
        await self.sandbox.create_for_task(task.task_id)

        print(f"🚀 بدأت مهمة [{task.task_id}]: {task.description[:50]}")

        # تشغيل في الخلفية
        asyncio.create_task(self._run_pipeline(task))

    # ── Pipeline كامل ──

    async def _run_pipeline(self, task: Task):
        """يُشغّل pipeline المهمة من الطبقة ٠ إلى ٨."""
        mem = self.memory.get_task(task.task_id)
        sb = self.sandbox.get(task.task_id)

        try:
            # ── conduct guard (before Layer 0): stay professional under abuse ──
            try:
                lang0 = self._detect_lang(task.description)
                g = self._skill_call("conduct_guard", "conduct_guard",
                                     "guard_response", task.description, lang0)
                task.task_card["conduct"] = g
                if g.get("hostile") and not g.get("do_task"):
                    # abuse only, no task: calm redirect, do nothing else
                    task.task_card["reply"] = g.get("reply_prefix", "")
                    task.status = TaskStatus.COMPLETED
                    task.completed_at = time.time()
                    return
            except Exception:
                pass

            # الطبقات بالتسلسل
            await self._layer_0(task, mem)
            await self._layer_1(task, mem, sb)
            await self._layer_2(task, mem)
            await self._layer_3(task, mem)
            await self._layer_4(task, mem)
            await self._layer_5(task, mem)
            await self._layer_6(task, mem)
            await self._layer_6_6(task, mem)
            await self._layer_6_5(task, mem)
            await self._layer_7(task, mem)
            await self._layer_8(task, mem)

            # اكتمال
            task.status = TaskStatus.COMPLETED
            task.completed_at = time.time()
            print(f"✅ مهمة [{task.task_id}] اكتملت في {task.elapsed():.0f}ث")

            # استخلاص دروس
            mem.distill_task_lessons([
                {"role": "system", "content": f"مهمة اكتملت: {task.description}"},
                {"role": "system", "content": f"المخرج: {task.output_path}"},
            ])

        except Exception as e:
            task.status = TaskStatus.FAILED
            task.error = str(e)
            print(f"❌ مهمة [{task.task_id}] فشلت: {e}")

        finally:
            # تنظيف
            await self.sandbox.destroy(task.task_id)
            self.memory.close_task(task.task_id)

            async with self._lock:
                self._active.pop(task.task_id, None)
                self._completed.append(task)
                # تشغيل مهمة من الطابور إن وُجدت
                if self._queue:
                    next_task = self._queue.pop(0)
                    await self._start_task(next_task)

    # ── الطبقات ──

    async def _layer_0(self, task: Task, mem: TaskMemory):
        """٠: التنسيق — تسجيل المهمة وإعداد السياق."""
        task.status = TaskStatus.LAYER_0
        mem.set_status(0, "بدأ التنسيق")
        await asyncio.sleep(0)  # yield للـ event loop

    async def _layer_1(self, task: Task, mem: TaskMemory, sb: TaskSandbox):
        """١: البنية التحتية — تجهيز sandbox والأدوات."""
        task.status = TaskStatus.LAYER_1
        mem.set_status(1, "تجهيز البنية التحتية")
        # تثبيت مكتبات إضافية إن لزم
        if sb:
            await sb.install("paperqa", "pymupdf4llm")

    async def _layer_2(self, task: Task, mem: TaskMemory):
        """٢: الإدخال — قراءة الملفات مع أرقام الصفحات."""
        task.status = TaskStatus.LAYER_2
        mem.set_status(2, "قراءة الملفات")
        from core.ocr import WeaverOCR
        ocr = WeaverOCR()
        for filepath in task.input_files:
            doc = ocr.read_with_pages(filepath)
            # حفظ محتوى كل صفحة في الذاكرة
            for page in doc.pages:
                mem.add_reference(
                    f"[{os.path.basename(filepath)}] {page.text[:200]}",
                    page=page.page,
                )

    @staticmethod
    def _detect_lang(text: str) -> str:
        """Cheap language guess for the conduct guard: Arabic if any Arabic
        letter is present, else English."""
        for ch in (text or ""):
            if "؀" <= ch <= "ۿ":
                return "ar"
        return "en"

    # explicit output-language directives (request words → language name the model
    # will write in). Order matters only for display; matching is substring-based.
    _LANG_NAMED = (
        ("العربية", ("بالعربية", "بالعربي", "باللغة العربية", "اكتبها بالعربية",
                     "اكتبه بالعربية", " عربي ", " عربى ", "in arabic", " arabic")),
        ("الإنجليزية", ("بالانجليزية", "بالإنجليزية", "بالانكليزية", "بالإنكليزية",
                        " انجليزي ", " إنجليزي ", "in english", " english")),
        ("الفرنسية", ("بالفرنسية", "بالفرنسي", "in french", " french")),
        ("الإسبانية", ("بالاسبانية", "بالإسبانية", "in spanish", " spanish")),
        ("الألمانية", ("بالالمانية", "بالألمانية", "in german", " german")),
        ("التركية", ("بالتركية", "in turkish", " turkish")),
        ("الروسية", ("بالروسية", "in russian", " russian")),
        ("الصينية", ("بالصينية", "in chinese", " chinese")),
        ("الأردية", ("بالاردية", "بالأردية", "in urdu", " urdu")),
        ("الفارسية", ("بالفارسية", "in persian", " persian", " farsi")),
        ("الهندية", ("بالهندية", "in hindi", " hindi")),
    )
    _LANG_SOURCE = (
        "اللغة الاصلية", "اللغة الأصلية", "بلغتها الاصلية", "بلغتها الأصلية",
        "بلغته الاصلية", "بلغته الأصلية", "بلغة الاصلية", "بلغة الأصلية",
        "بلغة الفيديو", "لغة الفيديو", "بلغة المصدر", "بلغة النص الاصلي",
        "بلغة النص الأصلي", "كما هي بلغتها", "بنفس اللغة", "بنفس لغة",
        "original language", "source language", "same language",
        "in its original", "keep the language", "keep it in",
    )

    @staticmethod
    def _task_scope(text):
        """Detect a SCOPE that limits a document task: 'references' (only find
        and list references/studies, no writing), 'outline' (only the structure),
        'part' (write only a specific part), or None (full document)."""
        t = " " + (text or "").lower() + " "
        refs = ("مراجع فقط", "المراجع فقط", "فقط المراجع", "فقط مراجع",
                "دراسات سابقة فقط", "فقط الدراسات", "الدراسات فقط", "فقط المصادر",
                "المصادر فقط", "قائمة مراجع", "قائمة المراجع", "قائمة مصادر",
                "اوجد مراجع", "أوجد مراجع", "ابحث عن مراجع", "ابحث لي عن مراجع",
                "جد مراجع", "هات مراجع", "اعطني مراجع", "أعطني مراجع",
                "اعطني مصادر", "اعطني دراسات", "دون كتابة الموضوع",
                "references only", "just references", "only references",
                "list of references", "find references", "find sources",
                "sources only", "only sources", "list sources", "bibliography")
        outline = ("هيكل فقط", "الهيكل فقط", "فقط الهيكل", "العناصر فقط",
                   "فقط العناصر", "الخطة فقط", "خطة البحث فقط", "فقط الخطة",
                   "عناصر البحث فقط", "جدول المحتويات", "الخطوط العريضة",
                   "outline only", "just an outline", "only an outline",
                   "structure only", "just the outline", "table of contents",
                   "only outline")
        part = ("اكتب فقط", "فقط اكتب", "جزء فقط", "فقط جزء", "المقدمة فقط",
                "فقط المقدمة", "قسم فقط", "فقط قسم", "فقرة فقط", "فقط الخاتمة",
                "الخاتمة فقط", "فصل فقط", "فقط هذا الجزء",
                "only the introduction", "only the conclusion",
                "just write the", "only write the", "write only the",
                "just the section", "only this part", "only this section")
        if any(k in t for k in refs):
            return "references"
        if any(k in t for k in outline):
            return "outline"
        if any(k in t for k in part):
            return "part"
        return None

    def _format_references_only(self, card, lang):
        """Build a plain numbered references list from the gathered sources."""
        srcs = card.get("sources") or []
        if not srcs:
            return ("لم يُعثر على مراجع/مصادر لهذا الموضوع الآن — جرّب لاحقاً أو "
                    "وسّع الصياغة." if lang == "ar"
                    else "No references/sources were found for this topic.")
        lines, seen, n = [], set(), 1
        for s in srcs:
            url = (s.get("url") or "").strip()
            title = (s.get("title") or url or "").strip()
            if not title or title in seen:
                continue
            seen.add(title)
            auth = ", ".join(s.get("authors") or [])
            year = str(s.get("year") or "").strip()
            doi = (s.get("doi") or "").strip()
            meta = " — ".join(x for x in (auth, year) if x)
            tail = (("doi:" + doi) if doi else url).strip()
            lines.append(f"{n}. {title}"
                         + (f" — {meta}" if meta else "")
                         + (f". {tail}" if tail else "."))
            n += 1
        return "\n".join(lines)

    @classmethod
    def _requested_output_lang(cls, text):
        """Detect an EXPLICIT output-language directive in the request. Returns
        ("source", None) for the content's own/original language, ("name", <lang>)
        when a language is named, or (None, None) when unspecified. Substring
        match on a space-padded, lowercased copy so word forms are tolerant."""
        t = " " + (text or "").lower() + " "
        if any(k in t for k in cls._LANG_SOURCE):
            return ("source", None)
        for lang, kws in cls._LANG_NAMED:
            if any(k in t for k in kws):
                return ("name", lang)
        return (None, None)

    @staticmethod
    def _sourcing_mode(text: str) -> str:
        """How the user wants sourcing handled — conservative: only an EXPLICIT
        request flips away from the default.
          "none"    → write WITHOUT any sources/references/studies.
          "uncited" → research FROM sources but DON'T document/cite them.
          "cited"   → default: research, cite in-text, and list references.
        """
        t = " " + (text or "").lower() + " "
        # sources are USED but must NOT be documented/cited
        uncited = (
            "بدون توثيق", "بلا توثيق", "دون توثيق", "من غير توثيق",
            "لا توثقها", "لا توثق", "بدون توثيقها", "دون توثيقها",
            "بدون ذكر المراجع", "دون ذكر المراجع", "بدون ذكر المصادر",
            "دون ذكر المصادر", "لا تذكر المراجع", "لا تذكر المصادر",
            "بدون ان توثقها", "بدون أن توثقها", "لكن لا توثقها",
            "without citing", "without documenting", "don't cite",
            "do not cite", "no in-text citation", "no in text citation",
            "uncited", "without a references list", "no references list",
        )
        # NO sources at all
        none_src = (
            "بدون مصادر", "بلا مصادر", "دون مصادر", "من غير مصادر",
            "بدون أي مصادر", "بدون اي مصادر", "بدون مراجع", "بلا مراجع",
            "دون مراجع", "من غير مراجع", "بدون دراسات", "بلا دراسات",
            "دون دراسات", "من غير دراسات", "بدون مصادر ومراجع",
            "without sources", "without references", "no sources",
            "no references", "no citations", "source-free", "without any sources",
        )
        if any(p in t for p in uncited):
            return "uncited"
        if any(p in t for p in none_src):
            return "none"
        return "cited"

    @staticmethod
    def _strip_citations(text: str) -> str:
        """Remove parenthesised in-text citations (…, YEAR) / (key, p. N) /
        (…، ص. N) from prose. Used in the no-citation writing modes so an
        accidental citation from the model never survives to the output."""
        import re
        if not text:
            return text
        text = re.sub(
            r"\s*\([^()]*(?:\b\d{4}\b|p\.?\s*\d+|ص\.?\s*\d+)[^()]*\)", "", text)
        return re.sub(r"[ \t]{2,}", " ", text)

    @staticmethod
    def _looks_conversational(text: str) -> bool:
        """True when a section body is a chat turn (greeting / clarifying
        question / options menu) instead of document content — so it can be
        retried or dropped. Conservative: needs a real chat marker, not just a
        question mark inside otherwise substantial prose."""
        t = (text or "").strip()
        if not t:
            return False
        head = t[:400]
        markers = (
            "أهلاً", "أهلًا", "اهلا", "مرحبا", "مرحباً", "عزيزي",
            "ما الذي تريد", "ماذا تريد", "يرجى التوضيح", "الرجاء التوضيح",
            "أحتاج أن أحدد", "أحتاج إلى تحديد", "هل تريد", "هل تفضل",
            "بحاجة إلى مزيد", "أخبرني", "قبل أن أبدأ", "قبل أن أكتب",
            "hello", "hi there", "could you clarify", "what would you like",
            "which of the following", "please specify", "let me know",
            "before i begin", "i need to know", "would you like",
        )
        low = head.lower()
        if any(m in head or m in low for m in markers):
            return True
        # an options menu near the top: "أ." / "ب." / "ج." or "a)" "b)" list
        import re
        if re.search(r"(^|\n)\s*[أ-د]\s*[\.\)\-]", head) and (
                "؟" in head or "?" in head):
            return True
        # very short and ends in a question → almost certainly a clarifying Q
        if len(t) < 200 and t.rstrip().endswith(("؟", "?")):
            return True
        return False

    @staticmethod
    def _skill_call(skill: str, module: str, func: str, *args, **kwargs):
        """Dynamically import capabilities/skills/<skill>/scripts/<module>.py
        and call <func>(*args, **kwargs). Raises on failure — callers guard it
        so the pipeline degrades to placeholder behaviour."""
        import os as _os, sys as _sys, importlib
        sp = _os.path.abspath(_os.path.join(
            _os.path.dirname(__file__), "..", "capabilities", "skills",
            skill, "scripts"))
        if sp not in _sys.path:
            _sys.path.insert(0, sp)
        mod = importlib.import_module(module)
        return getattr(mod, func)(*args, **kwargs)

    @staticmethod
    def _requested_format(text):
        """Detect an EXPLICIT output format in the request (Word/PDF/PowerPoint/
        Excel/HTML/Text/Markdown) — broad natural phrasings. None otherwise."""
        t = " " + (text or "").lower() + " "
        if any(k in t for k in (
                "بوربوينت", "باوربوينت", "بوربوينت", "باور بوينت", "بور بوينت",
                "powerpoint", "power point", "pptx", "ppt ", "عرض تقديمي",
                "عرض بوربوينت", "شرائح", "شريحة", "بريزنتيشن", " slides",
                " presentation", "deck")):
            return "PPTX"
        if any(k in t for k in (
                " csv", ".csv", "ملف csv", "سي اس في", "قيم مفصولة بفواصل",
                "comma separated", "comma-separated")):
            return "CSV"
        if any(k in t for k in (
                "اكسل", "إكسل", "اكسيل", "excel", "xlsx", "xls ", "جدول بيانات",
                "جدول اكسل", "شيت", "spreadsheet", "sheet")):
            return "XLSX"
        if any(k in t for k in (
                " html", " htm ", "اتش تي ام ال", "صفحة ويب", "صفحة انترنت",
                "صفحه ويب", "webpage", "web page", "html file")):
            return "HTML"
        if any(k in t for k in (
                " pdf", "pdf ", "بي دي اف", "بيدياف", "بي دي إف", "ملف pdf")):
            return "PDF"
        if any(k in t for k in (
                "وورد", "word", "docx", "doc ", "ملف وورد", "مستند وورد",
                "مايكروسوفت وورد", "ورد ", "word document")):
            return "DOCX"
        if any(k in t for k in (
                "ملف نصي", "ملف نصّي", "نص عادي", "نصي فقط", "txt", ".txt",
                "text file", "plain text", "as text")):
            return "TXT"
        if any(k in t for k in (
                "ماركداون", "ماركدوان", "markdown", ".md", " md ", "نص فقط")):
            return "INLINE"
        return None

    @staticmethod
    def _primary_format(task_card: dict) -> str:
        """The first requested output format as a lowercase string (docx/pptx/
        xlsx/pdf), tolerating either a list or a bare string in the card."""
        of = task_card.get("output_format", ["DOCX"])
        if isinstance(of, list):
            of = of[0] if of else "DOCX"
        return str(of).lower()

    def _placeholder_card(self, task: Task) -> dict:
        """The offline fallback task card (used when llm_fn is None or fails)."""
        return {
            "task_type": "بحث",
            "topic": task.description,
            "language": "ar",
            "citation_style": "APA",
            "output_format": ["DOCX"],
        }

    @staticmethod
    def _model_strength() -> str:
        """Estimate the running model's strength — "small" | "medium" | "large".

        Order: an explicit override (WEAVER_MODEL_STRENGTH) wins; otherwise the
        model NAME (WEAVER_MODEL) is matched against size hints. Purpose: let
        every layer adapt depth/temperature/length to the model's ceiling so a
        small model works reliably at its own peak, and a large one is used to
        its full depth — without changing any skill. Unknown → "medium"."""
        import os as _os
        ov = (_os.environ.get("WEAVER_MODEL_STRENGTH", "") or "").strip().lower()
        if ov in ("small", "weak", "low", "tiny", "ضعيف", "صغير"):
            return "small"
        if ov in ("medium", "mid", "متوسط"):
            return "medium"
        if ov in ("large", "strong", "high", "big", "كبير", "قوي"):
            return "large"
        name = (_os.environ.get("WEAVER_MODEL", "") or "").lower()
        large_kw = ("opus", "ultra", "pro", "70b", "72b", "65b", "405b", "110b",
                    "large", "huge", "32b", "34b", "-max", "gpt-4o", "gpt-4.1",
                    "o1", "o3", "sonnet-4", "sonnet-5", "opus-5")
        small_kw = ("flash", "mini", "nano", "lite", "tiny", "small", "0.5b",
                    "1b", "1.5b", "2b", "3b", "4b", "7b", "8b", "9b", "haiku")
        # a small/flash/mini variant is small even inside a large family
        # (e.g. gpt-4o-mini, gemini-flash) → check the small hints first
        if any(k in name for k in small_kw):
            return "small"
        if any(k in name for k in large_kw):
            return "large"
        return "medium"

    @staticmethod
    def _strength_profile(strength: str) -> dict:
        """Per-strength writing profile: model temperature, target words for a
        specialized intro, and a depth directive appended to the generic section
        prompt. Adapts OUTPUT to the model ceiling — never fabricates capability
        a small model lacks; it raises the reliable floor and unlocks depth on a
        capable model. Returns a dict always usable (unknown → medium)."""
        s = (strength or "medium").lower()
        if s == "small":
            return {
                "temp": 0.35,
                "intro_words": 220,
                "depth": (
                    "النموذج محدود الطاقة: اكتب بجُملٍ قصيرة واضحة ومباشرة، وركّز "
                    "على النقاط الجوهرية دون حشوٍ أو استطراد، ورتّب الأفكار في "
                    "فقراتٍ قصيرة. الدقّة والوضوح والالتزام بالمصادر أهمّ من الطول "
                    "(استهدف نحو 180–260 كلمة لهذا القسم)."),
                "depth_en": (
                    "The model has limited capacity: write short, clear, direct "
                    "sentences; focus on the essential points with no padding; "
                    "keep paragraphs short. Accuracy, clarity and staying on "
                    "sources matter more than length (aim ~180–260 words)."),
            }
        if s == "large":
            return {
                "temp": 0.6,
                "intro_words": 650,
                "depth": (
                    "استغلّ طاقة النموذج الكاملة: حلّل بعمق، واعرض وجهات النظر "
                    "المختلفة، واربط الأفكار ببعضها بنقدٍ علميّ وأمثلةٍ دقيقة، مع "
                    "التزامٍ صارمٍ بالمصادر (استهدف نحو 500–800 كلمة لهذا القسم)."),
                "depth_en": (
                    "Use the model's full capacity: analyze in depth, present "
                    "differing viewpoints, and connect ideas with scholarly "
                    "critique and precise examples, strictly grounded in the "
                    "sources (aim ~500–800 words)."),
            }
        return {"temp": 0.5, "intro_words": 400, "depth": "", "depth_en": ""}

    @staticmethod
    def _section_kind(title: str):
        """Classify a section title into a purpose-built writer kind:
        "intro" | "conclusion" | "results" | None (→ generic writer)."""
        t = (title or "").strip().lower()
        if not t:
            return None
        intro_kw = ("مقدمة", "المقدمة", "تمهيد", "introduction", "intro")
        concl_kw = ("خاتمة", "الخاتمة", "خلاصة", "الخلاصة", "استنتاج",
                    "الاستنتاجات", "التوصيات", "توصيات", "conclusion",
                    "recommendation", "closing")
        res_kw = ("النتائج", "نتائج", "تحليل النتائج", "عرض النتائج",
                  "results", "findings")
        if any(k in t for k in intro_kw):
            return "intro"
        if any(k in t for k in concl_kw):
            return "conclusion"
        if any(k in t for k in res_kw):
            return "results"
        return None

    def _write_section_specialized(self, title, card, lang, mode, no_ctx,
                                   prior_sections, prof):
        """Route a section to its purpose-built writer skill when the section
        type AND sourcing mode fit, returning prose text — or None to let the
        generic writer handle it. Purely additive: every skill call is guarded
        by the caller, so any miss falls back to the existing generic path.

        Bindings (skills already present, only wired here):
          intro       → research_intro.build_intro
          conclusion  → conclusion_writer.build_conclusion
          results     → results_formatter.format_results (also uses table_builder)
        """
        if not self.llm_fn:
            return None
        kind = self._section_kind(title)
        if not kind:
            return None
        topic = card.get("topic", "") or title
        if kind == "intro" and mode == "cited" and not no_ctx:
            refs = []
            for s in (card.get("sources") or [])[:12]:
                if isinstance(s, dict):
                    refs.append({
                        "key": s.get("key") or (s.get("title", "") or "")[:40],
                        "text": (s.get("content") or s.get("title", "") or "")[:160],
                        "page": s.get("page", "")})
            out = self._skill_call(
                "research_intro", "build_intro", "build_intro",
                topic, refs, int(prof.get("intro_words", 400)), lang, self.llm_fn)
            return (out or {}).get("text") or None
        if kind == "conclusion" and mode != "none":
            findings = []
            for sec in (prior_sections or [])[-8:]:
                b = (sec.get("body") or "").strip()
                if b:
                    first = b.split("\n", 1)[0].strip()[:200]
                    if first:
                        findings.append(first)
            out = self._skill_call(
                "conclusion_writer", "build_conclusion", "build_conclusion",
                topic, findings, lang, self.llm_fn)
            return (out or {}).get("text") or None
        if kind == "results" and mode != "none":
            out = self._skill_call(
                "results_formatter", "format_results", "format_results",
                [{"title": title, "note": ""}], lang, self.llm_fn)
            return (out or {}).get("text") or None
        return None

    # ── task.skills dispatch ─────────────────────────────────────────────────
    # _route() matches skills to the task; this turns that match into real
    # execution. A skill runs ONLY when the task actually selected it (its name
    # is in task.skills), keeping behaviour targeted. Skills already invoked at
    # a fixed point (structure/methodology/rewriters/formatters/builders and the
    # per-section writers above) are NOT re-run here — this dispatch adds the
    # remaining enrichment skills that had no wiring. Every handler is guarded,
    # idempotent, and additive; on any miss the draft is left unchanged.
    def _skill_handlers(self):
        """skill name → write-stage handler(self, task, card, lang, mem)->bool.
        A skill absent here is either wired elsewhere (a fixed layer point) or
        context-gated for a later increment; its match is simply skipped."""
        return {
            "literature_review": self._sk_literature_review,
        }

    def _dispatch_skills(self, task: Task, card: dict, lang: str, mem):
        """Run the matched skills (task.skills) that have a write-stage handler.
        Guarded per skill; a failure never breaks the draft."""
        handlers = self._skill_handlers()
        for name in list(task.skills or []):
            h = handlers.get(name)
            if not h:
                continue
            try:
                if h(task, card, lang, mem):
                    mem.set_status(6, f"مهارة موزّعة: {name} ✓")
            except Exception as e:
                mem.set_status(6, f"مهارة {name} (تخطّي: {e})")

    @staticmethod
    def _is_literature_title(title: str) -> bool:
        t = (title or "").strip().lower()
        return any(k in t for k in (
            "الدراسات السابقة", "دراسات سابقة", "أدبيات", "الأدبيات",
            "الإطار النظري", "مراجعة الأدبيات", "literature", "related work",
            "prior work", "background"))

    def _sk_literature_review(self, task: Task, card: dict, lang: str,
                              mem) -> bool:
        """Enrich an EXISTING literature/theoretical-framework section with a
        theme-organized view of the gathered sources (organize_by_theme). Never
        invents a section: if no literature section was written, it does nothing.
        Additive — appends beneath the section's current body."""
        sources = [s for s in (card.get("sources") or []) if isinstance(s, dict)]
        if len(sources) < 2 or not task.sections:
            return False
        lit_idx = next((i for i, s in enumerate(task.sections)
                        if self._is_literature_title(s.get("heading", ""))), None)
        if lit_idx is None:
            return False
        refs = [{"key": s.get("key") or (s.get("title", "") or "")[:40],
                 "text": (s.get("content") or s.get("title", "") or "")}
                for s in sources]
        groups = self._skill_call("literature_review", "organize_by_theme",
                                  "organize_by_theme", refs, None) or {}
        lines = []
        for theme, items in groups.items():
            if theme == "unclassified" or not items:
                continue
            keys = "؛ ".join((it.get("key") or "")[:60] for it in items[:6]) \
                if lang == "ar" else \
                "; ".join((it.get("key") or "")[:60] for it in items[:6])
            lines.append(f"- **{theme}**: {keys}")
        if not lines:
            return False
        header = ("\n\n**تنظيم الدراسات موضوعياً:**\n" if lang == "ar"
                  else "\n\n**Thematic grouping of studies:**\n")
        cur = task.sections[lit_idx].get("body", "") or ""
        task.sections[lit_idx]["body"] = cur + header + "\n".join(lines)
        # rebuild the chat/preview draft to reflect the enriched section
        task.draft = "\n\n".join(
            (f"{s.get('heading', '')}\n{s.get('body', '')}").strip()
            for s in task.sections if (s.get("heading") or s.get("body")))
        return True

    # ── statistical_analysis: real stats when a data file is attached ─────────
    @staticmethod
    def _data_files(task: Task):
        """Attached data files (csv/xlsx/xls) the stats skill can analyze."""
        exts = (".csv", ".xlsx", ".xls")
        return [f for f in (task.input_files or [])
                if isinstance(f, str) and f.lower().endswith(exts)]

    @staticmethod
    def _format_statistics(res, lang: str):
        """Render analyze()'s REAL computed numbers as a Markdown block
        (descriptives table + reliability + an honest 'computed, not estimated'
        note). Returns None on error/empty so nothing fake is ever injected."""
        if not isinstance(res, dict) or res.get("error"):
            return None
        dd = (res.get("descriptives") or {})
        desc = dd.get("descriptives") or {}
        n = dd.get("n") or res.get("n")
        if not desc:
            return None
        cols = list(desc.keys())
        order = ["count", "mean", "std", "min", "25%", "50%", "75%", "max"]
        first = desc[cols[0]] if cols else {}
        stats = [k for k in order if k in first] or list(first.keys())
        labels_ar = {"count": "العدد", "mean": "المتوسط",
                     "std": "الانحراف المعياري", "min": "الأدنى",
                     "25%": "الربيع الأول", "50%": "الوسيط",
                     "75%": "الربيع الثالث", "max": "الأعلى"}

        def _fmt(v):
            try:
                return str(round(float(v), 3))
            except Exception:
                return str(v)
        head_stat = "الإحصاء" if lang == "ar" else "Statistic"
        rows = ["| " + head_stat + " | " + " | ".join(str(c) for c in cols) + " |",
                "|" + "---|" * (len(cols) + 1)]
        for st in stats:
            label = labels_ar.get(st, st) if lang == "ar" else st
            rows.append("| " + " | ".join(
                [label] + [_fmt(desc[c].get(st, "")) for c in cols]) + " |")
        out = []
        out.append(("حجم العينة: %s. الإحصاءات الوصفية للمتغيّرات العددية:" % n)
                   if lang == "ar" else
                   ("Sample size: %s. Descriptive statistics for numeric "
                    "variables:" % n))
        out.append("\n".join(rows))
        rel = res.get("reliability")
        if isinstance(rel, dict) and "cronbach_alpha" in rel:
            out.append(
                ("**ثبات المقياس (كرونباخ ألفا):** %s (%s عبارة، ن=%s) — %s."
                 % (rel.get("cronbach_alpha"), rel.get("n_items", ""),
                    rel.get("n", ""), rel.get("interpretation", "")))
                if lang == "ar" else
                ("**Reliability (Cronbach's α):** %s (%s items, n=%s) — %s."
                 % (rel.get("cronbach_alpha"), rel.get("n_items", ""),
                    rel.get("n", ""), rel.get("interpretation", ""))))
        out.append("_" + ("الأرقام أعلاه محسوبة فعلياً من الملف المرفق، لم "
                          "تُقدَّر أو تُختلق." if lang == "ar" else
                          "The figures above are computed directly from the "
                          "attached file, not estimated.") + "_")
        return "\n\n".join(out)

    def _inject_statistics(self, task: Task, card: dict, lang: str, mem):
        """When a data file is attached, run statistical_analysis.analyze on it
        and inject the REAL computed results into the document — into a results
        section if one exists, else as its own 'التحليل الإحصائي' section. Never
        fabricates numbers: on a library/read error it adds an honest note only.
        Additive and fully guarded."""
        files = self._data_files(task)
        if not files:
            return
        path = files[0]
        try:
            res = self._skill_call("statistical_analysis", "survey_analysis",
                                   "analyze", path)
            # add scale reliability when Likert-type items are detected
            vt = (res or {}).get("variable_types") or {}
            likert = [c for c, t in vt.items()
                      if "likert" in str(t).lower() or "ليكرت" in str(t)]
            if isinstance(res, dict) and "error" not in res and len(likert) >= 2:
                res2 = self._skill_call("statistical_analysis",
                                        "survey_analysis", "analyze", path,
                                        likert)
                if isinstance(res2, dict) and "error" not in res2:
                    res = res2
        except Exception as e:
            mem.set_status(6, f"إحصاء (تخطّي: {e})")
            return
        card["statistics"] = res
        block = self._format_statistics(res, lang)
        head = "التحليل الإحصائي" if lang == "ar" else "Statistical Analysis"
        if not block:
            # honest, actionable note — never fake numbers
            err = res.get("error") if isinstance(res, dict) else "unknown"
            block = (("تعذّر تنفيذ التحليل الإحصائي على الملف المرفق: %s. "
                      "قد تحتاج تثبيت المكتبات: pip install pandas scipy." % err)
                     if lang == "ar" else
                     ("Could not run the statistical analysis on the attached "
                      "file: %s. You may need: pip install pandas scipy." % err))
        idx = next((i for i, s in enumerate(task.sections or [])
                    if self._section_kind(s.get("heading", "")) == "results"),
                   None)
        if idx is not None:
            cur = task.sections[idx].get("body", "") or ""
            sub = ("\n\n**التحليل الإحصائي:**\n\n" if lang == "ar"
                   else "\n\n**Statistical analysis:**\n\n")
            task.sections[idx]["body"] = cur + sub + block
        else:
            task.sections = (task.sections or []) + [{"heading": head,
                                                      "body": block}]
        task.draft = "\n\n".join(
            (f"{s.get('heading', '')}\n{s.get('body', '')}").strip()
            for s in task.sections if (s.get("heading") or s.get("body")))
        mem.set_status(6, "أُدرج التحليل الإحصائي (أرقام محسوبة فعلياً)")

    # ── quran_hadith_citation: correct marks for Islamic content ─────────────
    @staticmethod
    def _is_islamic_content(text: str) -> bool:
        """True when the text quotes/discusses Quran or Hadith (so the marks
        skill should enforce ﴿ ﴾ for verses and « » for hadith)."""
        t = text or ""
        kw = ("قال الله", "قال تعالى", "يقول الله", "سبحانه وتعالى", "عز وجل",
              "قال رسول الله", "قال النبي", "عن النبي", "صلى الله عليه وسلم",
              "ﷺ", "رواه البخاري", "رواه مسلم", "حديث شريف", "الحديث الشريف",
              "القرآن", "قرآن كريم", "آية كريمة", "الآية الكريمة",
              "﴾", "«", "السنة النبوية", "السيرة النبوية")
        return any(k in t for k in kw)

    # the writing directive appended for Islamic content (skill's conventions)
    _ISLAMIC_DIRECTIVE_AR = (
        "عند الاستشهاد بآية قرآنية: ضعها بين قوسي الآية ﴿ ﴾ (لا أقواس عادية ولا "
        "علامات اقتباس) وأتبِعها بالمصدر (السورة: رقم الآية). وعند الاستشهاد بحديث "
        "نبوي: ضعه بين علامتي « » (لا أقواس الآية) وأتبِعه بالتخريج (رواه فلان، "
        "الحكم). لا تخلط بين العلامتين إطلاقاً.")
    _ISLAMIC_DIRECTIVE_EN = (
        "When quoting a Quranic verse, enclose it in the ornamental brackets "
        "﴿ ﴾ (never normal quotes/parentheses) and follow it with (Surah: Ayah). "
        "When quoting a hadith, enclose it in « » (never the Quran brackets) and "
        "follow it with its takhrij. Never mix the two marks.")

    def _apply_islamic_marks(self, task: Task, card: dict, lang: str, mem):
        """For Islamic content, enforce the skill's marks at the TEXT level so
        every export format is correct: a verse introduced by an explicit Quran
        lead-in gets ﴿ ﴾, a hadith introduced by an explicit lead-in gets « ».
        ONLY the delimiter marks are changed — the quoted text itself is kept
        verbatim (never rewritten). Then the skill's validate_marks flags any
        remaining misuse. Additive, guarded, and a no-op for non-Islamic text.

        Note: this normalizes the MARKS across all formats; the skill's richer
        Word styling (bold verse, Kufyan font) via add_quran_verse/add_hadith
        needs a docx object and stays a later, docx-only step."""
        import re
        sections = task.sections or []
        joined = "\n".join((s.get("body", "") or "") for s in sections) \
            or (task.draft or "")
        if not self._is_islamic_content(joined + " " + str(card.get("topic", ""))):
            return
        # wire to the skill module: real marks + validator (no docx needed here)
        try:
            import importlib, sys as _sys, os as _os
            sp = _os.path.abspath(_os.path.join(
                _os.path.dirname(__file__), "..", "capabilities", "skills",
                "quran_hadith_citation", "scripts"))
            if sp not in _sys.path:
                _sys.path.insert(0, sp)
            qh = importlib.import_module("quran_hadith")
        except Exception as e:
            mem.set_status(6, f"تنسيق إسلامي (تخطّي: {e})")
            return
        QO, QC = qh.QURAN_OPEN, qh.QURAN_CLOSE
        HO, HC = qh.HADITH_OPEN, qh.HADITH_CLOSE
        q = '["“”]'      # straight or curly double quotes
        qlead = (r'(?:قال\s+الله\s+تعالى|قال\s+تعالى|قال\s+الله|'
                 r'يقول\s+الله(?:\s+تعالى)?|قال\s+عز\s+وجل)')
        hlead = (r'(?:قال\s+رسول\s+الله(?:\s*ﷺ|\s*صلى\s+الله\s+عليه\s+وسلم)?|'
                 r'قال\s+النبي(?:\s*ﷺ|\s*صلى\s+الله\s+عليه\s+وسلم)?|عن\s+النبي)')
        q_re = re.compile(r'(' + qlead + r'\s*[:：]?\s*)' + q +
                          r'([^"“”\n]{3,300})' + q)
        h_re = re.compile(r'(' + hlead + r'\s*[:：]?\s*)' + q +
                          r'([^"“”\n]{3,400})' + q)

        def _norm(txt):
            txt = q_re.sub(
                lambda m: f'{m.group(1)}{QO} {m.group(2).strip()} {QC}', txt)
            txt = h_re.sub(
                lambda m: f'{m.group(1)}{HO} {m.group(2).strip()} {HC}', txt)
            return txt

        changed = 0
        for s in sections:
            b = s.get("body", "") or ""
            nb = _norm(b)
            if nb != b:
                s["body"] = nb
                changed += 1
        if changed:
            task.draft = "\n\n".join(
                (f"{s.get('heading', '')}\n{s.get('body', '')}").strip()
                for s in sections if (s.get("heading") or s.get("body")))
        elif task.draft:
            task.draft = _norm(task.draft)
        # validate remaining marks (skill's own check) and record honestly
        try:
            v = qh.validate_marks(task.draft or joined) or {}
        except Exception:
            v = {}
        card["islamic_marks"] = v
        if v.get("warnings"):
            mem.set_status(6, "تنسيق إسلامي: " + "؛ ".join(v["warnings"]))
        else:
            mem.set_status(6, f"تنسيق إسلامي: علامات مضبوطة ({changed} تصحيح)")

    def _style_islamic_docx(self, path, card):
        """Post-pass on the built .docx: embolden Quran verses (﴿…﴾) and hadith
        («…») using the quran_hadith_citation skill's own _set_run, preserving
        the surrounding paragraph font/size. Touches only paragraphs that carry
        a complete mark span. Fully guarded — any failure (no python-docx, read
        error) leaves the file exactly as built. No-op for non-Islamic docs.

        This is the richer Word-only step the text-level _apply_islamic_marks
        deferred: marks are already correct in every format; here the verse and
        matn also become bold, per the skill's typographic rule."""
        if not card.get("islamic"):
            return
        try:
            import importlib, sys as _sys, os as _os, re as _re
            sp = _os.path.abspath(_os.path.join(
                _os.path.dirname(__file__), "..", "capabilities", "skills",
                "quran_hadith_citation", "scripts"))
            if sp not in _sys.path:
                _sys.path.insert(0, sp)
            qh = importlib.import_module("quran_hadith")
            from docx import Document
        except Exception:
            return
        QO, QC = qh.QURAN_OPEN, qh.QURAN_CLOSE
        HO, HC = qh.HADITH_OPEN, qh.HADITH_CLOSE
        span_re = _re.compile(
            "(" + _re.escape(QO) + ".*?" + _re.escape(QC) + "|"
            + _re.escape(HO) + ".*?" + _re.escape(HC) + ")")

        def _is_span(seg):
            return ((seg.startswith(QO) and seg.endswith(QC))
                    or (seg.startswith(HO) and seg.endswith(HC)))
        try:
            doc = Document(path)
        except Exception:
            return
        changed = False
        for p in doc.paragraphs:
            txt = p.text
            if not span_re.search(txt):
                continue
            base = p.runs[0] if p.runs else None
            base_font = base.font.name if base else None
            base_size = base.font.size if base else None
            for r in list(p.runs):
                r._element.getparent().remove(r._element)
            for seg in span_re.split(txt):
                if not seg:
                    continue
                run = p.add_run(seg)
                if _is_span(seg):
                    qh._set_run(run, bold=True,
                                font=(base_font or "Kufyan Arabic"))
                else:
                    if base_font:
                        run.font.name = base_font
                if base_size:
                    run.font.size = base_size
            changed = True
        if changed:
            try:
                doc.save(path)
            except Exception:
                pass

    def _route(self, task: Task):
        """Phase 3: from the understood task_card, compute ONCE the tools &
        skills the task needs, so later layers act only on what's required
        (each tool/skill invoked only when needed — no overlap)."""
        card = task.task_card
        of = card.get("output_format", [])
        of = of if isinstance(of, list) else [of]
        text = f"{card.get('topic','')} {task.description} " \
               f"{card.get('task_type','')} {' '.join(str(x) for x in of)}"
        if self.caps:
            task.tools = [t.name for t in self.caps.match_tools(text)]
            task.skills = [s.name for s in self.caps.match_skills(text)]
        else:
            task.tools, task.skills = [], []
        # sourcing mode decides whether we gather and/or document sources
        mode = card.get("sourcing_mode", "cited")
        # always-on skills by task type
        cs = str(card.get("citation_style", "")).upper()
        # a citation-style formatter runs ONLY when sources will be documented
        if mode == "cited" and cs and cs != "UNSPECIFIED":
            task.skills.append("apa_formatter" if cs == "APA" else "mla_formatter")
        task.skills.append("arabic_rewriter"
                           if card.get("language", "ar") == "ar"
                           else "english_rewriter")
        # gather live web sources for any task that needs references — its
        # triggers rarely appear in a plain "اكتب بحثاً…", so add it explicitly.
        # "cited" and "uncited" both gather (uncited uses them to inform the
        # text but won't cite them); "none" gathers nothing.
        needs_sources = mode != "none" and (
            card.get("needs_academic_search")
            or "academic_search" in task.tools
            or (mode == "cited" and cs and cs != "UNSPECIFIED")
            or card.get("reference_count")
            or str(card.get("task_type", "")).lower() in
            ("بحث", "research", "دراسة", "report", "تقرير", "مراجعة أدبيات",
             "literature review", "analysis", "تحليل", "أطروحة", "thesis"))
        if needs_sources:
            task.tools.append("web_search")
        # genuinely academic tasks also gather peer-reviewed sources (free
        # scholarly APIs). Not for "none" mode, and not for news/recency.
        acad_types = ("بحث", "research", "دراسة", "أطروحة", "thesis",
                      "مراجعة أدبيات", "literature review", "رسالة علمية",
                      "dissertation")
        if mode != "none" and (
                card.get("needs_academic_search")
                or str(card.get("task_type", "")).lower() in acad_types):
            task.tools.append("academic_search")
        # news/recency intent enables web_search even for non-academic tasks
        # (never academic_search — news isn't academic). Respect "none" mode below.
        if mode != "none" and self._is_recency_query(
                f"{card.get('topic','')} {task.description}"):
            task.tools.append("web_search")
        if mode == "none":
            # explicit no-sources request: strip every source-gathering tool
            task.tools = [t for t in task.tools
                          if t not in ("web_search", "academic_search")]
            card.pop("needs_academic_search", None)
        task.tools = list(dict.fromkeys(task.tools))   # dedupe, keep order
        task.skills = list(dict.fromkeys(task.skills))
        self._dedupe_tools(task)                        # collapse redundant tools

    # canonical provider for every registered tool. The pipeline ACTS only on
    # the layer-4 research tools ("active"); every other capability is already
    # served by a skill or by inline code ("skill"/"inline"), and a few have no
    # wired path yet ("inactive"). This map resolves the tool/skill duplication
    # WITHOUT removing any registry entry or file — it only records where each
    # capability really runs so a matched tool is never double-counted as a
    # separate action.
    _TOOL_DELEGATION = {
        # active — consulted by layer 4 as real actions
        "web_search": ("active", None),
        "academic_search": ("active", "_scholarly_search (inline)"),
        "web_extract": ("active", None),
        "web_document": ("active", None),
        "youtube": ("active", None),
        # served by a skill (export/format/scoring) — driven by output_format /
        # the relevant layer, not by task.tools
        "word": ("skill", "docx_builder"),
        "powerpoint": ("skill", "pptx_builder"),
        "excel": ("skill", "xlsx_builder"),
        "pdf": ("skill", "pdf_builder"),
        "doc_export": ("skill", "docx/pdf/pptx/xlsx_builder"),
        "chart": ("skill", "chart_builder"),
        "credibility_check": ("skill", "credibility_scorer"),
        # served by inline code elsewhere
        "doc_read": ("inline", "web.server._extract_bytes + core.ocr"),
        "memory_store": ("inline", "TaskMemory + config/chats"),
        # csv is now a real export format served in _export
        "csv": ("inline", "_export csv branch"),
        # present in the registry but with no wired path yet
        "diagram": ("inactive", None),
        "calendar": ("inactive", None),
        "scheduler": ("inactive", None),
        "mcp_connector": ("inactive", None),
    }

    def _dedupe_tools(self, task: Task):
        """Keep in task.tools only the tools the pipeline actually acts on
        ("active"); record every other matched tool under card['tool_delegation']
        with the skill/inline path that really serves it, then drop it from the
        action list. Unknown tools are kept untouched. Additive and safe: no
        registry entry or tool file is removed."""
        deleg, kept = {}, []
        for t in (task.tools or []):
            d = self._TOOL_DELEGATION.get(t)
            if d is None:
                kept.append(t)                 # unknown → leave as-is
            elif d[0] == "active":
                kept.append(t)
            else:
                deleg[t] = {"via": d[0], "by": d[1]}   # served elsewhere
        task.tools = list(dict.fromkeys(kept))
        if deleg:
            task.task_card["tool_delegation"] = deleg

    @staticmethod
    def _current_request(text):
        """Return only the CURRENT request, dropping any threaded conversation
        history the web/terminal prefixes as
        "[سياق المحادثة السابقة] … [الطلب الحالي] …". Used so a link or verb from
        an EARLIER turn never hijacks the current one — e.g. asking to summarize
        THIS conversation must not re-summarize a video linked earlier."""
        t = text or ""
        if "[الطلب الحالي]" in t:
            return t.rsplit("[الطلب الحالي]", 1)[-1]
        return t

    @staticmethod
    def _extract_slide_count(text):
        """Extract a requested slide count from the request. None if unstated.
        Feeds design_slides so "اعمل عرض 30 شريحة" honours 30. (tested)"""
        import re
        if not text:
            return None
        patterns = [
            r'(\d{1,3})\s*(?:شريحة|شرائح|slides?|slide)',
            r'(?:شريحة|شرائح|slides?|عرض)\D{0,10}?(\d{1,3})',
            r'(?:عدد|count|number)\D{0,15}?(\d{1,3})',
        ]
        for p in patterns:
            m = re.search(p, text, re.I)
            if m:
                n = int(m.group(1))
                if 1 <= n <= 100:
                    return n
        return None

    @staticmethod
    def extract_length_target(text):
        """→ {'words': int|None, 'pages': int|None} from the request; None if
        unstated. Pages → estimated words (~500 words/academic page) when the
        word count itself isn't given. (tested)"""
        import re
        if not text:
            return {"words": None, "pages": None}
        words = pages = None
        m = re.search(r'(\d{2,6})\s*(?:كلمة|كلمات|words?|word)', text, re.I)
        if m:
            words = int(m.group(1))
        m = re.search(r'(\d{1,4})\s*(?:صفحة|صفحات|pages?|page)', text, re.I)
        if m:
            pages = int(m.group(1))
        if words is None and pages:
            words = pages * 500
        return {"words": words, "pages": pages}

    @staticmethod
    def count_words(text):
        """Deterministic word count — Arabic + Latin, after light markdown
        stripping. This is arithmetic (len of tokens), never an estimate."""
        import re
        if not text:
            return 0
        clean = re.sub(r'[#*_`>\-]+', ' ', text)
        return len([w for w in clean.split() if any(c.isalnum() for c in w)])

    @staticmethod
    def verify_sections_coverage(required_titles, sections):
        """Return the required titles NOT covered by any written section
        (flexible two-way substring match). (tested)"""
        present = [(s.get("heading") or "").strip() for s in sections
                   if (s.get("heading") or s.get("body"))]
        present_l = [p.lower() for p in present if p]
        joined = " ".join(present_l)
        missing = []
        for req in required_titles:
            r = (req or "").strip().lower()
            if not r:
                continue
            if r in joined:
                continue
            if any(r in p or p in r for p in present_l):
                continue
            missing.append(req)
        return missing

    @staticmethod
    def _detect_youtube_intent(text):
        """Keyword heuristic for what to do with a YouTube link. Returns
        {"mode", "with_timing", "explicit"} where `explicit` is True only when a
        clear transcript/summary signal was found (so the caller can fall back to
        the model for genuinely ambiguous phrasings). Broad synonym lists so most
        natural wordings resolve here without an extra model call."""
        t = " " + (text or "").lower() + " "
        transcript_kw = (
            "فرّغ", "فرغ", "فرّغلي", "فرغلي", "تفريغ", "التفريغ", "فرِّغ",
            "النص الكامل", "النص كامل", "كامل النص", "نص الفيديو", "نصّ الفيديو",
            "النص الحرفي", "النص حرفي", "حرفي", "حرفياً", "حرفيا",
            "كلمة بكلمة", "كلمة كلمة", "انسخ النص", "انسخ الكلام", "انسخ لي النص",
            "اكتب ما قيل", "اكتب ما قِيل", "اكتب النص", "اكتب الكلام", "اكتب كل ما",
            "المحتوى النصي", "سكربت", "السكربت", "سكريبت", "السكريبت",
            "transcribe", "transcript", "full text", "verbatim",
            "word for word", "word-for-word", "captions", "subtitles", "script",
        )
        summary_kw = (
            "لخّص", "لخص", "لخّصلي", "لخصلي", "لخّص لي", "لخص لي", "ملخص", "ملخّص",
            "الملخص", "الملخّص", "اختصر", "اختصار", "باختصار", "أهم النقاط",
            "اهم النقاط", "أهم ما", "اهم ما", "النقاط الرئيسية", "النقاط المهمة",
            "الأفكار الرئيسية", "الافكار الرئيسية", "الخلاصة", "خلاصة", "زبدة",
            "لبّ الموضوع", "لب الموضوع", "عن ماذا يتحدث", "عن ماذا يتكلم",
            "ماذا يقول", "وش يقول", "ايش يقول", "شنو يقول", "فكرة الفيديو",
            "summary", "summarize", "summarise", "tl;dr", "tldr", "key points",
            "main points", "gist", "overview", "in short",
        )
        timing_kw = (
            "مع التوقيت", "مع التوقيتات", "بالتوقيت", "بالتوقيتات", "التوقيت",
            "توقيت", "الطوابع الزمنية", "طوابع زمنية", "الطابع الزمني",
            "الدقائق", "الدقيقة", "الثواني", "بالدقائق", "الوقت لكل",
            "timestamp", "timestamps", "with timing", "with time", "with times",
            "time codes", "timecodes", "time-stamps",
        )
        has_t = any(k in t for k in transcript_kw)
        has_s = any(k in t for k in summary_kw)
        with_timing = any(k in t for k in timing_kw)
        if has_t and has_s:
            mode = "both"
        elif has_t:
            mode = "transcript"
        else:
            mode = "summary"          # summary-only OR ambiguous default
        return {"mode": mode, "with_timing": with_timing,
                "explicit": bool(has_t or has_s)}

    def _classify_youtube_intent_llm(self, text):
        """Ask the model to classify a YouTube request into
        summary / transcript / timing — understands ANY phrasing. Returns
        {"mode", "with_timing"} or None when unavailable or unusable. Used only
        when the keyword heuristic is not decisive, to keep model calls rare."""
        if not self.llm_fn:
            return None
        try:
            from core.llm import extract_json
            prompt = (
                "صنّف طلب المستخدم المتعلّق بفيديو يوتيوب. قد يريد المستخدم: "
                "ملخصاً (summary)، أو تفريغاً حرفياً كاملاً للنص (transcript)، أو "
                "كليهما معاً، وقد يريد طوابع زمنية/توقيتاً (timing). افهم أيّ صياغة "
                "مهما اختلفت اللهجة أو الأسلوب. أعِد JSON فقط بلا أي نص آخر بالشكل:\n"
                '{"summary": true|false, "transcript": true|false, '
                '"timing": true|false}\n\nطلب المستخدم:\n' + (text or "")[:800]
            )
            data = extract_json(self.llm_fn(prompt, system=self.system_main,
                                            temperature=0.0)) or {}
            s = bool(data.get("summary"))
            tr = bool(data.get("transcript"))
            tm = bool(data.get("timing"))
            if not s and not tr:
                return None
            mode = "both" if (s and tr) else ("transcript" if tr else "summary")
            return {"mode": mode, "with_timing": tm}
        except Exception:
            return None

    async def _layer_3(self, task: Task, mem: TaskMemory):
        """٣: الفهم — تحليل المهمة وبناء بطاقتها ثم توجيه الأدوات/المهارات."""
        task.status = TaskStatus.LAYER_3
        mem.set_status(3, "تحليل المهمة")

        # ── YouTube link → dedicated transcript path, BEFORE the model. Builds a
        #    minimal card and returns early, so the video is summarized/
        #    transcribed instead of being treated as a research task (which
        #    produced empty مباحث/مطالب). Fully guarded: any failure, or a link
        #    with no captions, continues the normal path unchanged.
        try:
            import re as _re
            from capabilities.tools import tool_youtube as _yt
            # Look for a YouTube link (and read its intent) in the CURRENT request
            # ONLY — never the threaded conversation history. Otherwise a video
            # linked in an EARLIER turn hijacks an unrelated request now (e.g.
            # "لخّص ما قمنا بعمله في هذه المحادثة" would re-summarize that video
            # instead of the conversation).
            _cur = self._current_request(task.description)
            yurl = None
            for _u in _re.findall(r"https?://\S+", _cur):
                _u = _u.rstrip('.,)"\'،؛')
                if _yt.is_youtube_url(_u):
                    yurl = _u
                    break
            # Follow-up with NO link in the current message but a clear video
            # request ("فرّغه"، "لخّص الفيديو"، "أريده بالكامل") → reuse the LAST
            # YouTube link from the conversation. Guarded so a conversation-summary
            # request never re-summarizes an old video.
            if not yurl:
                _cl = " " + _cur.lower() + " "
                _refers_video = (self._detect_youtube_intent(_cur).get("explicit")
                                 or any(k in _cl for k in (
                                     "الفيديو", "الفيدو", "المقطع", "الحلقة",
                                     "المحاضرة", "بالكامل", "الكامل", " كامل ",
                                     " video ", " full ", " complete ")))
                if _refers_video:
                    _last = None
                    for _u in _re.findall(r"https?://\S+", task.description or ""):
                        _u = _u.rstrip('.,)"\'،؛')
                        if _yt.is_youtube_url(_u):
                            _last = _u        # keep the LAST match
                    yurl = _last
            if yurl:
                kw = self._detect_youtube_intent(_cur)
                if kw.get("explicit"):
                    intent = {"mode": kw["mode"],
                              "with_timing": kw["with_timing"]}
                else:
                    # ambiguous wording → let the model understand ANY phrasing;
                    # fall back to the heuristic (default summary) if unavailable.
                    intent = (self._classify_youtube_intent_llm(_cur)
                              or {"mode": kw["mode"],
                                  "with_timing": kw["with_timing"]})
                # ALWAYS fetch the video's ORIGINAL captions (never auto-translated
                # Arabic) — the transcript stays in the video's own language; the
                # summary is translated by the model per the output-language rule.
                res = await _yt.run({"url": yurl, "lang": "ar",
                                     "with_timing": intent["with_timing"],
                                     "prefer_original": True})
                if getattr(res, "ok", False) and (res.data or {}).get("text"):
                    # Output language for the SUMMARY: an explicit request wins;
                    # "source/original" → the video's own language; otherwise the
                    # language the request itself is written in (never assumed).
                    _lk, _ln = self._requested_output_lang(_cur)
                    task.task_card = {
                        "task_type": "youtube_" + intent["mode"],
                        "topic": "تلخيص/تفريغ فيديو يوتيوب",
                        "language": ("ar" if (_lk == "name" and _ln == "العربية")
                                     or (_lk is None
                                         and self._detect_lang(_cur) == "ar")
                                     else "en"),
                        "output_format": ["INLINE"],
                        "sourcing_mode": "none",
                        "youtube": {
                            "url": yurl, "mode": intent["mode"],
                            "with_timing": intent["with_timing"],
                            "transcript": res.data["text"],
                            "transcript_plain": (res.data or {}).get(
                                "text_plain", ""),
                            "out_lang_kind": _lk,      # None | "name" | "source"
                            "out_lang_name": _ln,      # e.g. "العربية"/"English"
                            "req_lang": self._detect_lang(_cur),
                        },
                    }
                    task.skills = []
                    task.tools = []
                    mem.set_status(3, f"يوتيوب: {intent['mode']} — نص "
                                      f"{len(res.data['text'])} حرف")
                    return
                # Fetch failed (no captions / throttled / library missing).
                # If the request is ALSO a research task (has other triggers),
                # let the normal path proceed. But for a pure YouTube ask, stay
                # in the youtube lane with an honest message instead of emitting
                # a spurious research report (the original complaint).
                _t = _cur.lower()
                _also_research = any(k in _t for k in _TASK_TRIGGERS)
                if not _also_research:
                    _why = getattr(res, "error", "") or "لا يوجد نص/ترجمة متاح"
                    task.task_card = {
                        "task_type": "youtube_" + intent["mode"],
                        "topic": "تفريغ/تلخيص فيديو يوتيوب",
                        "language": "ar",
                        "output_format": ["INLINE"],
                        "sourcing_mode": "none",
                        "youtube": {
                            "url": yurl, "mode": intent["mode"],
                            "with_timing": intent["with_timing"],
                            "transcript": "",
                            "error": str(_why),
                        },
                    }
                    task.skills = []
                    task.tools = []
                    mem.set_status(3, f"يوتيوب: تعذّر جلب النص ({_why})")
                    return
                mem.set_status(3, "يوتيوب: لا نص متاح — المسار العادي")
        except Exception:
            pass

        if self.llm_fn:
            from pipeline.prompts import PROMPT_LAYER_3_UNDERSTAND
            from core.llm import extract_json
            prompt = PROMPT_LAYER_3_UNDERSTAND.format(
                task_description=task.description)
            try:
                raw = self.llm_fn(prompt, system=self.system_main,
                                  temperature=0.2)
                task.task_card = extract_json(raw)
            except Exception as e:
                mem.set_status(3, f"فهم (تخطّي للنموذج: {e})")
                task.task_card = self._placeholder_card(task)
        else:
            task.task_card = self._placeholder_card(task)

        # normalize output_format to a list (prompt/canonical form)
        of = task.task_card.get("output_format")
        if isinstance(of, str):
            task.task_card["output_format"] = [of]
        elif not of:
            task.task_card["output_format"] = ["DOCX"]
        # an explicit format in the request wins (احفظه Word/PDF/بوربوينت/اكسل)
        _rf = self._requested_format(self._current_request(task.description))
        if _rf:
            task.task_card["output_format"] = [_rf]

        # Output-language rule: an explicit request wins; otherwise the language
        # of the task INSTRUCTIONS (the current request text), never the chat
        # language. The "follow the pasted file/link" case is finalized in
        # _read_pasted_urls once the file's own language is known.
        try:
            _cur3 = self._current_request(task.description)
            _lk3, _ln3 = self._requested_output_lang(_cur3)
            if _lk3 == "name":
                task.task_card["language"] = ("ar" if _ln3 == "العربية" else "en")
                task.task_card["lang_locked"] = True
            elif _lk3 == "source":
                task.task_card["lang_locked"] = "source"   # resolved from content
            else:
                task.task_card["language"] = self._detect_lang(_cur3)
        except Exception:
            pass

        # Safety net: never let an INJECTED context header (memory / attachment /
        # instruction block) become the document topic or filename. If the model
        # picked one up, fall back to the current request text.
        try:
            _tp = str(task.task_card.get("topic") or "").strip()
            _bad = ("مهام سابقة", "محتوى الملف", "تعليمات مهمة", "ذاكرة:",
                    "للسياق فقط", "الملفات المرفقة", "reference only",
                    "for context only")
            if _tp.startswith("[") or any(k in _tp for k in _bad):
                _clean = self._current_request(task.description).strip()
                # keep only the part before any injected marker
                for _m in ("\n[", "[محتوى الملف", "[للسياق", "[ذاكرة",
                           "[تعليم", "[سياق المحادثة"):
                    _i = _clean.find(_m)
                    if _i > 0:
                        _clean = _clean[:_i].strip()
                if _clean:
                    task.task_card["topic"] = _clean[:150]
        except Exception:
            pass

        # how the user wants sourcing handled (cited / uncited / none). Detected
        # from the RAW request so an explicit "بدون مصادر" / "دون توثيقها" is
        # honoured even if the model didn't surface it in the card.
        task.task_card["sourcing_mode"] = self._sourcing_mode(task.description)

        # scope that LIMITS the task: references-only / outline-only / part-only.
        task.task_card["scope"] = self._task_scope(
            self._current_request(task.description))

        # requested slide count (e.g. "اعمل عرض 30 شريحة") → reaches
        # design_slides. Absent → None → default behaviour unchanged.
        try:
            _sc = self._extract_slide_count(task.description)
            if _sc and isinstance(task.task_card, dict):
                task.task_card["slide_count"] = _sc
        except Exception:
            pass

        # requested length (words/pages) → recorded for the writer/export and
        # read by layer 6.6. Absent → nothing set → behaviour unchanged.
        try:
            _lt = self.extract_length_target(task.description)
            if isinstance(task.task_card, dict):
                if _lt.get("words"):
                    task.task_card["target_words"] = _lt["words"]
                if _lt.get("pages"):
                    task.task_card["target_pages"] = _lt["pages"]
        except Exception:
            pass

        # Phase 3: route tools & skills once
        self._route(task)

        # references-only must actually gather sources → force the search tools.
        if task.task_card.get("scope") == "references":
            for _t in ("web_search", "academic_search"):
                if _t not in task.tools:
                    task.tools.append(_t)
            task.task_card["needs_academic_search"] = True

    async def _layer_4(self, task: Task, mem: TaskMemory):
        """٤: البحث — أكاديمي (PaperQA) + بحث ويب حي (SearXNG). يُشغَّل ما وُجّهت
        إليه طبقة الفهم فقط، ونتائج الويب تُخزَّن كمصادر للطبقتين ٥ و٦."""
        task.status = TaskStatus.LAYER_4
        # read any URL pasted in the request as a PRIMARY source (page/YouTube),
        # regardless of routing — so "لخّص هذا الفيديو <url>" reads the video.
        await self._read_pasted_urls(task, mem)
        # when the user pasted a link, FOCUS on it: a link summary is never an
        # academic-paper task (skip academic), and if we actually read the link
        # we skip the generic topic web search too (it only adds off-topic
        # "how to summarize videos" noise). If the link couldn't be read, the
        # web search stays as a fallback.
        pasted_present = bool(task.task_card.get("pasted_present"))
        pasted_ok = bool(task.task_card.get("pasted_reads"))
        want_academic = (("academic_search" in task.tools
                          or task.task_card.get("needs_academic_search"))
                         and not pasted_present)
        want_web = ("web_search" in task.tools) and not pasted_ok
        if not want_academic and not want_web:
            if pasted_ok:
                mem.set_status(4, "الاعتماد على الرابط المُدرَج — تخطّي البحث العام")
            else:
                mem.set_status(4, "لا يلزم بحث — تخطّي")
            return
        if want_academic:
            from pipeline.layers.layer_4_research import run as _layer4_run
            await _layer4_run(task, mem)          # PaperQA (if installed)
            await self._academic_search(task, mem)  # free scholarly APIs
        if want_web:
            await self._web_search(task, mem)

    @staticmethod
    def _is_recency_query(text) -> bool:
        """True when the request wants fresh/current information (news, latest,
        today…), in Arabic or English. Used to switch on live web search and
        the recency-oriented ranking."""
        t = " " + (text or "").lower() + " "
        kws = ("أخبار", "آخر", "اليوم", "الآن", "حالياً", "حاليا", "مؤخراً",
               "مؤخرا", "أحدث", "جديد", "هذا الأسبوع", "هذا الشهر",
               "news", "latest", "today", "now", "recent", "breaking",
               "current", "this week", "this month")
        return any(k in t for k in kws)

    @staticmethod
    def _search_directives(text):
        """Explicit search constraints in the request: (site_domain, df).
        `site_domain` restricts results to one site; `df` is a date range
        (d/w/m/y). Either may be None."""
        import re
        tl = (text or "").lower()
        site = None
        m = re.search(r'site:\s*([a-z0-9.\-]+\.[a-z]{2,})', tl)
        if m:
            site = m.group(1)
        else:
            m = re.search(r'(?:موقع|from|on)\s+'
                          r'([a-z0-9\-]+\.[a-z]{2,}(?:\.[a-z]{2,})?)', tl)
            if m:
                site = m.group(1)
        df = None
        if any(k in tl for k in ("اليوم", "today", "آخر يوم", "اخر يوم",
                                 "past day", "last 24")):
            df = "d"
        elif any(k in tl for k in ("هذا الأسبوع", "هذا الاسبوع", "آخر أسبوع",
                                   "اخر اسبوع", "this week", "past week")):
            df = "w"
        elif any(k in tl for k in ("هذا الشهر", "آخر شهر", "اخر شهر",
                                   "this month", "past month")):
            df = "m"
        elif any(k in tl for k in ("هذا العام", "هذه السنة", "آخر سنة",
                                   "اخر سنة", "this year", "past year")):
            df = "y"
        return (site, df)

    @staticmethod
    def _augment_query_with_date(query, lang):
        """Append the current year (and month for daily/'today' intent) to a
        query so engines favour fresh results. The year/month are computed
        dynamically from datetime.now() — never hard-coded."""
        import datetime
        q = (query or "").strip()
        if not q:
            return q
        now = datetime.datetime.now()
        year = str(now.year)
        if year in q:
            return q                     # already dated — leave it
        daily = any(w in q.lower() for w in
                    ("اليوم", "الآن", "today", "now", "breaking", "عاجل"))
        if daily:
            if lang == "ar":
                months_ar = ["يناير", "فبراير", "مارس", "أبريل", "مايو",
                             "يونيو", "يوليو", "أغسطس", "سبتمبر", "أكتوبر",
                             "نوفمبر", "ديسمبر"]
                return q + " " + months_ar[now.month - 1] + " " + year
            return q + " " + now.strftime("%B") + " " + year
        return q + " " + year

    @staticmethod
    def _sort_results_by_recency(results):
        """Stable-sort results newest-first. Scores each by an explicit year
        (2020-2030) and relative-time phrases ('قبل ساعة/يوم', 'hours/days ago')
        in title+content; items with no time signal keep their relative order
        (stable sort). Never raises."""
        import re
        import datetime  # noqa: F401 (kept for clarity/extension)
        if not results:
            return results

        def score(r):
            try:
                t = (str(r.get("title", "")) + " "
                     + str(r.get("content", ""))).lower()
            except Exception:
                return 0
            best = 0
            for m in re.findall(r"\b(20[2-3]\d)\b", t):
                y = int(m)
                if 2020 <= y <= 2030:
                    best = max(best, 1000 + (y - 2000) * 10)
            rel = (
                (r"(?:قبل|منذ)\s*(?:دقيقة|دقائق|ساعة|ساعات)", 1400),
                (r"\b(?:minute|hour)s?\s+ago\b|just now", 1400),
                (r"(?:قبل|منذ)\s*(?:يوم|يومين|أيام)", 1350),
                (r"\bday(?:s)?\s+ago\b|yesterday|أمس", 1330),
                (r"(?:قبل|منذ)\s*(?:أسبوع|أسابيع)", 1300),
                (r"\bweek(?:s)?\s+ago\b", 1300),
                (r"(?:قبل|منذ)\s*(?:شهر|أشهر)", 1200),
                (r"\bmonth(?:s)?\s+ago\b", 1200),
            )
            for pat, w in rel:
                if re.search(pat, t):
                    best = max(best, w)
            return best
        # sorted() is stable, so equal-score items keep their original order
        return sorted(results, key=lambda r: -score(r))

    @staticmethod
    def _searx_query(instance: str, query: str, lang: str, limit: int,
                     timeout: int = 8, time_range=None, sort_by_date=False):
        """Direct SearXNG JSON search. Returns a list of {title,url,content}
        or None when the instance is unreachable / returns nothing usable.

        Optional (backward compatible — old calls without them still work):
          * time_range: "day"/"week"/"month"/"year" → SearXNG time filter.
          * sort_by_date: intent flag kept for callers; SearXNG's JSON API has
            no direct sort param, so newest-first ordering is applied later by
            _sort_results_by_recency.
        Uses _http_get so it survives broken device DNS (DoH retry)."""
        import urllib.parse
        import json as _json
        instance = (instance or "").rstrip("/")
        if not instance:
            return None
        params = {"q": query, "format": "json", "categories": "general"}
        if lang:
            params["language"] = lang
        if time_range in ("day", "week", "month", "year"):
            params["time_range"] = time_range
        url = instance + "/search?" + urllib.parse.urlencode(params)
        # a real browser UA: many public SearXNG instances reject bot-like UAs
        raw = WeaverOrchestrator._http_get(
            url, {"User-Agent": ("Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                                 "AppleWebKit/537.36 (KHTML, like Gecko) "
                                 "Chrome/122.0.0.0 Safari/537.36"),
                  "Accept": "application/json"}, timeout)
        if not raw:
            return None
        try:
            data = _json.loads(raw)
        except Exception:
            return None
        return [{"title": r.get("title", ""), "url": r.get("url", ""),
                 "content": r.get("content", "")}
                for r in (data.get("results", []) or [])[:limit]]

    @staticmethod
    def _doh_resolve(host: str, timeout: int = 10):
        """Resolve a hostname to an IPv4 via DNS-over-HTTPS, using resolvers
        addressed BY IP — so it needs NO working system DNS. This is the fix
        for Termux/Android where getaddrinfo fails with 'No address associated
        with hostname' even though HTTPS itself works. Returns an IP or None."""
        import urllib.parse
        import urllib.request
        import json as _json
        if not host:
            return None
        # Cloudflare (1.1.1.1) and Google (8.8.8.8) both present valid certs for
        # their own IPs, so https-by-IP validates without any name lookup.
        for base in ("https://1.1.1.1/dns-query",
                     "https://8.8.8.8/resolve",
                     "https://1.0.0.1/dns-query"):
            try:
                url = base + "?name=" + urllib.parse.quote(host) + "&type=A"
                req = urllib.request.Request(
                    url, headers={"accept": "application/dns-json"})
                with urllib.request.urlopen(req, timeout=timeout) as r:
                    data = _json.loads(r.read().decode("utf-8"))
                for ans in (data.get("Answer") or []):
                    if ans.get("type") == 1 and ans.get("data"):
                        return str(ans["data"]).strip()
            except Exception:
                continue
        return None

    @classmethod
    def _http_get(cls, url: str, headers: dict, timeout: int = 15):
        """HTTP GET that survives broken system DNS. Tries the normal resolver
        first; on a name-resolution failure it resolves the host via DoH
        (_doh_resolve) and retries by pinning getaddrinfo to that IP (TLS SNI /
        Host stay correct because the hostname is preserved). Returns decoded
        text, or None on any failure."""
        import urllib.request
        import urllib.error
        import urllib.parse
        import socket
        req = urllib.request.Request(url, headers=headers or {})
        try:
            with urllib.request.urlopen(req, timeout=timeout) as r:
                return r.read().decode("utf-8", "replace")
        except Exception as e:
            reason = getattr(e, "reason", None)
            msg = (str(reason) + " " + str(e)).lower()
            is_dns = isinstance(reason, socket.gaierror) or isinstance(
                e, socket.gaierror) or ("address associated" in msg) or (
                "name or service" in msg) or ("name resolution" in msg) or (
                "getaddrinfo" in msg)
            if not is_dns:
                return None
        # DNS path: resolve via DoH and pin it
        host = urllib.parse.urlparse(url).hostname
        ip = cls._doh_resolve(host) if host else None
        if not ip:
            return None
        orig = socket.getaddrinfo

        def _pinned(h, p, *a, **k):
            if h == host:
                return [(socket.AF_INET, socket.SOCK_STREAM, 6, "", (ip, p))]
            return orig(h, p, *a, **k)

        socket.getaddrinfo = _pinned
        try:
            with urllib.request.urlopen(req, timeout=timeout) as r:
                return r.read().decode("utf-8", "replace")
        except Exception:
            return None
        finally:
            socket.getaddrinfo = orig

    @staticmethod
    def _ddg_search(query: str, lang: str, limit: int, timeout: int = 12,
                    df=None):
        """Direct DuckDuckGo search — NO server required (works on the phone as
        is). Hits the html.duckduckgo.com endpoint over plain HTTP, decoding
        DDG's redirect links. Prefers UniWeb/curl_impersonate (real browser
        fingerprint, beats bot-blocking) and falls back to urllib. Returns a
        list of {title,url,content} or None when nothing usable comes back.

        Optional (backward compatible): df is DuckDuckGo's time filter —
        'd' (day), 'w' (week), 'm' (month), 'y' (year)."""
        import urllib.parse
        import urllib.request
        import html as _html
        import re
        q = (query or "").strip()
        if not q:
            return None
        params = {"q": q}
        if lang == "ar":
            params["kl"] = "xa-ar"      # region/language hint (best-effort)
        if df in ("d", "w", "m", "y"):
            params["df"] = df           # recency filter
        endpoint = ("https://html.duckduckgo.com/html/?"
                    + urllib.parse.urlencode(params))
        ua = ("Mozilla/5.0 (Linux; Android 13; Pixel 7) AppleWebKit/537.36 "
              "(KHTML, like Gecko) Chrome/122.0.0.0 Mobile Safari/537.36")
        raw = None
        # 1) UniWeb (curl_impersonate) — best chance against anti-bot
        try:
            import os as _os, sys as _sys
            uw = _os.path.abspath(_os.path.join(
                _os.path.dirname(__file__), "..", "engines", "uniweb-core"))
            if uw not in _sys.path:
                _sys.path.insert(0, uw)
            import uniweb as _uniweb
            html = _uniweb.fetch(endpoint)
            if html and isinstance(html, str) and "result" in html:
                raw = html
        except Exception:
            raw = None
        # 2) urllib fallback — DNS-safe (works even when the phone's resolver
        #    fails: it retries over DNS-over-HTTPS). This is the path that makes
        #    search work on Termux/Android with broken getaddrinfo.
        if not raw:
            raw = WeaverOrchestrator._http_get(endpoint, {
                "User-Agent": ua, "Accept": "text/html",
                "Accept-Language": ("ar,en;q=0.8" if lang == "ar"
                                    else "en-US,en;q=0.8")}, timeout)
        if not raw:
            return None

        def _clean(s: str) -> str:
            return _html.unescape(re.sub(r"\s+", " ", re.sub(r"<[^>]+>", "", s))).strip()

        def _real_url(href: str) -> str:
            href = _html.unescape(href)
            m = re.search(r"[?&]uddg=([^&]+)", href)
            if m:
                return urllib.parse.unquote(m.group(1))
            if href.startswith("//"):
                return "https:" + href
            return href

        results = []
        for m in re.finditer(
                r'<a[^>]+class="[^"]*result__a[^"]*"[^>]+href="([^"]+)"[^>]*>'
                r'(.*?)</a>', raw, re.S):
            url = _real_url(m.group(1))
            if not url.startswith("http"):
                continue
            results.append({"title": _clean(m.group(2)), "url": url,
                            "content": ""})
            if len(results) >= limit:
                break
        # attach snippets (aligned by document order, best-effort)
        snips = re.findall(r'class="result__snippet"[^>]*>(.*?)</a>', raw, re.S)
        for i, s in enumerate(snips[:len(results)]):
            results[i]["content"] = _clean(s)
        return results or None

    @staticmethod
    def _bing_search(query: str, lang: str, limit: int, timeout: int = 12):
        """Serverless Bing HTML search. Returns {title,url,content} list or None.
        Bing has a huge, independent index, so it broadens results beyond DDG.
        Degrades safely (parse/network failure → None). DNS-safe via _http_get."""
        import urllib.parse
        import html as _html
        import re
        q = (query or "").strip()
        if not q:
            return None
        params = {"q": q}
        if lang == "ar":
            params["setlang"] = "ar"
            params["cc"] = "XA"
        url = "https://www.bing.com/search?" + urllib.parse.urlencode(params)
        ua = ("Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
              "(KHTML, like Gecko) Chrome/122.0.0.0 Safari/537.36")
        raw = WeaverOrchestrator._http_get(
            url, {"User-Agent": ua, "Accept": "text/html",
                  "Accept-Language": ("ar,en;q=0.8" if lang == "ar"
                                      else "en-US,en;q=0.8")}, timeout)
        if not raw:
            return None

        def _clean(s):
            return _html.unescape(re.sub(r"\s+", " ",
                                         re.sub(r"<[^>]+>", "", s or ""))).strip()
        results = []
        # tolerate nested tags between <h2> and its <a> (spans, etc.)
        for blk in re.findall(r'<li class="b_algo".*?</li>', raw, re.S):
            a = (re.search(r'<h2[^>]*>.*?<a[^>]+href="([^"]+)"[^>]*>(.*?)</a>',
                           blk, re.S)
                 or re.search(r'<a[^>]+class="[^"]*tilk[^"]*"[^>]+'
                              r'href="([^"]+)"[^>]*>(.*?)</a>', blk, re.S))
            if not a:
                continue
            u = _html.unescape(a.group(1))
            title = _clean(a.group(2))
            if not u.startswith("http") or not title:
                continue
            if "bing.com" in u or "microsofttranslator" in u:
                continue
            cap = (re.search(r'<p class="b_[^"]*"[^>]*>(.*?)</p>', blk, re.S)
                   or re.search(r'<p[^>]*>(.*?)</p>', blk, re.S))
            results.append({"title": title, "url": u,
                            "content": _clean(cap.group(1) if cap else "")})
            if len(results) >= limit:
                break
        return results or None

    @staticmethod
    def _mojeek_search(query: str, lang: str, limit: int, timeout: int = 10):
        """Serverless Mojeek HTML search. Mojeek has its OWN independent crawler
        (not Google/Bing), so it genuinely broadens results. Returns
        {title,url,content} list or None; degrades safely. DNS-safe via
        _http_get."""
        import urllib.parse
        import html as _html
        import re
        q = (query or "").strip()
        if not q:
            return None
        url = "https://www.mojeek.com/search?" + urllib.parse.urlencode({"q": q})
        ua = ("Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
              "(KHTML, like Gecko) Chrome/122.0.0.0 Safari/537.36")
        raw = WeaverOrchestrator._http_get(
            url, {"User-Agent": ua, "Accept": "text/html",
                  "Accept-Language": ("ar,en;q=0.8" if lang == "ar"
                                      else "en-US,en;q=0.8")}, timeout)
        if not raw:
            return None

        def _clean(s):
            return _html.unescape(re.sub(r"\s+", " ",
                                         re.sub(r"<[^>]+>", "", s or ""))).strip()
        mc = re.search(r'<ul class="results-standard">(.*?)</ul>', raw, re.S)
        body = mc.group(1) if mc else raw
        results = []
        for blk in re.findall(r'<li[^>]*>(.*?)</li>', body, re.S):
            # the real result title is an <a class="title"> (or the <h2> anchor)
            a = (re.search(r'<a[^>]+class="[^"]*title[^"]*"[^>]+'
                           r'href="(https?://[^"]+)"[^>]*>(.*?)</a>', blk, re.S)
                 or re.search(r'<h2[^>]*>\s*<a[^>]+href="(https?://[^"]+)"'
                              r'[^>]*>(.*?)</a>', blk, re.S))
            if not a:
                continue
            u, title = a.group(1), _clean(a.group(2))
            if not title or "mojeek.com" in u:      # skip logo/nav/empty rows
                continue
            snip = re.search(r'<p class="s"[^>]*>(.*?)</p>', blk, re.S)
            results.append({"title": title, "url": u,
                            "content": _clean(snip.group(1) if snip else "")})
            if len(results) >= limit:
                break
        return results or None

    @staticmethod
    def _startpage_search(query: str, lang: str, limit: int, timeout: int = 8):
        """Serverless Startpage HTML search (Google results, privacy proxy).
        BEST-EFFORT: Startpage has strong anti-bot protection, so it often
        returns nothing — that's fine, it simply contributes no results.
        DNS-safe via _http_get; never raises."""
        import urllib.parse
        import html as _html
        import re
        q = (query or "").strip()
        if not q:
            return None
        url = ("https://www.startpage.com/sp/search?"
               + urllib.parse.urlencode({"query": q}))
        ua = ("Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
              "(KHTML, like Gecko) Chrome/122.0.0.0 Safari/537.36")
        raw = WeaverOrchestrator._http_get(
            url, {"User-Agent": ua, "Accept": "text/html",
                  "Accept-Language": ("ar,en;q=0.8" if lang == "ar"
                                      else "en-US,en;q=0.8")}, timeout)
        if not raw:
            return None

        def _clean(s):
            return _html.unescape(re.sub(r"\s+", " ",
                                         re.sub(r"<[^>]+>", "", s or ""))).strip()
        results = []
        for m in re.finditer(
                r'<a[^>]+class="[^"]*result[-_]?(?:link|title)[^"]*"[^>]+'
                r'href="(https?://[^"]+)"[^>]*>(.*?)</a>', raw, re.S):
            u = _html.unescape(m.group(1))
            if "startpage.com" in u:
                continue
            results.append({"title": _clean(m.group(2)), "url": u,
                            "content": ""})
            if len(results) >= limit:
                break
        snips = re.findall(r'<p[^>]+class="[^"]*description[^"]*"[^>]*>(.*?)</p>',
                           raw, re.S)
        for i, s in enumerate(snips[:len(results)]):
            results[i]["content"] = _clean(s)
        return results or None

    @classmethod
    def _multi_engine_search(cls, query, lang, limit, df=None):
        """SearXNG-like breadth WITHOUT a server: query several independent
        engines IN PARALLEL (DuckDuckGo + Bing + Mojeek + Startpage) and
        merge/dedupe by URL, interleaved so the mix stays diverse. Parallel, so
        total time ≈ the slowest engine, not the sum. Each engine degrades
        safely — a dead/blocked one just contributes nothing. Returns a merged
        list or None."""
        import itertools
        import concurrent.futures as _cf
        engines = [
            ("ddg", lambda: cls._ddg_search(query, lang, limit, df=df)),
            ("bing", lambda: cls._bing_search(query, lang, limit)),
            ("mojeek", lambda: cls._mojeek_search(query, lang, limit)),
            ("startpage", lambda: cls._startpage_search(query, lang, limit)),
        ]
        results_by_name = {}
        ex = _cf.ThreadPoolExecutor(max_workers=len(engines))
        try:
            futmap = {ex.submit(fn): name for name, fn in engines}
            done, _pending = _cf.wait(futmap, timeout=16)
            for fut, name in futmap.items():
                if fut in done:
                    try:
                        results_by_name[name] = fut.result() or []
                    except Exception:
                        results_by_name[name] = []
                else:
                    results_by_name[name] = []       # too slow → skip
        except Exception:
            # last-resort sequential fallback
            for name, fn in engines:
                try:
                    results_by_name[name] = fn() or []
                except Exception:
                    results_by_name[name] = []
        finally:
            ex.shutdown(wait=False, cancel_futures=True)
        lists = [results_by_name.get(name, []) for name, _ in engines]
        merged, seen = [], set()
        for r in itertools.chain.from_iterable(
                itertools.zip_longest(*lists)) if lists else []:
            if not r:
                continue
            u = (r.get("url") or "").split("#")[0].rstrip("/")
            if not u or u in seen or not (r.get("title") or "").strip():
                continue                              # drop dupes + empty rows
            seen.add(u)
            merged.append(r)
            if len(merged) >= max(limit, 8):
                break
        return merged or None

    # ── free academic sources (no API key) ────────────────────────────────
    # Each returns a list of {title,url,content,authors,year,doi,source} or
    # None; all go through _http_get (DoH-safe) and degrade safely. They index
    # Arabic works too, so Arabic queries return Arabic papers.
    _ACAD_UA = "WeaverWrite/1.0 (mailto:research@weaver.local)"

    @staticmethod
    def _openalex_search(query, lang, limit=6, timeout=12):
        import urllib.parse
        import json as _json
        q = (query or "").strip()
        if not q:
            return None
        url = "https://api.openalex.org/works?" + urllib.parse.urlencode(
            {"search": q, "per_page": min(int(limit) or 6, 10)})
        raw = WeaverOrchestrator._http_get(
            url, {"User-Agent": WeaverOrchestrator._ACAD_UA,
                  "Accept": "application/json"}, timeout)
        if not raw:
            return None
        try:
            data = _json.loads(raw)
        except Exception:
            return None
        out = []
        for w in (data.get("results") or [])[:limit]:
            title = w.get("title") or w.get("display_name") or ""
            if not title:
                continue
            doi = (w.get("doi") or "").replace("https://doi.org/", "")
            oa = (w.get("open_access") or {}).get("oa_url")
            url_ = oa or w.get("doi") or w.get("id") or ""
            auths = [(a.get("author") or {}).get("display_name", "")
                     for a in (w.get("authorships") or [])[:4]]
            abx = ""
            inv = w.get("abstract_inverted_index")
            if isinstance(inv, dict):
                pos = {}
                for word, ps in inv.items():
                    for p in ps:
                        pos[p] = word
                abx = " ".join(pos[k] for k in sorted(pos))[:400]
            out.append({"title": title, "url": url_, "content": abx,
                        "authors": [a for a in auths if a],
                        "year": str(w.get("publication_year") or ""),
                        "doi": doi, "source": "openalex"})
        return out or None

    @staticmethod
    def _crossref_search(query, lang, limit=6, timeout=12):
        import urllib.parse
        import json as _json
        import re
        q = (query or "").strip()
        if not q:
            return None
        url = "https://api.crossref.org/works?" + urllib.parse.urlencode(
            {"query": q, "rows": min(int(limit) or 6, 10)})
        raw = WeaverOrchestrator._http_get(
            url, {"User-Agent": WeaverOrchestrator._ACAD_UA,
                  "Accept": "application/json"}, timeout)
        if not raw:
            return None
        try:
            items = ((_json.loads(raw).get("message") or {}).get("items")) or []
        except Exception:
            return None
        out = []
        for it in items[:limit]:
            title = (it.get("title") or [""])[0]
            if not title:
                continue
            doi = it.get("DOI", "")
            url_ = it.get("URL") or (("https://doi.org/" + doi) if doi else "")
            year = ""
            dp = ((it.get("issued") or {}).get("date-parts")
                  or (it.get("published") or {}).get("date-parts"))
            if dp and dp[0]:
                year = str(dp[0][0])
            auths = [(a.get("given", "") + " " + a.get("family", "")).strip()
                     for a in (it.get("author") or [])[:4]]
            abx = re.sub(r"<[^>]+>", "", it.get("abstract", "") or "")[:400]
            out.append({"title": title, "url": url_, "content": abx,
                        "authors": [a for a in auths if a], "year": year,
                        "doi": doi, "source": "crossref"})
        return out or None

    @staticmethod
    def _arxiv_search(query, lang, limit=6, timeout=12):
        import urllib.parse
        import html as _html
        import re
        q = (query or "").strip()
        if not q:
            return None
        url = "http://export.arxiv.org/api/query?" + urllib.parse.urlencode(
            {"search_query": "all:" + q, "max_results": min(int(limit) or 6, 10)})
        raw = WeaverOrchestrator._http_get(
            url, {"User-Agent": WeaverOrchestrator._ACAD_UA,
                  "Accept": "application/atom+xml"}, timeout)
        if not raw:
            return None

        def _c(s):
            return _html.unescape(re.sub(r"\s+", " ", s or "")).strip()
        out = []
        for m in re.finditer(r"<entry>(.*?)</entry>", raw, re.S):
            e = m.group(1)
            t = re.search(r"<title>(.*?)</title>", e, re.S)
            if not t:
                continue
            idm = re.search(r"<id>(.*?)</id>", e, re.S)
            summ = re.search(r"<summary>(.*?)</summary>", e, re.S)
            pub = re.search(r"<published>(\d{4})", e)
            auths = re.findall(r"<name>(.*?)</name>", e, re.S)
            out.append({"title": _c(t.group(1)),
                        "url": (idm.group(1).strip() if idm else ""),
                        "content": (_c(summ.group(1))[:400] if summ else ""),
                        "authors": [_c(a) for a in auths[:4]],
                        "year": (pub.group(1) if pub else ""),
                        "doi": "", "source": "arxiv"})
            if len(out) >= limit:
                break
        return out or None

    @staticmethod
    def _semanticscholar_search(query, lang, limit=6, timeout=12):
        import urllib.parse
        import json as _json
        q = (query or "").strip()
        if not q:
            return None
        url = ("https://api.semanticscholar.org/graph/v1/paper/search?"
               + urllib.parse.urlencode(
                   {"query": q, "limit": min(int(limit) or 6, 10),
                    "fields": "title,abstract,year,authors,url,openAccessPdf"}))
        raw = WeaverOrchestrator._http_get(
            url, {"User-Agent": WeaverOrchestrator._ACAD_UA,
                  "Accept": "application/json"}, timeout)
        if not raw:
            return None
        try:
            data = _json.loads(raw).get("data") or []
        except Exception:
            return None
        out = []
        for p in data[:limit]:
            title = p.get("title") or ""
            if not title:
                continue
            pdf = (p.get("openAccessPdf") or {}).get("url")
            out.append({"title": title, "url": pdf or p.get("url") or "",
                        "content": (p.get("abstract") or "")[:400],
                        "authors": [a.get("name", "")
                                    for a in (p.get("authors") or [])[:4]],
                        "year": str(p.get("year") or ""), "doi": "",
                        "source": "semanticscholar"})
        return out or None

    @staticmethod
    def _doaj_search(query, lang, limit=6, timeout=12):
        import urllib.parse
        import json as _json
        q = (query or "").strip()
        if not q:
            return None
        url = ("https://doaj.org/api/search/articles/" + urllib.parse.quote(q)
               + "?" + urllib.parse.urlencode({"pageSize": min(int(limit) or 6, 10)}))
        raw = WeaverOrchestrator._http_get(
            url, {"User-Agent": WeaverOrchestrator._ACAD_UA,
                  "Accept": "application/json"}, timeout)
        if not raw:
            return None
        try:
            results = _json.loads(raw).get("results") or []
        except Exception:
            return None
        out = []
        for r in results[:limit]:
            b = r.get("bibjson") or {}
            title = b.get("title") or ""
            if not title:
                continue
            url_ = ""
            for L in (b.get("link") or []):
                if L.get("url"):
                    url_ = L["url"]
                    break
            auths = [a.get("name", "") for a in (b.get("author") or [])[:4]]
            out.append({"title": title, "url": url_,
                        "content": (b.get("abstract") or "")[:400],
                        "authors": [a for a in auths if a],
                        "year": str(b.get("year") or ""), "doi": "",
                        "source": "doaj"})
        return out or None

    @staticmethod
    def _europepmc_search(query, lang, limit=6, timeout=12):
        import urllib.parse
        import json as _json
        q = (query or "").strip()
        if not q:
            return None
        url = ("https://www.ebi.ac.uk/europepmc/webservices/rest/search?"
               + urllib.parse.urlencode(
                   {"query": q, "format": "json",
                    "pageSize": min(int(limit) or 6, 10)}))
        raw = WeaverOrchestrator._http_get(
            url, {"User-Agent": WeaverOrchestrator._ACAD_UA,
                  "Accept": "application/json"}, timeout)
        if not raw:
            return None
        try:
            results = ((_json.loads(raw).get("resultList") or {})
                       .get("result")) or []
        except Exception:
            return None
        out = []
        for r in results[:limit]:
            title = r.get("title") or ""
            if not title:
                continue
            doi = r.get("doi", "")
            url_ = (("https://doi.org/" + doi) if doi
                    else (("https://europepmc.org/abstract/"
                           + str(r.get("source", "")) + "/" + str(r.get("id", "")))
                          if r.get("id") else ""))
            out.append({"title": title, "url": url_, "content": "",
                        "authors": [r.get("authorString", "")],
                        "year": str(r.get("pubYear") or ""), "doi": doi,
                        "source": "europepmc"})
        return out or None

    @classmethod
    def _scholarly_search(cls, query, lang, limit, timeout=14):
        """Query all free scholarly sources IN PARALLEL (OpenAlex, Crossref,
        arXiv, Semantic Scholar, DOAJ, Europe PMC) and merge/dedupe by DOI/URL/
        title. Total time ≈ the slowest source. Each degrades safely. Returns a
        merged list or None."""
        import concurrent.futures as _cf
        import itertools
        engines = [
            ("openalex", lambda: cls._openalex_search(query, lang, limit)),
            ("crossref", lambda: cls._crossref_search(query, lang, limit)),
            ("arxiv", lambda: cls._arxiv_search(query, lang, limit)),
            ("s2", lambda: cls._semanticscholar_search(query, lang, limit)),
            ("doaj", lambda: cls._doaj_search(query, lang, limit)),
            ("europepmc", lambda: cls._europepmc_search(query, lang, limit)),
        ]
        res = {}
        ex = _cf.ThreadPoolExecutor(max_workers=len(engines))
        try:
            fm = {ex.submit(fn): n for n, fn in engines}
            done, _pending = _cf.wait(fm, timeout=timeout)
            for f, n in fm.items():
                try:
                    res[n] = (f.result() or []) if f in done else []
                except Exception:
                    res[n] = []
        except Exception:
            for n, fn in engines:
                try:
                    res[n] = fn() or []
                except Exception:
                    res[n] = []
        finally:
            ex.shutdown(wait=False, cancel_futures=True)
        lists = [res.get(n, []) for n, _ in engines]
        # relevance terms from the query (drop off-topic noise like random
        # physics papers arXiv returns for common words). Arabic tokens ≥3,
        # Latin tokens ≥4; ignore a short stop-list.
        terms = cls._query_terms(query)
        cap = max(int(limit) or 8, 8)
        merged, backup, seen = [], [], set()
        for r in itertools.chain.from_iterable(itertools.zip_longest(*lists)):
            if not r:
                continue
            key = (r.get("doi") or r.get("url") or r.get("title") or "")
            key = key.strip().lower().rstrip("/")
            if not key or key in seen or not (r.get("title") or "").strip():
                continue
            seen.add(key)
            blob = (str(r.get("title", "")) + " "
                    + str(r.get("content", ""))).lower()
            if not terms or any(t in blob for t in terms):
                merged.append(r)
            else:
                backup.append(r)          # off-topic → only if nothing relevant
            if len(merged) >= cap:
                break
        final = merged if merged else backup
        return final[:cap] or None

    @staticmethod
    def _query_terms(query):
        """Meaningful lowercased query terms for relevance filtering: Arabic
        tokens ≥3 chars, Latin tokens ≥4, minus a tiny stop-list."""
        import re
        stop = {"عن", "في", "من", "على", "the", "and", "for", "with", "about",
                "أثر", "تأثير", "دراسة", "بحث", "تقرير", "اكتب", "حول"}
        out = set()
        for tok in re.split(r"[\s,،.:؛()\[\]\"']+", (query or "").lower()):
            tok = tok.strip()
            if not tok or tok in stop:
                continue
            is_ar = any("؀" <= c <= "ۿ" for c in tok)
            if (is_ar and len(tok) >= 3) or (not is_ar and len(tok) >= 4):
                out.add(tok)
        return out

    async def _academic_search(self, task: Task, mem: TaskMemory):
        """Layer 4 academic path: gather peer-reviewed / open-access sources
        from the free scholarly APIs and add them to the task's sources + RAG
        memory. Degrades safely (no network / all down → nothing added)."""
        card = task.task_card
        query = (card.get("topic") or task.description or "").strip()
        if not query:
            return
        lang = "ar" if card.get("language", "ar") == "ar" else "en"
        limit = int(card.get("reference_count") or 8)
        try:
            results = self._scholarly_search(query, lang, limit)
        except Exception:
            results = None
        if not results:
            mem.set_status(4, "بحث أكاديمي: لا نتائج (تدهور آمن)")
            return
        srcs = card.setdefault("sources", [])
        for r in results:
            title = r.get("title", "")
            url = r.get("url", "")
            auth = ", ".join(r.get("authors") or [])
            year = r.get("year", "")
            doi = r.get("doi", "")
            content = r.get("content") or ""
            srcs.append({"key": (title or url)[:60], "url": url, "title": title,
                         "content": content, "authors": r.get("authors") or [],
                         "year": year, "doi": doi,
                         "source": r.get("source", ""), "academic": True,
                         "full": False})
            mem.add_reference(
                f"[أكاديمي/{r.get('source','')}] {title} — {auth} ({year}) "
                f"{('doi:'+doi) if doi else ''} {content[:200]} ({url})",
                source_key=(url or doi or title))
        card["academic_reads"] = len(results)
        mem.set_status(4, f"بحث أكاديمي: {len(results)} مصدر محكّم")

    _URL_RE = None

    async def _read_pasted_urls(self, task: Task, mem: TaskMemory):
        """Read any URL the user pasted in the request (a web page or a YouTube
        video) as a PRIMARY source, via _extract_full (which routes YouTube to
        tool_youtube). Adds them to the task's sources + RAG memory. Safe: no
        URLs, or a failed read, just adds nothing."""
        import re
        # the dedicated YouTube path (Layer 3) already handled it — don't re-read
        if task.task_card.get("youtube"):
            return
        if self._URL_RE is None:
            type(self)._URL_RE = re.compile(r'https?://[^\s)>\]\"\'،]+')
        # only URLs in the CURRENT request — not ones pasted in earlier turns —
        # so an old link doesn't get re-read into an unrelated request now.
        urls = self._URL_RE.findall(self._current_request(task.description))
        card = task.task_card
        card["pasted_present"] = bool(urls)
        if not urls:
            return
        self._yt_lang = "ar" if card.get("language", "ar") == "ar" else "en"
        srcs = card.setdefault("sources", [])
        read = 0
        first_content = ""
        for u in urls[:3]:
            u = u.rstrip('.,)"،')
            try:
                txt = await self._extract_full(u)
            except Exception:
                txt = None
            if txt and txt.strip():
                full = txt.strip()
                if not first_content:
                    first_content = full
                srcs.append({"key": u[:60], "url": u, "title": u,
                             "content": full, "full": True, "pasted": True})
                # feed the FULL content into RAG in chunks (not a 300-char
                # snippet) so the writer can actually summarize the page/video.
                for i in range(0, min(len(full), 9000), 1500):
                    mem.add_reference(f"[رابط مُدرَج {u}] {full[i:i+1500]}",
                                      source_key=u)
                read += 1
        if read:
            card["pasted_reads"] = read
            mem.set_status(4, f"قراءة {read} رابط مُدرَج في الطلب")
            # "follow the file/link" rule: when the user did NOT lock a language
            # explicitly, the output follows the SOURCE content's own language.
            if card.get("lang_locked") is not True and first_content:
                card["language"] = self._detect_lang(first_content)

    async def _tool_web_search(self, query: str, lang: str, limit: int):
        """Fallback: the packaged web_search tool. Returns a results list
        (possibly empty) and never raises."""
        try:
            from capabilities.tools import tool_web_search
        except Exception:
            return []
        inputs = {"query": query, "language": lang, "limit": limit}
        inst = os.environ.get("WEAVER_SEARXNG_URL", "").strip()
        if inst:
            inputs["instance"] = inst
        try:
            res = await tool_web_search.run(inputs)
        except Exception:
            return []
        if not getattr(res, "ok", False):
            return []
        return [{"title": r.get("title", ""), "url": r.get("url", ""),
                 "content": r.get("content", "")}
                for r in (res.data or {}).get("results", [])]

    async def _extract_full(self, url: str):
        """Read a page's full text. Order: (0) UniWeb browser (curl_impersonate)
        → (1) tool_web_document (HTML via trafilatura, text PDFs via pdfplumber,
        SCANNED PDFs & images via OCR) → (2) plain trafilatura → None. Every
        branch degrades safely when a library/service is missing."""
        if not url:
            return None
        # 0a) YouTube links → dedicated tool_youtube (transcript/Whisper). Fully
        #     guarded: any failure just falls through to the normal path below,
        #     so non-YouTube links are completely unaffected.
        try:
            from capabilities.tools import tool_youtube
            if tool_youtube.is_youtube_url(url):
                lang = getattr(self, "_yt_lang", None) or "ar"
                res = await tool_youtube.run({"url": url, "lang": lang})
                if getattr(res, "ok", False):
                    txt = (res.data or {}).get("text")
                    if txt and len(txt.strip()) > 200:
                        return txt
        except Exception:
            pass
        # 0) UniWeb browser (curl_impersonate: real browser fingerprint, beats
        #    bot-blocking). firecrawl is removed; needs curl_cffi on the device.
        try:
            import os as _os, sys as _sys
            uw = _os.path.abspath(_os.path.join(
                _os.path.dirname(__file__), "..", "engines", "uniweb-core"))
            if uw not in _sys.path:
                _sys.path.insert(0, uw)
            import uniweb as _uniweb
            html = _uniweb.fetch(url)
            if html and isinstance(html, str) and len(html.strip()) > 200:
                # clean the fetched HTML to article text via trafilatura
                try:
                    from trafilatura import extract as _tex
                    txt = _tex(html, output_format="markdown",
                               include_comments=False, include_tables=True)
                    if txt and txt.strip():
                        return txt
                except Exception:
                    pass
                return html
        except Exception:
            pass
        # 1) web_document: HTML + text-PDF + scanned-PDF(OCR) + image(OCR)
        try:
            from capabilities.tools import tool_web_document
            res = await tool_web_document.run({"url": url, "ocr_lang": "ara+eng"})
            if getattr(res, "ok", False):
                d = res.data or {}
                if d.get("text"):
                    return d["text"]
                pages = d.get("pages") or []
                joined = "\n\n".join(p.get("text", "") for p in pages
                                     if p.get("text"))
                if joined.strip():
                    return joined
        except Exception:
            pass
        # 2) fallback: plain HTML extractor (trafilatura only)
        try:
            from capabilities.tools import tool_web_extract
            res = await tool_web_extract.run({"url": url, "format": "markdown"})
            if getattr(res, "ok", False):
                return (res.data or {}).get("text")
        except Exception:
            pass
        # 3) last resort: DNS-safe raw GET (survives broken phone DNS) then
        #    clean with trafilatura if available, else return the raw HTML.
        try:
            ua = ("Mozilla/5.0 (Linux; Android 13; Pixel 7) AppleWebKit/537.36 "
                  "(KHTML, like Gecko) Chrome/122.0.0.0 Mobile Safari/537.36")
            html = self._http_get(url, {"User-Agent": ua, "Accept": "text/html"})
            if html and len(html.strip()) > 200:
                try:
                    from trafilatura import extract as _tex
                    txt = _tex(html, output_format="markdown",
                               include_comments=False, include_tables=True)
                    if txt and txt.strip():
                        return txt
                except Exception:
                    pass
                return html
        except Exception:
            pass
        return None

    async def _web_search(self, task: Task, mem: TaskMemory):
        """Live web research — the search backbone of Layer 4.

        Layer wiring (the end-to-end flow, documented here):
          * Layer 3 (_route): detects a news/recency intent and turns on
            web_search (even for non-academic tasks; academic_search stays off).
          * Layer 4 (_layer_4 → _web_search): engine attempt order is
              SearXNG primary → SearXNG fallbacks → multi-engine (DuckDuckGo +
              Bing + Mojeek + Startpage, queried in parallel and merged) →
              built-in public SearXNG → packaged tool.
            The top 3 links are then read in full (HTML, text/scanned PDFs via
            OCR, images); the rest are kept as snippets. For recency queries it
            injects the current date into the query, applies a time filter
            (SearXNG time_range / DuckDuckGo df), and re-orders results
            newest-first before reading them.
          * Every engine degrades safely — a down service or missing library
            just yields fewer/no sources, never an error.
        """
        card = task.task_card
        query = (card.get("topic") or task.description or "").strip()
        if not query:
            return
        lang = "ar" if card.get("language", "ar") == "ar" else "en"
        limit = int(card.get("reference_count") or 8)

        # news/recency intent → date-augmented query + time filter + recency sort
        is_recency = self._is_recency_query(
            (card.get("topic") or "") + " " + (task.description or ""))
        sx_time = ddg_df = None
        if is_recency:
            query = self._augment_query_with_date(query, lang)
            sx_time = "week"          # SearXNG time_range
            ddg_df = "w"              # DuckDuckGo df = past week

        # 1) SearXNG primary (env override, else the packaged default port 8888)
        instance = os.environ.get("WEAVER_SEARXNG_URL", "").strip() or "http://127.0.0.1:8888"
        results = self._searx_query(instance, query, lang, limit,
                                    time_range=sx_time, sort_by_date=is_recency)
        used = "searxng:" + instance
        # 2) SearXNG fallbacks (comma-separated list, tried in order)
        if not results:
            for fb in [u.strip() for u in
                       os.environ.get("WEAVER_SEARXNG_FALLBACKS", "").split(",")
                       if u.strip()]:
                r = self._searx_query(fb, query, lang, limit,
                                      time_range=sx_time, sort_by_date=is_recency)
                if r:
                    results = r
                    used = "searxng:" + fb
                    break
        # 3) serverless MULTI-ENGINE (DuckDuckGo + Bing + Mojeek + Startpage,
        #    queried in parallel and merged) — SearXNG-like breadth, NO server.
        if not results:
            multi = self._multi_engine_search(query, lang, limit, df=ddg_df)
            if multi:
                results = multi
                used = "multi-engine"
        # 4) built-in public SearXNG fallbacks — AUTOMATIC, zero setup. Reached
        #    only if everything above failed, so DuckDuckGo's fast path is never
        #    slowed. Each dead instance is memoized (skipped next time); short
        #    timeout so a hanging server can't stall the pipeline.
        if not results:
            for fb in _DEFAULT_SEARXNG_FALLBACKS:
                if fb in _SEARX_DEAD:
                    continue
                r = self._searx_query(fb, query, lang, limit, timeout=5,
                                      time_range=sx_time, sort_by_date=is_recency)
                if r:
                    results = r
                    used = "searxng-default:" + fb
                    break
                _SEARX_DEAD.add(fb)
        # 5) fall back to the packaged tool as a last resort
        if not results:
            results = await self._tool_web_search(query, lang, limit)
            used = "web_search"
        if not results:
            mem.set_status(4, "بحث ويب: لا نتائج (تدهور آمن)")
            return

        # recency queries: put the newest results first before reading top 3
        if is_recency:
            results = self._sort_results_by_recency(results)

        srcs = card.setdefault("sources", [])
        full_reads = 0
        for i, r in enumerate(results):
            url = r.get("url", "")
            title = r.get("title", "")
            snippet = r.get("content", "")
            content = snippet
            is_full = False
            if i < 3:  # read the top 3 links in full
                text = await self._extract_full(url)
                if text:
                    content = text
                    is_full = True
                    full_reads += 1
            srcs.append({"key": (title or url)[:60], "url": url, "title": title,
                         "content": content, "full": is_full})
            mem.add_reference(f"[ويب] {title} — {content[:300]} ({url})",
                              source_key=url)
        card["web_full_reads"] = full_reads
        mem.set_status(4, f"بحث ويب ({used}): {len(results)} نتيجة، "
                          f"قراءة كاملة لـ {full_reads} صفحة")

    async def _layer_5(self, task: Task, mem: TaskMemory):
        """٥: المصداقية — تمرير كل مصدر عبر check_source وإسقاط المرفوض."""
        task.status = TaskStatus.LAYER_5
        mem.set_status(5, "تقييم مصداقية المصادر")

        sources = task.task_card.get("sources") or []
        if not sources:
            return  # لا مصادر مُنظّمة لتصفيتها — أبقِ السلوك الافتراضي
        try:
            import os as _os, sys as _sys
            sp = _os.path.abspath(_os.path.join(
                _os.path.dirname(__file__), "..", "capabilities", "skills",
                "credibility_scorer", "scripts"))
            if sp not in _sys.path:
                _sys.path.insert(0, sp)
            import source_reliability as _sr
        except Exception as e:
            mem.set_status(5, f"مصداقية (تخطّي: {e})")
            return
        lang = task.task_card.get("language", "ar")
        kept, dropped = [], []
        for s in sources:
            url = s.get("url", "") if isinstance(s, dict) else str(s)
            r = _sr.check_source(url, task.task_card, lang)
            if r.get("allowed"):
                kept.append(s)
            else:
                dropped.append({"source": s, "reason": r.get("reason"),
                                "alternative": r.get("alternative")})
        task.task_card["sources"] = kept
        task.task_card["credibility"] = {"kept": len(kept), "dropped": dropped}
        mem.set_status(5, f"مصداقية: قُبل {len(kept)}، رُفض {len(dropped)}")

    def _descriptive_titles(self, topic, sections_plan, lang):
        """Replace the abstract structural labels ("المبحث 1"/"المطلب 1.1"/
        "Section 1"/"Subsection 1.2") with DESCRIPTIVE, topic-specific titles the
        model proposes — preserving each slot's position/key/level, and giving
        every section a DISTINCT sub-topic (which is what stops sibling sections
        from repeating the same generic definition).

        Uses a SIMPLE numbered line-per-item request (one title per line), which
        a weak model handles far more reliably than nested JSON — a strong model
        produces it just as well, so quality is never capped. Intro/conclusion/
        references are left untouched. Guarded: no model, an empty/garbled reply
        → the original plan is returned unchanged."""
        if not self.llm_fn or not sections_plan:
            return sections_plan
        import re
        abstract_re = re.compile(
            r'^\s*(?:المبحث|المطلب|Section|Subsection)\b', re.I)
        # collect the abstract body slots IN ORDER, each with a role label
        slots, main_no = [], 0        # slots: list of (plan_index, role_text)
        for i, s in enumerate(sections_plan):
            title = (s.get("title") or s.get("heading") or "").strip()
            if not abstract_re.match(title):
                continue
            lvl = int(s.get("level", 1) or 1)
            if lvl <= 1:
                main_no += 1
                role = (f"مبحث رئيسي رقم {main_no}" if lang == "ar"
                        else f"main section #{main_no}")
            else:
                role = (f"مطلب فرعي تحت المبحث {main_no}" if lang == "ar"
                        else f"subsection under section {main_no}")
            slots.append((i, role))
        if not slots:
            return sections_plan          # nothing abstract to rename
        roles_block = "\n".join(f"{n + 1}. {r}"
                                for n, (idx, r) in enumerate(slots))
        if lang == "ar":
            prompt = (
                f"أريد عناوين وصفية دقيقة لبحث علمي عن: «{topic}».\n"
                f"لكل بندٍ في القائمة التالية اكتب عنواناً وصفياً واحداً يخصّ "
                f"الموضوع فعلاً، ومختلفاً عن البقية (لا تعريفات عامة مكرّرة)، بلا "
                f"كلمتَي «مبحث»/«مطلب»:\n{roles_block}\n\n"
                f"أعِد {len(slots)} سطراً فقط، سطراً واحداً لكل عنوان وبنفس "
                f"الترتيب، كلٌّ يبدأ برقمه هكذا: «1. العنوان».")
        else:
            prompt = (
                f"I need precise descriptive titles for research on: \"{topic}\".\n"
                f"For each item below, write ONE descriptive, topic-specific "
                f"title, distinct from the others (no repeated general "
                f"definitions), without the words 'Section'/'Subsection':\n"
                f"{roles_block}\n\nReturn exactly {len(slots)} lines, one title "
                f"per line in the same order, each starting with its number: "
                f"\"1. Title\".")
        try:
            raw = self.llm_fn(prompt, system=self.system_main,
                              temperature=0.3) or ""
        except Exception:
            return sections_plan
        # parse: prefer "N. title" numbering; else take non-empty lines in order
        titles = {}
        for line in raw.splitlines():
            line = line.strip()
            if not line:
                continue
            m = re.match(r'^[\(\[]?(\d{1,3})[\)\].\-:،]\s*(.+)$', line)
            if m:
                n = int(m.group(1))
                if 1 <= n <= len(slots):
                    titles[n - 1] = m.group(2).strip()
        if not titles:
            lines = [re.sub(r'^[\-\*•\d\.\)\(:،\s]+', '', l).strip()
                     for l in raw.splitlines() if l.strip()]
            lines = [l for l in lines if l]
            for k, l in enumerate(lines[:len(slots)]):
                titles[k] = l
        if not titles:
            return sections_plan

        def _clean(t):
            t = (t or "").strip().strip('"“”«»').strip()
            t = re.sub(r'^(?:المبحث|المطلب|Section|Subsection)\b[\s:،.\d]*', '',
                       t, flags=re.I).strip()
            return t
        plan = [dict(s) for s in sections_plan]     # copy, don't mutate input
        for k, (idx, role) in enumerate(slots):
            if k in titles:
                ct = _clean(titles[k])
                if ct:
                    plan[idx]["title"] = ct
        return plan

    @staticmethod
    def _clean_section_body(body, title):
        """Tidy a written section body: drop a leading duplicate of its own
        heading (the "المطلب 1.3: …" leak), and strip leaked markdown heading
        markers (## …) that the docx builder would otherwise show literally.
        Prose content is preserved. Safe on empty input."""
        import re
        if not body:
            return body
        b = body.lstrip()
        t = (title or "").strip()
        # strip a leading duplicate of the heading ONLY when it reads as a
        # heading (title then ":"/"："/line-break/end) — never when the title
        # naturally opens the first sentence (e.g. "التركيب … هو الوحدة …").
        if t and b.startswith(t):
            _rest = b[len(t):]
            _after = _rest.lstrip()
            if _rest[:1] == "\n" or _after[:1] in (":", "：") or _after == "":
                b = _rest.lstrip(" :：،.-\n")
        # a leading abstract label glued to the start ("المطلب 1.3: ")
        b = re.sub(r'^\s*(?:المبحث|المطلب|Section|Subsection)\s*[\d.]*\s*'
                   r'[:：]?\s*', '', b)
        # markdown heading markers at line starts → keep text, drop the hashes
        b = re.sub(r'(?m)^[ \t]*#{1,6}[ \t]*', '', b)
        return b.strip()

    async def _layer_6(self, task: Task, mem: TaskMemory):
        """٦: الصياغة — بناء البنية ثم المنهجية ثم كتابة كل قسم.
        كل خطوة تستخدم مهارة/قالباً موجوداً؛ عند غياب النموذج تبقى مسودة فارغة."""
        task.status = TaskStatus.LAYER_6
        mem.set_status(6, "صياغة البحث")
        card = task.task_card
        lang = card.get("language", "ar")
        # adapt every writer to the running model's ceiling (small/medium/large)
        # so it works reliably at its own peak. Set once; read by this layer and
        # by weak_model_support below. Additive — default stays "medium".
        try:
            card.setdefault("model_strength", self._model_strength())
        except Exception:
            pass
        # detect Islamic content once so the writer is guided to use the correct
        # Quran/Hadith marks (enforced again after writing by _apply_islamic_marks)
        try:
            card.setdefault("islamic", self._is_islamic_content(
                f"{card.get('topic','')} {task.description}"))
        except Exception:
            pass

        # ── YouTube path: produce the summary / transcript directly, with NO
        #    research structure or methodology, then return early.
        yt = card.get("youtube")
        if yt:
            transcript = (yt.get("transcript") or "").strip()
            mode = yt.get("mode", "summary")
            # Fetch failed upstream (no captions / throttled / lib missing):
            # emit a clear, honest, actionable message — never an empty reply
            # and never a fake research report.
            if not transcript and yt.get("error"):
                why = yt.get("error")
                body_txt = (
                    f"تعذّر جلب نص هذا الفيديو من يوتيوب ({why}).\n\n"
                    "**الأسباب المحتملة:**\n\n"
                    "- الفيديو لا يحتوي ترجمة/نصاً تلقائياً (captions مُعطّلة).\n"
                    "- يوتيوب حجب الطلب مؤقتاً بسبب طلبات متتالية — انتظر قليلاً "
                    "ثم أعد المحاولة، ويُفضّل عدم فتح عدة محادثات في آنٍ واحد.\n"
                    "- مكتبة youtube-transcript-api غير مثبّتة على الجهاز "
                    "(pip install youtube-transcript-api)."
                )
                task.sections = [{"heading": "تعذّر التفريغ", "body": body_txt}]
                # web/terminal reply as real Markdown (heading + body)
                task.draft = "## تعذّر التفريغ\n\n" + body_txt
                mem.set_status(6, "يوتيوب: تعذّر جلب النص")
                return
            # Clean (no-timestamp) text drives the summary, so the summary is real
            # prose — never a copy of the timestamped transcript.
            plain = (yt.get("transcript_plain") or "").strip() or transcript
            # Resolve the SUMMARY output language per the rule: explicit name >
            # source/original language of the video > the language the request
            # itself is written in.
            _lk = yt.get("out_lang_kind")
            _ln = yt.get("out_lang_name")
            _en = (yt.get("req_lang") or "ar") == "en"
            if _lk == "source":
                _lang_instr = "بنفس لغة النص أدناه (لغته الأصلية)"
            elif _lk == "name":
                _lang_instr = "باللغة " + (_ln or "العربية")
            else:
                _lang_instr = "بالإنجليزية" if _en else "بالعربية"
            _h_sum = "Summary" if (_lk == "name" and _ln == "الإنجليزية") \
                or (_lk is None and _en) else "الخلاصة"
            _h_tr = "Video transcript" if _h_sum == "Summary" \
                else "التفريغ النصي للفيديو"
            sections = []
            # 1) TRANSCRIPT FIRST (when requested) — always verbatim, its own lang
            if mode in ("transcript", "both"):
                sections.append({"heading": _h_tr,
                                 "body": transcript, "kind": "transcript"})
            # 2) SUMMARY AFTER the transcript (only when requested)
            if mode in ("summary", "both"):
                summ = ""
                if self.llm_fn and plain:
                    try:
                        summ = self.llm_fn(
                            "لخّص النص التالي " + _lang_instr + " في نقاط واضحة "
                            "ومرتّبة، دون مقدمة بحثية أو مباحث أو مراجع. اجعل كل "
                            "نقطة سطراً يبدأ بـ \"- \":\n\n" + plain[:12000],
                            system=self.system_main, temperature=0.4) or ""
                    except Exception as e:
                        mem.set_status(6, f"يوتيوب (تخطّي التلخيص: {e})")
                summ = summ.strip()
                if summ:
                    sections.append({"heading": _h_sum, "body": summ,
                                     "kind": "summary"})
                elif mode == "summary":
                    # summary-only and generation failed → give the clean text
                    # (not the timestamped dump) with an honest note.
                    note = (plain[:4000] + ("…" if len(plain) > 4000 else ""))
                    sections.append({
                        "heading": _h_tr,
                        "body": "تعذّر توليد ملخّص تلقائي الآن؛ في ما يلي نصّ "
                                "الفيديو:\n\n" + note, "kind": "summary"})
                else:
                    # "both": transcript is already shown → just note the miss.
                    sections.append({"heading": _h_sum,
                                     "body": "تعذّر توليد الملخّص تلقائياً الآن.",
                                     "kind": "summary"})
            task.sections = [{"heading": s["heading"], "body": s["body"]}
                             for s in sections]
            # Web/terminal reply (task.draft) as real Markdown so the chat UI
            # renders it like the exported file: "##" headings, each verbatim
            # "[MM:SS]" transcript line on its own line, and a faint "---" divider
            # before the summary that follows the transcript. task.sections keeps
            # its own (already-correct) formatting for the exported file.
            parts = []
            for i, s in enumerate(sections):
                body = s["body"]
                if s.get("kind") == "transcript":
                    body = "\n\n".join(ln.strip() for ln in body.split("\n")
                                       if ln.strip())
                block = f"## {s['heading']}\n\n{body}"
                if i > 0 and sections[i - 1].get("kind") == "transcript":
                    block = "---\n\n" + block
                parts.append(block)
            task.draft = "\n\n".join(parts).strip()
            mem.set_status(6, f"يوتيوب: أُنتج ({mode}، {len(sections)} قسم)")
            return

        # ── scope limits (references-only / part-only handled here; outline-only
        #    after the structure is built below) ──
        scope = card.get("scope")
        if scope == "references":
            head = "المراجع والدراسات" if lang == "ar" else "References"
            body = self._format_references_only(card, lang)
            task.sections = [{"heading": head, "body": body}]
            task.draft = f"## {head}\n\n{body}"
            mem.set_status(6, "إخراج: مراجع/دراسات فقط")
            return

        # techniques for the model strength — shown in the tool-call/thinking UI,
        # never written into the output document. We already write per-section.
        try:
            card["reliability"] = self._skill_call(
                "weak_model_support", "weak_model_support", "reliability_plan",
                card.get("model_strength", "medium"))
        except Exception:
            pass

        # 1) البنية — تخطّى إذا كانت المهمة تُحدّد بنيتها بنفسها (build_structure→None)
        sections_plan = card.get("sections")
        if not sections_plan:
            try:
                plan = self._skill_call("research_structure", "structures",
                                        "build_structure", card, lang)
                if plan and plan.get("sections"):
                    sections_plan = plan["sections"]
                    card["sections"] = sections_plan
                    card.setdefault("tier", plan.get("tier"))
            except Exception as e:
                mem.set_status(6, f"بنية (تخطّي: {e})")
        if not sections_plan:
            sections_plan = [{"title": card.get("topic", "") or task.description,
                              "level": 1}]

        # give the abstract "المبحث/المطلب" slots DESCRIPTIVE, topic-specific
        # titles so each section chunk has a real sub-topic to write about (the
        # writer can't produce content for a meaningless "المطلب 1.1"). Additive
        # and guarded: on any miss the original structural labels are kept.
        try:
            sections_plan = self._descriptive_titles(
                card.get("topic", "") or task.description, sections_plan, lang)
            card["sections"] = sections_plan
        except Exception as e:
            mem.set_status(6, f"عناوين وصفية (تخطّي: {e})")

        # outline-only → output just the structure, don't write any bodies
        if scope == "outline":
            head = "هيكل العمل" if lang == "ar" else "Outline"
            lines = []
            for sec in sections_plan:
                lvl = int(sec.get("level", 1) or 1)
                title = sec.get("title") or sec.get("heading") or ""
                if title:
                    lines.append(("  " * max(0, lvl - 1)) + "- " + title)
            body = "\n".join(lines)
            task.sections = [{"heading": head, "body": body}]
            task.draft = f"## {head}\n\n{body}"
            mem.set_status(6, "إخراج: هيكل فقط")
            return

        # part-only → write just the one part the user asked for (single section)
        if scope == "part":
            sections_plan = [{"title": self._current_request(task.description),
                              "level": 1}]

        # 2) المنهجية — إن لزمت وغابت
        try:
            has_m = self._skill_call("research_methodology", "methodology",
                                     "has_methodology", card)
            if (not has_m and card.get("task_type", "") in
                    ("بحث", "research", "دراسة", "thesis", "report", "تقرير",
                     "analysis", "تحليل")):
                m = self._skill_call("research_methodology", "methodology",
                                     "build_methodology", card, lang)
                if m:
                    card["methodology"] = m
        except Exception as e:
            mem.set_status(6, f"منهجية (تخطّي: {e})")

        # 3) كتابة كل قسم بحقن سياقات RAG الخاصة به
        rag = mem.get_references(card.get("topic", "") or task.description,
                                 limit=20) or []
        rag_ctx = "\n".join(str(x) for x in rag)
        # fallback: the semantic memory search can miss even when sources WERE
        # gathered (they live in card["sources"]). Build the context straight
        # from them so the section writers actually receive the evidence — this
        # is what lets the cited path (and the specialized writers) run instead
        # of silently degrading to the no-sources path.
        if (not rag_ctx.strip()) and card.get("sources"):
            _lines = []
            for s in (card.get("sources") or [])[:20]:
                if isinstance(s, dict):
                    _t = s.get("title") or ""
                    _c = (s.get("content") or "")[:200]
                    _k = s.get("key") or s.get("doi") or ""
                    _u = s.get("url") or ""
                    _line = f"[{_k}] {_t} — {_c} ({_u})".strip()
                    if _line.strip("[] —()"):
                        _lines.append(_line)
                elif str(s).strip():
                    _lines.append(str(s))
            if _lines:
                rag_ctx = "\n".join(_lines)
        no_ctx = (not rag_ctx) or rag_ctx.strip() in ("", "(none)")
        mode = card.get("sourcing_mode", "cited")
        # In "cited" mode with NO retrieved context, don't refuse — write from
        # the model's knowledge and flag it so a clear note is added later.
        if mode == "cited" and no_ctx:
            card["sources_unavailable"] = True
        prof = self._strength_profile(card.get("model_strength", "medium"))
        parts, out_sections = [], []
        for sec in sections_plan:
            title = sec.get("title") or sec.get("heading") or ""
            body = ""
            # ── bound specialized section writers (skills already present, wired
            #    here) — additive: on any miss the generic writer below runs
            #    unchanged, keeping full backward compatibility ──
            if self.llm_fn:
                try:
                    _spec = self._write_section_specialized(
                        title, card, lang, mode, no_ctx, out_sections, prof)
                except Exception as e:
                    _spec = None
                    mem.set_status(6, f"مهارة قسم (تخطّي: {e})")
                if _spec:
                    body = _spec
            if self.llm_fn and not body:
                from pipeline import prompts as _p
                _topic = card.get("topic", "") or task.description
                # prior-sections context so each section knows what was already
                # written and does NOT repeat the general definition (the direct
                # cause of sibling sections all restating the same opening). Helps
                # a strong model cohere and a weak one avoid loops alike.
                _prior = ""
                _pi = [f"- {o.get('heading', '')}: "
                       f"{' '.join((o.get('body', '') or '').split())[:110]}"
                       for o in out_sections[-6:] if (o.get('body') or '').strip()]
                if _pi:
                    _prior = ((
                        "أقسامٌ كُتبت قبل هذا القسم — لا تُعِد تعريف الموضوع العام "
                        "ولا محتواها، وركّز حصراً على الزاوية الخاصة بهذا القسم:\n"
                        if lang == "ar" else
                        "Sections already written — do NOT repeat the general "
                        "definition or their content; focus only on THIS "
                        "section's angle:\n") + "\n".join(_pi) + "\n")
                # MODEL-AGNOSTIC safety net: if the title is still an abstract
                # structural label (a weak model may not have produced a
                # descriptive one), frame it with the topic so the writer knows
                # what this section is about — otherwise it writes nothing.
                if title.strip().startswith(
                        ("المبحث", "المطلب", "Section", "Subsection")):
                    section_name = (
                        f"«{title}» ضمن بحث عن: {_topic} — اكتب المحتوى العلمي "
                        f"المناسب لموضع هذا القسم (خلفية/تفصيل/تحليل بحسب موقعه)، "
                        f"متماسكاً ومرتبطاً بالموضوع مباشرة"
                        if lang == "ar" else
                        f"\"{title}\" within research on: {_topic} — write the "
                        f"scientific content appropriate to this section's role")
                else:
                    section_name = title
                if mode == "uncited":
                    prompt = _p.PROMPT_LAYER_6_WRITE_UNCITED.format(
                        section_name=section_name, topic=card.get("topic", ""),
                        length=card.get("page_count", ""),
                        rag_contexts=rag_ctx or "(none)", prior_content=_prior)
                    system = _p.SYSTEM_PROMPT_WRITE_NO_SOURCES
                elif mode == "none" or no_ctx:
                    # explicit no-sources request, OR sources were required but
                    # none could be retrieved — write from knowledge, no refusal
                    prompt = _p.PROMPT_LAYER_6_WRITE_NO_SOURCES.format(
                        section_name=section_name, topic=card.get("topic", ""),
                        length=card.get("page_count", ""), prior_content=_prior)
                    system = _p.SYSTEM_PROMPT_WRITE_NO_SOURCES
                else:
                    prompt = _p.PROMPT_LAYER_6_WRITE.format(
                        section_name=section_name, topic=card.get("topic", ""),
                        citation_style=card.get("citation_style", ""),
                        length=card.get("page_count", ""),
                        rag_contexts=rag_ctx or "(none)", prior_content=_prior)
                    # dedicated WRITING system prompt: forbids clarifying
                    # questions/greetings that a chatty model would emit
                    system = _p.SYSTEM_PROMPT_WRITE
                # adapt depth/length + temperature to the model's ceiling
                _depth = prof.get("depth") if lang == "ar" else prof.get("depth_en")
                if _depth:
                    prompt = prompt + "\n\n" + _depth
                # guide Quran/Hadith marks for Islamic content
                if card.get("islamic"):
                    prompt = prompt + "\n\n" + (self._ISLAMIC_DIRECTIVE_AR
                                               if lang == "ar"
                                               else self._ISLAMIC_DIRECTIVE_EN)
                try:
                    body = self.llm_fn(prompt, system=system,
                                       temperature=prof.get("temp", 0.5))
                except Exception as e:
                    mem.set_status(6, f"كتابة قسم (تخطّي: {e})")
                # guard: a conversational model may answer with a greeting /
                # clarifying question / options menu instead of content. Detect
                # it and retry ONCE with a blunt content-only instruction.
                if self._looks_conversational(body):
                    firm = (prompt + "\n\n"
                            + ("اكتب نص هذا القسم كاملاً ومباشرةً الآن. ممنوع منعاً "
                               "باتاً: التحية، طرح أي سؤال، طلب توضيح، أو عرض "
                               "خيارات. ابدأ بالمحتوى فوراً."
                               if lang == "ar" else
                               "Write the full text of this section directly "
                               "now. Absolutely no greeting, no question, no "
                               "request for clarification, no options. Begin "
                               "with the content immediately."))
                    try:
                        retry = self.llm_fn(firm, system=_p.SYSTEM_PROMPT_WRITE,
                                            temperature=0.4)
                        if retry and not self._looks_conversational(retry):
                            body = retry
                        elif self._looks_conversational(body):
                            body = ""   # drop the chat turn rather than ship it
                    except Exception:
                        pass
                # MODEL-AGNOSTIC safety net: a weak model may return an EMPTY
                # (non-conversational) body. One last blunt, direct attempt so no
                # section ships as a bare heading.
                if not (body or "").strip():
                    try:
                        _db = self.llm_fn(
                            (f"اكتب محتوى قسم «{title}» من بحث علمي عن: {_topic}. "
                             "اكتب فقرات علمية مباشرة (نحو 150–300 كلمة) دون "
                             "عنوان ودون أي سؤال أو تحية."
                             if lang == "ar" else
                             f"Write the content of section \"{title}\" of "
                             f"research on: {_topic}. Direct scientific "
                             "paragraphs (~150–300 words), no heading, no "
                             "questions, no greeting."),
                            system=_p.SYSTEM_PROMPT_WRITE,
                            temperature=0.4) or ""
                        if _db.strip() and not self._looks_conversational(_db):
                            body = _db.strip()
                    except Exception:
                        pass
            # tidy leaked heading duplicates / markdown markers before shipping
            body = self._clean_section_body(body, title)
            parts.append((f"{title}\n{body}").strip())
            out_sections.append({"heading": title, "body": body})
        task.draft = "\n\n".join(p for p in parts if p)
        task.sections = out_sections
        mem.set_status(6, f"صياغة: {len(out_sections)} قسم ({mode})")
        # run matched enrichment skills (task.skills) that have a write-stage
        # handler — turns skill routing into real execution. Additive/guarded.
        self._dispatch_skills(task, card, lang, mem)
        # when a data file (csv/xlsx) is attached, run REAL statistics on it and
        # inject the computed results (never invented). Additive/guarded.
        self._inject_statistics(task, card, lang, mem)
        # enforce correct Quran/Hadith marks for Islamic content (text-level,
        # all formats). Additive/guarded; no-op for non-Islamic text.
        self._apply_islamic_marks(task, card, lang, mem)

    async def _layer_6_6(self, task, mem):
        """٦.٦: تحقق الطول والتغطية — بين الكتابة (6) والأنسنة (6.5).

        يمنح النظام ما يفعله المحرّر يدوياً: يتأكّد أن كل قسم مطلوب كُتب فعلاً
        (فيكتب الناقص)، وأن طول النص قريب من الهدف المطلوب (فيوسّع القصير).
        حلقة تصحيح واحدة، آمنة (تعود مبكّراً بلا تعديل عند غياب طول/أقسام/نموذج).

        صدق القياس: عدّ الكلمات حتمي ودقيق (len(split))، لا تقدير. أمّا الصفحات
        فتقديرية (~500 كلمة/صفحة) ما لم يُفتح الملف بعد التصدير — والتحقق الفعلي
        من الصفحات بعد التصدير مهمّة منفصلة لاحقة، لا تُنفَّذ هنا."""
        try:
            card = task.task_card or {}
            # تخطَّ أوضاع يوتيوب/التفريغ ومهام بلا كتابة
            yt = card.get("youtube")
            if yt and yt.get("mode") in ("transcript", "both"):
                return
            if not task.sections:
                return

            # ── (أ) تحقق تغطية الأقسام ──
            plan = card.get("sections") or []
            required = [(s.get("title") or s.get("heading") or "") for s in plan]
            required = [r for r in required if r]
            if required:
                missing = self.verify_sections_coverage(required, task.sections)
                if missing and self.llm_fn:
                    for title in missing:
                        try:
                            body = self.llm_fn(
                                f"اكتب قسم «{title}» لهذا الموضوع: "
                                f"{card.get('topic', task.description)}. "
                                f"اكتب المحتوى مباشرة دون عنوان.",
                                system=getattr(self, "system_write", None)
                                or getattr(self, "system_main", None),
                                temperature=0.5) or ""
                        except Exception:
                            body = ""
                        if body.strip():
                            task.sections.append({"heading": title,
                                                  "body": body.strip()})
                    mem.set_status(66, f"تغطية: أُضيف {len(missing)} قسم ناقص")

            # ── (ب) تحقق الطول (كلمات) ──
            target = self.extract_length_target(task.description)
            tw = target.get("words")
            if tw:
                actual = self.count_words(task.draft or "")
                lo, hi = int(tw * 0.9), int(tw * 1.15)
                if actual < lo and self.llm_fn:
                    # النص أقصر من المطلوب → وسّع أضعف الأقسام (الأقصر)
                    deficit = tw - actual
                    secs = sorted(task.sections,
                                  key=lambda s: self.count_words(s.get("body", "")))
                    for s in secs[: max(1, len(secs) // 2)]:
                        if deficit <= 0:
                            break
                        try:
                            more = self.llm_fn(
                                f"وسّع الفقرة التالية بعمق أكبر وتفصيل دقيق "
                                f"(أضِف نحو {min(deficit, 300)} كلمة) دون تكرار "
                                f"ودون حشو:\n\n{s.get('body', '')}",
                                system=getattr(self, "system_write", None)
                                or getattr(self, "system_main", None),
                                temperature=0.5) or ""
                        except Exception:
                            more = ""
                        if more.strip():
                            added = (self.count_words(more)
                                     - self.count_words(s.get("body", "")))
                            s["body"] = more.strip()
                            deficit -= max(0, added)
                    mem.set_status(66, f"طول: وُسّع النص نحو الهدف {tw}")

                # أعد بناء draft بعد أي تعديل
                task.draft = "\n\n".join(
                    (f"{s.get('heading', '')}\n{s.get('body', '')}").strip()
                    for s in task.sections if (s.get("heading") or s.get("body")))

                final = self.count_words(task.draft)
                card["word_count_actual"] = final
                card["word_count_target"] = tw
                mem.set_status(66, f"طول نهائي: {final} كلمة (هدف {tw})")
        except Exception as e:
            mem.set_status(66, f"تحقق الطول/التغطية (تخطّي: {e})")

    async def _layer_6_5(self, task: Task, mem: TaskMemory):
        """٦.٥: إعادة الصياغة والتنظيف — أنسنة النص وإزالة البصمة الآلية.

        تُطبّق صامتةً: تحمي الاستشهادات، تستبدل كلمات AI بمرادفات بشرية،
        وتزيل البصمات البصرية (الشرطات الطويلة، الرموز الزخرفية، الخلط اللغوي)
        حسب نوع الملف — مع إبقاء الرموز في عروض PowerPoint.
        """
        task.status = TaskStatus.LAYER_6_5
        mem.set_status(65, "إعادة الصياغة والتنظيف")

        # verbatim YouTube transcript must stay word-for-word → skip humanizing
        # (transcript-only, and "both" whose transcript half must not change).
        yt = task.task_card.get("youtube")
        if yt and yt.get("mode") in ("transcript", "both"):
            mem.set_status(65, "تفريغ حرفي — تخطّي الأنسنة")
            return

        lang = task.task_card.get("language", "ar")
        fmt = self._primary_format(task.task_card)
        file_type = {"pptx": "pptx", "xlsx": "xlsx", "pdf": "pdf"}.get(fmt, "docx")

        import os, sys
        # اختر السكربت حسب اللغة
        skill = "arabic_rewriter" if lang == "ar" else "english_rewriter"
        fname = "rewrite_ar" if lang == "ar" else "rewrite_en"
        scripts = os.path.join(os.path.dirname(__file__), "..", "capabilities",
                               "skills", skill, "scripts")
        scripts = os.path.abspath(scripts)
        if scripts not in sys.path:
            sys.path.insert(0, scripts)
        try:
            mod = __import__(fname)
            draft = task.draft or task.task_card.get("draft", "")
            if draft:
                # Protect citations from the AI-fingerprint cleaner (which would
                # otherwise strip "(Smith, 2023)" as a Latin-in-Arabic mix). We
                # mask them, humanize, then restore them intact.
                masked, cites = self._mask_citations(draft)
                result = mod.humanize_text(masked, file_type=file_type)
                task.draft = self._unmask_citations(result["text"], cites)
                task.task_card["humanized"] = True
                task.task_card["cleaning_issues"] = result.get("issues", [])
        except Exception as e:
            mem.set_status(65, f"إعادة الصياغة (تخطّي: {e})")

    # citation guards — keep (Author, Year) / (key, p. N) / (…، ص. N) intact
    _CITE_RE = None

    @classmethod
    def _mask_citations(cls, text: str):
        """Replace parenthesised citations with digit-only placeholders so the
        humanizer's Latin/decoration cleaning can't damage them."""
        import re
        if cls._CITE_RE is None:
            cls._CITE_RE = re.compile(
                r"\([^()]*(?:\b\d{4}\b|p\.?\s*\d+|ص\.?\s*\d+)[^()]*\)")
        cites = []

        def _sub(m):
            cites.append(m.group(0))
            return "" + str(len(cites) - 1) + ""
        return cls._CITE_RE.sub(_sub, text), cites

    @staticmethod
    def _unmask_citations(text: str, cites: list) -> str:
        import re
        if not cites:
            return text
        return re.sub(r"(\d+)",
                      lambda m: cites[int(m.group(1))]
                      if int(m.group(1)) < len(cites) else m.group(0), text)

    @staticmethod
    def _allowed_keys(task: Task) -> list:
        """Citation keys that really exist in the retrieved references."""
        keys = []
        for s in task.task_card.get("sources", []) or []:
            if isinstance(s, dict) and s.get("key"):
                keys.append(s["key"])
        pq = task.task_card.get("paperqa_result", {}) or {}
        for c in (pq.get("citations") or []):
            if isinstance(c, dict) and c.get("key"):
                keys.append(c["key"])
        return keys

    async def _layer_7(self, task: Task, mem: TaskMemory):
        """٧: التحقق من التوثيق — PaperQA truth-check ثم strict-RAG صارم:
        يُسقط أي استشهاد مفتاحه غير موجود فعلاً في المراجع المسترجَعة."""
        from pipeline.layers.layer_7_verify import run as _layer7_run
        await _layer7_run(task, mem)

        # No-citation modes: the text must carry NO in-text citations. Strip any
        # the model produced anyway (none / uncited / sources-were-unavailable).
        card = task.task_card
        if task.draft and (card.get("sourcing_mode") in ("none", "uncited")
                           or card.get("sources_unavailable")):
            task.draft = self._strip_citations(task.draft)
            task.sections = [{**s, "body": self._strip_citations(s.get("body", ""))}
                             for s in (task.sections or [])]
            return

        allowed = self._allowed_keys(task)
        # نُطبّق strict-RAG فقط حين توجد مفاتيح فعلية — وإلا فقد نحذف كل شيء
        if task.draft and allowed:
            try:
                res = self._skill_call("weak_model_support", "weak_model_support",
                                       "enforce_strict_rag", task.draft, allowed)
                task.draft = res["text"]
                task.task_card["citations_removed"] = res.get("removed", [])
                if res.get("removed"):
                    mem.set_status(7, f"حُذف {len(res['removed'])} استشهاد مُختلَق")
            except Exception as e:
                mem.set_status(7, f"تحقق صارم (تخطّي: {e})")


    @staticmethod
    def _resolve_output_dir() -> str:
        """Where finished files are written. Priority:
        1) WEAVER_OUTPUT_DIR (explicit override),
        2) the phone's shared storage in a "Weaver Write" folder — on
           Termux/Android (~/storage/shared, /storage/emulated/0, /sdcard),
        3) the project's outputs/ folder (desktop / when storage isn't set up).
        The chosen directory is created if missing."""
        import os
        env = os.environ.get("WEAVER_OUTPUT_DIR", "").strip()
        if env:
            d = os.path.expanduser(env)
            try:
                os.makedirs(d, exist_ok=True)
                return d
            except OSError:
                pass
        for base in (os.path.expanduser("~/storage/shared"),
                     "/storage/emulated/0", "/sdcard"):
            if os.path.isdir(base):
                d = os.path.join(base, "Weaver Write")
                try:
                    os.makedirs(d, exist_ok=True)
                    return d
                except OSError:
                    continue
        root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        d = os.path.join(root, "outputs")
        os.makedirs(d, exist_ok=True)
        return d

    def _export_fallback(self, out_dir: str, safe: str, task: Task) -> str:
        """Always writes a REAL file to disk (Markdown) even when a format's
        library is missing — so an output always exists."""
        import os
        out = os.path.join(out_dir, safe + ".md")
        body = task.draft or ""
        if not body and task.sections:
            body = "\n\n".join(f"# {s.get('heading','')}\n{s.get('body','')}"
                               for s in task.sections)
        with open(out, "w", encoding="utf-8") as f:
            f.write(body or "(لا يوجد محتوى بعد — لم يُضبط مفتاح النموذج)")
        return out

    @staticmethod
    def _resolve_font(card: dict) -> str:
        """Resolve the document font through fonts-core (engines/fonts-core).
        Keeps the requested name (Office renders it) but validates it against
        the bundled families. Falls back to a sane per-language default."""
        import os as _os, sys as _sys
        lang = card.get("language", "ar")
        requested = card.get("font") or (
            "Kufyan Arabic" if lang == "ar" else "Times New Roman")
        try:
            fc = _os.path.abspath(_os.path.join(
                _os.path.dirname(__file__), "..", "engines", "fonts-core"))
            if fc not in _sys.path:
                _sys.path.insert(0, fc)
            import fonts as _fonts
            info = _fonts.resolve_named_font(requested)
            return info.get("requested") or requested
        except Exception:
            return requested

    def _maybe_chart(self, task: Task, out_dir: str):
        """Build a chart PNG via chart_builder when the task provides a chart
        spec, or when charts are requested and tabular data exists. Returns the
        image path or None. Degrades safely if matplotlib is missing."""
        import os as _os, re as _re
        card = task.task_card
        spec = card.get("chart")
        if not spec:
            extras = card.get("extras") or {}
            data = card.get("data")
            if extras.get("charts") and isinstance(data, list) and len(data) >= 2:
                labels, vals = [], []
                for r in data:
                    if isinstance(r, (list, tuple)) and len(r) >= 2:
                        try:
                            vals.append(float(r[1]))
                            labels.append(str(r[0]))
                        except (TypeError, ValueError):
                            labels, vals = [], []
                            break
                if labels and vals and len(labels) == len(vals):
                    spec = {"type": "bar",
                            "data": {"labels": labels, "values": vals},
                            "title": card.get("topic", "")}
        if not spec or not spec.get("data"):
            return None
        base = _re.sub(r"\W+", "", (card.get("topic") or "chart"))[:30] or "chart"
        png = _os.path.join(out_dir, "chart_" + base + ".png")
        try:
            res = self._skill_call(
                "chart_builder", "build_chart", "build_chart",
                spec.get("type", "bar"), spec["data"], png,
                title=spec.get("title", ""), lang=card.get("language", "ar"))
            if isinstance(res, dict) and res.get("ok") is False:
                return None
            return png if _os.path.exists(png) else None
        except Exception:
            return None

    def _export(self, task: Task) -> str:
        """Route to the right builder by output_format and WRITE the file to
        the resolved output directory (the phone's "Weaver Write" folder on
        Android; see _resolve_output_dir). No download links. Any builder
        failure (e.g. a missing library) degrades to a real Markdown file."""
        import os, re
        card = task.task_card
        lang = card.get("language", "ar")
        fmt = self._primary_format(card)
        out_dir = self._resolve_output_dir()
        topic = (card.get("topic") or task.description or "document").strip()
        safe = re.sub(r'[\\/:*?"<>|]+', "", topic)
        safe = re.sub(r"\s+", "_", safe)[:60] or "document"
        title = topic
        sections = task.sections or [{"heading": title, "body": task.draft}]
        references = (card.get("paperqa_result") or {}).get("references")
        font = self._resolve_font(card)

        # generate a chart when requested/derivable and append it as an image
        chart_png = self._maybe_chart(task, out_dir)
        if chart_png:
            sections = sections + [{
                "heading": ("الرسم البياني" if lang == "ar" else "Chart"),
                "body": "",
                "image": {"path": chart_png, "caption": card.get("topic", "")}}]
            card["chart_path"] = chart_png

        try:
            if fmt == "docx":
                out = os.path.join(out_dir, safe + ".docx")
                cover = (card.get("cover") if self._skill_call(
                    "docx_builder", "docx_frontmatter", "should_add_cover", card)
                    else None)
                toc_pos = self._skill_call(
                    "docx_builder", "docx_frontmatter", "resolve_toc_position",
                    card) or "after_cover"
                self._skill_call(
                    "docx_builder", "docx_advanced", "build_rich_docx",
                    title=title, sections=sections, output_path=out, lang=lang,
                    font=font, references=references, toc=bool(card.get("toc")),
                    cover=cover, toc_position=toc_pos)
                # rich Word styling for Quran/Hadith (bold verse/matn via the
                # quran_hadith_citation skill's own _set_run). Guarded/no-op.
                self._style_islamic_docx(out, card)
                return out
            if fmt == "pdf":
                out = os.path.join(out_dir, safe + ".pdf")
                self._skill_call("pdf_builder", "build_pdf", "build_pdf",
                                 sections=sections, output_path=out,
                                 title=title, lang=lang, references=references)
                return out
            if fmt == "pptx":
                out = os.path.join(out_dir, safe + ".pptx")
                # turn each section body into clean slide points (strip markdown
                # so the deck reads cleanly — build_deck renders text directly)
                def _slide_points(body):
                    pts = []
                    for ln in (body or "").split("\n"):
                        ln = ln.strip()
                        if not ln:
                            continue
                        ln = re.sub(r'^[-*•]\s*', '', ln)     # bullet markers
                        ln = re.sub(r'^#{1,6}\s*', '', ln)    # md headings
                        ln = re.sub(r'[*_`]+', '', ln).strip()  # inline md
                        if ln:
                            pts.append(ln)
                    return pts
                plan_sections = [{"title": s.get("heading", ""),
                                  "points": _slide_points(s.get("body", ""))}
                                 for s in sections]
                # ── creative path first ("the Claude way"): llm_deck_generator
                #    authors free HTML/CSS per slide, then the (now multi-slide)
                #    bridge converts to native PPTX. Needs llm_fn + the Node
                #    engine; on ANY miss we fall through to the template builder
                #    below — unchanged — so nothing breaks without them.
                if self.llm_fn:
                    try:
                        gen = self._skill_call(
                            "pptx_builder", "llm_deck_generator",
                            "generate_and_convert", title, plan_sections, out,
                            lang=lang, llm_fn=self.llm_fn)
                        if (gen and gen.get("ok") and os.path.exists(out)
                                and os.path.getsize(out) > 0):
                            return out
                    except Exception:
                        pass
                slides = None
                # slide_designer: build the slide PLAN (title/points/visual,
                # ≤5 points per slide, academic skeleton when empty). Its
                # "points" key is exactly what build_deck consumes. Additive —
                # on any miss we fall back to a direct mapping below.
                try:
                    plan = self._skill_call(
                        "slide_designer", "design_slides", "design_slides",
                        title, plan_sections, card.get("slide_count"),
                        lang, self.llm_fn)
                    if plan and plan.get("slides"):
                        slides = plan["slides"]
                except Exception:
                    slides = None
                if not slides:
                    # fallback with the CORRECT "points" key build_deck reads
                    # (the previous "bullets" key was silently ignored)
                    slides = [{"title": s["title"], "points": s["points"]}
                              for s in plan_sections]
                self._skill_call("pptx_builder", "build_pptx", "build_pptx",
                                 slides=slides, output_path=out, lang=lang,
                                 title=title)
                return out
            if fmt == "xlsx":
                out = os.path.join(out_dir, safe + ".xlsx")
                data = card.get("data") or [[s.get("heading", ""),
                                             s.get("body", "")]
                                            for s in sections]
                headers = card.get("headers") or (
                    ["القسم", "المحتوى"] if lang == "ar"
                    else ["Section", "Content"])
                self._skill_call("xlsx_builder", "build_xlsx", "build_xlsx",
                                 data=data, output_path=out, headers=headers,
                                 lang=lang)
                return out
            if fmt in ("csv",):
                out = os.path.join(out_dir, safe + ".csv")
                import csv as _csv
                data = card.get("data") or [[s.get("heading", ""),
                                             re.sub(r'\s+', ' ',
                                                    (s.get("body", "") or ""))]
                                            for s in sections]
                headers = card.get("headers") or (
                    ["القسم", "المحتوى"] if lang == "ar"
                    else ["Section", "Content"])
                # utf-8-sig so Excel opens Arabic correctly
                with open(out, "w", encoding="utf-8-sig", newline="") as f:
                    w = _csv.writer(f)
                    if headers:
                        w.writerow(headers)
                    for row in data:
                        w.writerow(row if isinstance(row, (list, tuple))
                                   else [row])
                return out
            if fmt in ("txt", "text"):
                out = os.path.join(out_dir, safe + ".txt")
                body = "\n\n".join(
                    ((s.get("heading", "") + "\n") if s.get("heading") else "")
                    + (s.get("body", "") or "") for s in sections).strip()
                body = re.sub(r'^[ \t]*#{1,6}[ \t]*', '', body, flags=re.M)
                body = re.sub(r'[*_`]+', '', body)
                with open(out, "w", encoding="utf-8") as f:
                    f.write(body)
                return out
            if fmt in ("html", "htm"):
                out = os.path.join(out_dir, safe + ".html")
                with open(out, "w", encoding="utf-8") as f:
                    f.write(self._sections_to_html(title, sections, lang))
                return out
        except Exception:
            return self._export_fallback(out_dir, safe, task)
        return self._export_fallback(out_dir, safe, task)

    @staticmethod
    def _sections_to_html(title, sections, lang="ar"):
        """A clean standalone HTML document from [{heading, body}] sections."""
        import html as _h
        import re as _r
        rtl = "rtl" if lang == "ar" else "ltr"
        parts = ["<!doctype html><html lang=\"" + ("ar" if lang == "ar" else "en")
                 + "\" dir=\"" + rtl + "\"><head><meta charset=\"utf-8\">"
                 "<meta name=\"viewport\" content=\"width=device-width,"
                 "initial-scale=1\"><title>" + _h.escape(title or "") + "</title>"
                 "<style>body{font-family:system-ui,'Segoe UI',Arial,sans-serif;"
                 "line-height:1.8;max-width:820px;margin:24px auto;padding:0 16px;"
                 "color:#1a1a1a}h1,h2,h3{line-height:1.3}a{color:#1a56db}"
                 "@media(prefers-color-scheme:dark){body{background:#111;color:#eee}"
                 "a{color:#7aa2ff}}</style></head><body>"]
        if title:
            parts.append("<h1>" + _h.escape(title) + "</h1>")
        for s in sections:
            hd = (s.get("heading") or "").strip()
            if hd:
                parts.append("<h2>" + _h.escape(hd) + "</h2>")
            for para in _r.split(r"\n\s*\n", (s.get("body") or "").strip()):
                para = para.strip()
                if not para:
                    continue
                lines = [ln.strip() for ln in para.split("\n") if ln.strip()]
                if lines and all(_r.match(r'^[-*•]\s+', ln) for ln in lines):
                    parts.append("<ul>" + "".join(
                        "<li>" + _h.escape(_r.sub(r'^[-*•]\s+', '', ln)) + "</li>"
                        for ln in lines) + "</ul>")
                else:
                    parts.append("<p>" + "<br>".join(_h.escape(ln)
                                                     for ln in lines) + "</p>")
        parts.append("</body></html>")
        return "".join(parts)

    def _source_note(self, task: Task):
        """Prepend a short, honest note when the document was written without
        external sources: either because the user asked for that ("none"), or
        because sources were required but none could be retrieved on the device
        ("cited" + sources_unavailable). The "uncited" mode gets no note — the
        user deliberately chose not to document sources."""
        card = task.task_card
        mode = card.get("sourcing_mode", "cited")
        lang = card.get("language", "ar")
        note = None
        if mode == "none":
            note = ("أُعدّ هذا المستند دون مصادر خارجية بناءً على طلبك."
                    if lang == "ar" else
                    "This document was prepared without external sources, as "
                    "requested.")
        elif mode != "uncited" and card.get("sources_unavailable"):
            note = ("تعذّر الوصول إلى مصادر خارجية أثناء الإعداد، فحُرّر المحتوى "
                    "من المعرفة العامة."
                    if lang == "ar" else
                    "External sources could not be retrieved, so the content was "
                    "written from general knowledge.")
        # a link was pasted but couldn't be read (e.g. missing library / no
        # captions) — say so honestly instead of silently topic-searching.
        if card.get("pasted_present") and not card.get("pasted_reads"):
            extra = ("تعذّرت قراءة الرابط المُدرَج (قد يلزم تثبيت "
                     "youtube-transcript-api أو أنّ المحتوى غير متاح)، فاعتمد "
                     "المحتوى على بحث عام حول الموضوع."
                     if lang == "ar" else
                     "The pasted link could not be read (a library may be "
                     "missing or the content is unavailable), so this is based "
                     "on a general search of the topic.")
            note = (note + " " + extra) if note else extra
        if not note:
            return
        head = "ملاحظة" if lang == "ar" else "Note"
        if task.sections and (task.sections[0].get("heading") or "") == head:
            return
        task.sections = [{"heading": head, "body": note}] + (task.sections or [])
        if task.draft:
            task.draft = note + "\n\n" + task.draft
        card["source_note"] = note

    @staticmethod
    def _is_ref_heading(h: str) -> bool:
        h = (h or "").strip().lower()
        return any(k in h for k in ("مراجع", "مصادر", "references", "works cited",
                                    "bibliography"))

    def _append_references(self, task: Task):
        """Build the full reference list from the retrieved sources via the
        citation-style skill (apa_formatter / mla_formatter) and put it as the
        LAST section of the report, replacing any placeholder references
        heading. No sources → nothing added."""
        card = task.task_card
        # no-citation modes never get a references list
        if card.get("sourcing_mode") in ("none", "uncited"):
            return
        sources = card.get("sources") or []
        pq_refs = (card.get("paperqa_result") or {}).get("references")
        if not sources and not pq_refs:
            return
        lang = card.get("language", "ar")
        style = str(card.get("citation_style", "APA")).upper()
        skill = "mla_formatter" if style == "MLA" else "apa_formatter"
        module = "format_mla" if style == "MLA" else "format_apa"
        try:
            refs = self._skill_call(skill, module, "build_bibliography",
                                    sources, lang, pq_refs)
        except Exception:
            # minimal fallback list if the skill can't be loaded
            lines = []
            for i, s in enumerate(sources, 1):
                if isinstance(s, dict):
                    lines.append(f"{i}. {s.get('title') or s.get('key') or ''} "
                                 f"{s.get('url','')}".strip())
            refs = "\n".join(lines)
            if pq_refs:
                refs = (refs + "\n" + str(pq_refs)).strip()
        if not (refs or "").strip():
            return
        head = "قائمة المراجع" if lang == "ar" else "References"
        # drop any earlier placeholder references section, then append the real one
        task.sections = [s for s in (task.sections or [])
                         if not self._is_ref_heading(s.get("heading", ""))]
        task.sections.append({"heading": head, "body": refs})
        # also reflect it at the end of the chat draft
        if task.draft:
            task.draft = task.draft.rstrip() + "\n\n" + head + "\n" + refs
        card["references_list"] = refs

    async def _layer_8(self, task: Task, mem: TaskMemory):
        """٨: الإخراج — كتابة الملف النهائي على القرص في outputs/."""
        task.status = TaskStatus.LAYER_8
        mem.set_status(8, "توليد الملف النهائي")
        # honest note when the document was written without external sources
        try:
            self._source_note(task)
        except Exception as e:
            mem.set_status(8, f"ملاحظة المصادر (تخطّي: {e})")
        # append the full reference list at the very end of the report
        try:
            self._append_references(task)
        except Exception as e:
            mem.set_status(8, f"قائمة المراجع (تخطّي: {e})")
        # إضافة تقرير التحقق للوثيقة النهائية
        try:
            from pipeline.layers.layer_7_verify import format_verification_report
            verify_text = format_verification_report(
                task.task_card, lang=task.task_card.get("language", "ar")
            )
            if verify_text:
                mem.add_reference(f"[تقرير التحقق]\n{verify_text}", source_key="layer_8")
        except Exception:
            pass
        # كتابة الملف الفعلي على القرص
        try:
            task.output_path = self._export(task)
            mem.set_status(8, f"أُخرج الملف: {task.output_path}")
        except Exception as e:
            mem.set_status(8, f"إخراج (تخطّي: {e})")

    # ── تشغيل متزامن لطلب واحد عبر خط الأنابيب الكامل ──

    def _result(self, task: Task) -> dict:
        """Shape one finished task into a reply dict for the chat / terminal."""
        card = task.task_card
        reply = card.get("reply") or task.draft or ""
        return {
            "reply": reply,
            "output_path": task.output_path,
            "topic": card.get("topic"),
            "task_type": card.get("task_type"),
            "language": card.get("language"),
            "output_format": card.get("output_format"),
            "tools": task.tools,
            "skills": task.skills,
            "status": getattr(task.status, "value", str(task.status)),
        }

    def _emit(self, kind: str, label: str = "", detail: str = ""):
        """Push a progress event to the optional progress callback (used by the
        streaming chat endpoint to show tool-use steps live and in order)."""
        cb = getattr(self, "_progress", None)
        if not cb:
            return
        try:
            cb({"t": kind, "label": label, "detail": detail})
        except Exception:
            pass

    async def run_once(self, description: str, input_files: list = None,
                       sandbox: bool = False, progress=None) -> dict:
        """Run ONE request through the full pipeline (layers 0→8) and return the
        reply + output file path. Used by the web chat and the terminal so every
        request goes through the whole system. Isolated: its own task memory,
        created and closed here. Sandbox is off by default (text chat needs no
        package installs); pass sandbox=True for tasks with input files.
        `progress(ev)` receives step events for a live tool-use timeline."""
        self._progress = progress
        ar = (self._detect_lang(description) == "ar")
        L = (lambda a, e: a if ar else e)  # localized label helper
        task = Task(description=description, input_files=input_files or [])
        task.started_at = time.time()
        mem = self.memory.create_task(task.task_id)
        sb = None
        try:
            if sandbox:
                try:
                    sb = await self.sandbox.create_for_task(task.task_id)
                except Exception:
                    sb = None

            # conduct guard (before Layer 0): stay professional under abuse
            try:
                g = self._skill_call("conduct_guard", "conduct_guard",
                                     "guard_response", description,
                                     self._detect_lang(description))
                task.task_card["conduct"] = g
                if g.get("hostile") and not g.get("do_task"):
                    task.task_card["reply"] = g.get("reply_prefix", "")
                    task.status = TaskStatus.COMPLETED
                    return self._result(task)
            except Exception:
                pass

            await self._layer_0(task, mem)
            try:
                await self._layer_1(task, mem, sb)
            except Exception as e:
                mem.set_status(1, f"بنية تحتية (تخطّي: {e})")
            try:
                await self._layer_2(task, mem)
            except Exception as e:
                mem.set_status(2, f"إدخال (تخطّي: {e})")

            self._emit("step", L("فهم الطلب", "Understanding the request"))
            await self._layer_3(task, mem)
            self._emit("detail", "",
                       L("الأدوات: ", "Tools: ") + ", ".join(task.tools or []))

            if "academic_search" in task.tools or task.task_card.get("needs_academic_search"):
                self._emit("step", L("بحث أكاديمي", "Academic search"))
            if "web_search" in task.tools:
                self._emit("step", L("بحث في الويب", "Searching the web"))
            await self._layer_4(task, mem)
            _nsrc = len(task.task_card.get("sources", []) or [])
            _nfull = task.task_card.get("web_full_reads", 0)
            if _nsrc:
                d = L(f"{_nsrc} مصدر", f"{_nsrc} sources")
                if _nfull:
                    d += L(f" (قراءة كاملة لـ {_nfull})", f" ({_nfull} read in full)")
                self._emit("detail", "", d)

            self._emit("step", L("فحص مصداقية المصادر", "Checking source credibility"))
            await self._layer_5(task, mem)

            self._emit("step", L("كتابة المحتوى", "Writing the content"))
            await self._layer_6(task, mem)
            self._emit("detail", "",
                       L(f"{len(task.sections or [])} قسم",
                         f"{len(task.sections or [])} sections"))

            await self._layer_6_6(task, mem)

            self._emit("step", L("تنظيف وأنسنة النص", "Cleaning up the text"))
            await self._layer_6_5(task, mem)

            self._emit("step", L("التحقق من التوثيق", "Verifying citations"))
            await self._layer_7(task, mem)

            self._emit("step", L("توليد الملف", "Generating the file"))
            await self._layer_8(task, mem)
            if task.output_path:
                self._emit("detail", "", task.output_path)

            task.status = TaskStatus.COMPLETED
            task.completed_at = time.time()
            return self._result(task)
        finally:
            self._progress = None
            try:
                if sb is not None:
                    await self.sandbox.destroy(task.task_id)
            except Exception:
                pass
            self.memory.close_task(task.task_id)

    # ── حالة النظام ──

    def status(self) -> dict:
        """حالة النظام الكاملة."""
        return {
            "active_tasks": len(self._active),
            "queued_tasks": len(self._queue),
            "completed_tasks": len(self._completed),
            "max_parallel": MAX_TASKS,
            "slots_available": MAX_TASKS - len(self._active),
            "tasks": {
                tid: {
                    "description": t.description[:40],
                    "status": t.status.value,
                    "elapsed": f"{t.elapsed():.0f}ث",
                }
                for tid, t in self._active.items()
            },
        }

    async def shutdown(self):
        """إيقاف نظيف للنظام."""
        await self.sandbox.destroy_all()
        self.memory.close_all()


import os  # needed for layer_2


# ── intent: is this a document/generation task, or a quick question? ──
# A quick question is answered directly by the model (fast, no file). A task
# with a creation intent (write/report/presentation/analysis/export …) goes
# through the full pipeline. Bilingual triggers; substring match.
_TASK_TRIGGERS = (
    # Arabic — creation verbs
    "اكتب", "أكتب", "اكتبي", "اعمل", "أعمل", "اصنع", "أنشئ", "انشئ", "صمم",
    "صمّم", "جهّز", "جهز", "حضّر", "حضر", "ولّد", "ولد", "أخرج", "اخرج", "لخّص",
    "لخص", "حلّل", "حلل",
    # Arabic — document nouns
    "بحث", "بحثاً", "مقال", "مقالة", "تقرير", "دراسة", "أطروحة", "رسالة علمية",
    "عرض", "بوربوينت", "شرائح", "ملف", "مستند", "وثيقة", "صفحة", "صفحات",
    "مراجع", "مرجع", "استشهاد", "جدول", "جداول", "رسم بياني", "مخطط", "واجب",
    "ملخص", "خطة", "سيرة ذاتية", "تحليل بيانات",
    # English — creation verbs + document nouns
    "write", "create", "generate", "make ", "design", "draft", "compose",
    "essay", "report", "article", "research", "paper", "presentation",
    "slides", "powerpoint", "deck", "document", "docx", "pptx", "xlsx", "pdf",
    "references", "citation", "table", "chart", "analyze data", "analyse data",
    "thesis", "dissertation", "summariz", "summaris", "outline", "resume",
    "cv ", "assignment",
)


def _is_conversation_recap(text: str) -> bool:
    """True when the user is asking to recap/summarize THIS conversation (not to
    generate a document). Such a request belongs on the quick path, which sees
    the full recent history (~30 turns) and answers directly — not the pipeline,
    which only receives the last few turns and would also save a stray file.
    Requires BOTH a summary intent AND a reference to the conversation, and no
    pasted link or explicit file-format request (those mean something else)."""
    t = (text or "").lower()
    summary_words = ("لخّص", "لخص", "ملخص", "ملخّص", "لخصلي", "لخّصلي", "اختصر",
                     "خلاصة", "summar", "recap", "tl;dr", "tldr", "استعرض ما")
    convo_words = ("المحادثة", "محادثتنا", "محادثه", "الحوار", "النقاش", "الشات",
                   "الدردشة", "ما قمنا", "ما فعلنا", "ما تحدثنا", "ما دار",
                   "ما جرى", "كلامنا", "حديثنا", "ما اتفقنا", "conversation",
                   "this chat", "our chat", "what we did", "this session",
                   "our discussion", "this thread")
    doc_words = ("ملف", "وورد", "word", "pdf", "docx", "مستند", "وثيقة",
                 "بوربوينت", "pptx", "عرض تقديمي", "اكسل", "excel", "تقرير")
    has_url = ("http://" in t) or ("https://" in t)
    return (any(w in t for w in summary_words)
            and any(w in t for w in convo_words)
            and not any(w in t for w in doc_words)
            and not has_url)


def is_document_task(text: str) -> bool:
    """True when the message asks to produce/analyse a document (→ full
    pipeline); False for a quick conversational question (→ direct answer)."""
    t = (text or "").lower()
    # A recap of THIS conversation is a quick conversational answer, not a
    # document to generate → keep it on the fast path (full recent history, no
    # stray file). This overrides the "summarize" trigger below.
    if _is_conversation_recap(text):
        return False
    if any(trig in t for trig in _TASK_TRIGGERS):
        return True
    # A pasted YouTube link is ALWAYS a pipeline task: it must reach the
    # dedicated youtube path in _layer_3 (summary / transcript / timing).
    # Otherwise a "فرّغ ..." request (whose verbs are not task-triggers) falls
    # to the quick-chat path — which has no transcript tool — and returns wrong
    # output or an empty provider reply. Lazy + guarded; non-YouTube text is
    # completely unaffected.
    try:
        import re as _re
        from capabilities.tools import tool_youtube as _yt
        for _u in _re.findall(r"https?://\S+", text or ""):
            if _yt.is_youtube_url(_u.rstrip('.,)"\'،؛')):
                return True
    except Exception:
        pass
    return False


# ── synchronous entry point (used by web/server.py and weaver.py) ──
_SHARED_ORCH = None

import threading as _threading  # noqa: E402
import heapq as _heapq  # noqa: E402
import itertools as _itertools  # noqa: E402


class _PriorityGate:
    """A concurrency gate of `limit` slots with a PRIORITY wait queue: when all
    slots are busy, waiting callers are admitted highest-priority-first (FIFO on
    ties) as slots free — not in arrival order. Thread-safe; used to gate the
    sync pipeline across all request threads."""

    def __init__(self, limit):
        self.limit = limit
        self._lock = _threading.Lock()
        self._running = 0
        self._heap = []                      # (-priority, seq, Event)
        self._seq = _itertools.count()

    def acquire(self, priority: int = 0):
        with self._lock:
            if self._running < self.limit:
                self._running += 1
                return
            ev = _threading.Event()
            _heapq.heappush(self._heap, (-int(priority), next(self._seq), ev))
        ev.wait()   # a releaser hands us the slot (running already counts us)

    def release(self):
        with self._lock:
            if self._heap:
                _, _, ev = _heapq.heappop(self._heap)  # highest priority next
                ev.set()                                # slot handed over
            else:
                self._running = max(0, self._running - 1)

    def free_slots(self):
        with self._lock:
            return max(0, self.limit - self._running)

    def waiting(self):
        with self._lock:
            return len(self._heap)


# At most MAX_TASKS pipelines run at once across all request threads; extras
# wait in a priority queue (highest priority admitted first).
_PIPELINE_GATE = _PriorityGate(MAX_TASKS)

# words that bump a request's priority (so "عاجل …" jumps the queue)
_URGENT_WORDS = ("عاجل", "مستعجل", "أولوية عالية", "urgent", "asap",
                 "high priority", "بسرعة")


def task_priority(text: str) -> int:
    """Priority for a request: higher = admitted sooner when the 5 slots are
    full. Bumped by urgent keywords; default 0."""
    t = (text or "").lower()
    return 10 if any(w in t for w in _URGENT_WORDS) else 0


def pipeline_slots():
    """Free parallel slots right now (for diagnostics)."""
    return _PIPELINE_GATE.free_slots()


def _md_to_sections(md):
    """Split ready markdown into [{heading, body}] by its '#'/'##' headings."""
    import re
    secs, cur_h, cur_b = [], "", []
    for ln in (md or "").split("\n"):
        m = re.match(r"^\s{0,3}#{1,6}\s+(.*)$", ln)
        if m:
            if cur_h or "".join(cur_b).strip():
                secs.append({"heading": cur_h,
                             "body": "\n".join(cur_b).strip()})
            cur_h, cur_b = m.group(1).strip(), []
        else:
            cur_b.append(ln)
    if cur_h or "".join(cur_b).strip():
        secs.append({"heading": cur_h, "body": "\n".join(cur_b).strip()})
    secs = [s for s in secs if s.get("heading") or s.get("body")]
    return secs or [{"heading": "", "body": (md or "").strip()}]


def _derive_title(md, fallback="مستند"):
    """A document title from the content's first heading, else its first line."""
    import re
    for ln in (md or "").split("\n"):
        m = re.match(r"^\s{0,3}#{1,6}\s+(.*)$", ln)
        if m and m.group(1).strip():
            return m.group(1).strip()[:80]
    for ln in (md or "").split("\n"):
        s = ln.strip().lstrip("#").strip()
        if s:
            return s[:80]
    return fallback


def _content_to_slides(llm_fn, content, lang="ar"):
    """Ask the model to reshape ready content into presentation slides. Returns
    [{heading, body}] (one per slide, body = bullet lines) or None."""
    if not llm_fn:
        return None
    try:
        from core.llm import extract_json
        prompt = (
            "حوّل المحتوى التالي إلى شرائح عرض تقديمي واضحة. أعِد JSON فقط: "
            "{\"slides\":[{\"title\":\"عنوان قصير\",\"bullets\":[\"نقطة\",...]}]}"
            " — عناوين موجزة و3 إلى 6 نقاط قصيرة لكل شريحة، دون أي شرح خارج JSON:\n\n"
            if lang == "ar" else
            "Turn the following into clear presentation slides. Return JSON only: "
            "{\"slides\":[{\"title\":\"short\",\"bullets\":[\"point\",...]}]} — "
            "concise titles, 3-6 short bullets each, nothing outside JSON:\n\n"
        ) + (content or "")[:9000]
        data = extract_json(llm_fn(prompt, temperature=0.3)) or {}
        slides = data.get("slides") or []
        out = []
        for s in slides:
            if not isinstance(s, dict):
                continue
            bl = [str(b).strip() for b in (s.get("bullets") or []) if str(b).strip()]
            out.append({"heading": str(s.get("title", "")).strip(),
                        "body": "\n".join("- " + b for b in bl)})
        return out or None
    except Exception:
        return None


def _content_to_table(llm_fn, content, lang="ar"):
    """Ask the model to extract a table from content. Returns
    {"headers":[...], "rows":[[...],...]} or None."""
    if not llm_fn:
        return None
    try:
        from core.llm import extract_json
        prompt = (
            "استخرج من المحتوى التالي جدولاً منظّماً. أعِد JSON فقط: "
            "{\"headers\":[\"عمود1\",\"عمود2\",...],\"rows\":[[\"..\",\"..\"],...]}"
            " — إن كان المحتوى نقاطاً، اجعل عمودين: \"النقطة\" و\"التفصيل\". دون "
            "أي نص خارج JSON:\n\n"
            if lang == "ar" else
            "Extract a structured table from the content. Return JSON only: "
            "{\"headers\":[...],\"rows\":[[...],...]} — if it's bullet points use "
            "two columns \"Point\" and \"Detail\". Nothing outside JSON:\n\n"
        ) + (content or "")[:9000]
        data = extract_json(llm_fn(prompt, temperature=0.2)) or {}
        headers = data.get("headers") or []
        rows = data.get("rows") or []
        rows = [[str(c) for c in r] for r in rows if isinstance(r, list) and r]
        if headers and rows:
            return {"headers": [str(h) for h in headers], "rows": rows}
        return None
    except Exception:
        return None


def export_content_to_file(content, fmt="docx", lang="ar", title=None):
    """Export READY markdown content to a file (docx/pdf/pptx/xlsx) directly, with
    NO research pipeline — for "أخرج/حوّل الناتج السابق إلى وورد/pdf". The filename
    comes from the content's own title, never the command. Returns the output path
    or None. Safe/degrading."""
    import tempfile
    content = (content or "").strip()
    if not content:
        return None
    _PIPELINE_GATE.acquire(0)
    fd, db = tempfile.mkstemp(prefix="weaver_exp_", suffix=".db")
    os.close(fd)
    try:
        orch = WeaverOrchestrator(db_path=db)
        fmtU = str(fmt).upper()
        ttl = title or _derive_title(content)
        task = Task(description=ttl)
        task.draft = content
        task.task_card = {
            "topic": ttl, "language": lang,
            "output_format": [fmtU], "sourcing_mode": "none",
        }
        # smart restructuring so PowerPoint/Excel aren't just a text dump
        if fmtU == "PPTX":
            task.sections = (_content_to_slides(getattr(orch, "llm_fn", None),
                                                content, lang)
                             or _md_to_sections(content))
        elif fmtU in ("XLSX", "CSV"):
            tbl = _content_to_table(getattr(orch, "llm_fn", None), content, lang)
            if tbl:
                task.task_card["headers"] = tbl.get("headers")
                task.task_card["data"] = tbl.get("rows")
            task.sections = _md_to_sections(content)
        else:
            task.sections = _md_to_sections(content)
        try:
            return orch._export(task)
        except Exception:
            return None
    finally:
        try:
            os.remove(db)
        except OSError:
            pass
        _PIPELINE_GATE.release()


def quick_live_context_ex(msg, lang="ar", max_chars=6000):
    """Live context for the QUICK/chat path (used by web + terminal): read any
    URL pasted in the message (a page or a YouTube video) and — for news/recency
    questions — run a quick multi-engine web search. Returns a tuple
    (context_str, sources) where `sources` is an ORDERED list of {"title","url"}
    (newest first) gathered from the SAME search, so the caller can list clickable
    source links under a news answer without searching again. Fully synchronous
    and degrading (returns ("", []) on any failure)."""
    import asyncio
    import re
    parts = []
    sources = []
    try:
        urls = re.findall(r'https?://[^\s)>\]\"\'،]+', msg or "")
    except Exception:
        urls = []
    orch = object.__new__(WeaverOrchestrator)   # bare: only _extract_full used
    try:
        orch._yt_lang = lang
    except Exception:
        pass
    # 1) pasted URLs → read them (page / YouTube transcript)
    for u in (urls or [])[:2]:
        u = u.rstrip('.,)"،')
        try:
            txt = asyncio.run(orch._extract_full(u))
        except Exception:
            txt = None
        if txt and txt.strip():
            parts.append(f"[محتوى الرابط: {u}]\n{txt.strip()[:4000]}")
    # 2) news/recency OR an explicit site/date search (no URL pasted) → search
    _site, _df_dir = WeaverOrchestrator._search_directives(msg)
    _recency = WeaverOrchestrator._is_recency_query(msg)
    if not urls and (_recency or _site or _df_dir):
        try:
            q = (WeaverOrchestrator._augment_query_with_date(msg, lang)
                 if _recency else (msg or "").strip())
            if _site:
                q = q + " site:" + _site
            _df = _df_dir or ("w" if _recency else None)
            results = WeaverOrchestrator._multi_engine_search(
                q, lang, 12, df=_df) or []
            if _recency:
                results = WeaverOrchestrator._sort_results_by_recency(results)
            lines, n = [], 1
            for r in results[:12]:
                title = (r.get("title") or "").strip()
                if not title:
                    continue
                url = (r.get("url") or "").strip()
                snip = (r.get("content") or "").strip()[:220]
                # NUMBER each result so the model can cite [n] → the matching
                # source; sources are collected in the SAME order.
                lines.append(f"[{n}] {title} — {snip} ({url})")
                if url:
                    sources.append({"title": title, "url": url})
                n += 1
            if lines:
                _hdr = ("[نتائج بحث حيّة مرقّمة، الأحدث أولاً]" if _recency
                        else ("[نتائج بحث حيّة مرقّمة من موقع " + _site + "]"
                              if _site else "[نتائج بحث حيّة مرقّمة]"))
                _rule = ("\nقاعدة: لا تذكر خبراً/معلومة إلا إن كانت مدعومة "
                         "بأحد المصادر المرقّمة أعلاه، واذكر رقم مصدرها [n] بعدها، "
                         "واكتفِ بمصدر واحد لكل خبر إلا إن ورد فعلاً في أكثر من "
                         "مصدر. لا تنسب خبراً لمصدر لا يحتويه.")
                parts.append(_hdr + "\n" + "\n".join(lines) + _rule)
        except Exception:
            pass
    ctx = "\n\n".join(parts).strip()
    return ctx[:max_chars], sources


def quick_live_context(msg, lang="ar", max_chars=6000):
    """Backward-compatible wrapper: returns only the context string."""
    try:
        return quick_live_context_ex(msg, lang, max_chars)[0]
    except Exception:
        return ""


def run_pipeline_sync(description: str, input_files: list = None,
                      llm_fn=None, progress=None, priority: int = 0) -> dict:
    """Run one request through the full pipeline and return the reply dict.
    Safe to call from a synchronous context (a threaded HTTP handler, or the
    CLI): it spins its own event loop and its own isolated task memory. Each
    call builds the LLM client fresh from config/.env, so a key added at
    runtime is picked up without a restart. `progress(ev)` streams step
    events. At most MAX_TASKS (5) run concurrently; extras wait in a PRIORITY
    queue — higher `priority` is admitted first."""
    import asyncio
    import tempfile

    _PIPELINE_GATE.acquire(priority)
    try:
        fd, db = tempfile.mkstemp(prefix="weaver_", suffix=".db")
        os.close(fd)
        orch = WeaverOrchestrator(db_path=db, llm_fn=llm_fn)
        try:
            return asyncio.run(orch.run_once(description, input_files,
                                             sandbox=bool(input_files),
                                             progress=progress))
        finally:
            try:
                os.remove(db)
            except OSError:
                pass
    finally:
        _PIPELINE_GATE.release()
